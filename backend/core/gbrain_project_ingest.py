from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

import yaml

from core.extractor_classifier import ExtractorClassification, classify_source_file
from core.email_structured_extraction import (
    EmailStructuredExtractionResult,
    extract_email_attachments,
    extract_email_structured_markdown,
)
from core.gbrain import (
    GBrainSettings,
    ensure_project_gbrain_environment,
    load_gbrain_settings,
    project_source_id_for_workspace,
    project_source_paths_for_workspace,
)
from core.gbrain_ingest import (
    DOCX_EXTENSIONS,
    MEDIA_EXTENSIONS,
    MEDIA_TRANSCRIPTION_REQUIRED,
    PDF_EXTENSIONS,
    PDF_STRUCTURED_EXTRACTION_REQUIRED,
    TEXT_EXTENSIONS,
    _commit_derived_changes,
    _env_bool,
    _markdown_table,
    _pdf_image_sidecar_dirs,
    _relative_posix,
    _safe_filename,
    _sha256_file,
    _split_frontmatter,
    _write_markdown,
)
from core.meeting_structured_extraction import (
    MeetingStructuredExtractionResult,
    extract_meeting_structured_markdown,
    find_transcript_sidecar,
    find_transcript_sidecars_for_media_files,
)
from core.media_transcription import MediaTranscriptionResult, transcribe_media_to_markdown
from core.pdf_structured_extraction import PDFStructuredExtractionResult, extract_pdf_structured_markdown
from core.image_structured_extraction import ImageStructuredExtractionResult, extract_image_structured_markdown


PROJECT_DIR_CATEGORY_MAP = {
    "01-合同与报价": "contracts",
    "02-图纸与技术资料": "technical",
    "03-会议纪要": "meetings",
    "04-变更与签证": "changes",
    "05-生产与发货": "production",
    "06-现场与客诉": "site",
    "99-未归档文件": "unfiled",
}
PROJECT_RUNTIME_DIRS = {"derived", "manifests", ".trash", ".git", "__pycache__"}
PROJECT_INGEST_MANIFEST_NAME = "project-source-ingest-manifest.json"
EMAIL_ATTACHMENT_DIR_SUFFIX = ".attachments"


@dataclass(frozen=True)
class ProjectCompiledSource:
    source_path: Path
    status: str
    target_path: Path | None = None
    error: str | None = None
    source_sha256: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


PDFExtractor = Callable[..., PDFStructuredExtractionResult]
MeetingExtractor = Callable[..., MeetingStructuredExtractionResult]
MediaTranscriber = Callable[..., MediaTranscriptionResult]
EmailExtractor = Callable[..., EmailStructuredExtractionResult]
ImageExtractor = Callable[..., ImageStructuredExtractionResult]


def compile_project_workspace_sources(
    workspace: Any,
    settings: GBrainSettings | None = None,
    *,
    pdf_extractor: PDFExtractor | None = None,
    meeting_extractor: MeetingExtractor | None = None,
    media_transcriber: MediaTranscriber | None = None,
    email_extractor: EmailExtractor | None = None,
    image_extractor: ImageExtractor | None = None,
    enable_pdf_structured_extraction: bool | None = None,
    enable_media_transcription: bool | None = None,
    enable_email_extraction: bool | None = None,
    enable_image_extraction: bool | None = None,
    enable_email_attachment_recursion: bool | None = None,
) -> dict[str, Any]:
    settings = settings or load_gbrain_settings()
    environment = ensure_project_gbrain_environment(workspace, settings)
    paths = project_source_paths_for_workspace(workspace)
    source_id = project_source_id_for_workspace(workspace)
    started_at = _utc_now()
    pdf_enabled = (
        _env_bool("GBRAIN_PROJECT_PDF_STRUCTURED_EXTRACTION_ENABLED", True)
        if enable_pdf_structured_extraction is None
        else enable_pdf_structured_extraction
    )
    media_transcription_enabled = (
        _env_bool("GBRAIN_PROJECT_MEDIA_TRANSCRIPTION_ENABLED", True)
        if enable_media_transcription is None
        else enable_media_transcription
    )
    email_extraction_enabled = (
        _env_bool("GBRAIN_PROJECT_EMAIL_EXTRACTION_ENABLED", True)
        if enable_email_extraction is None
        else enable_email_extraction
    )
    image_extraction_enabled = (
        _env_bool("GBRAIN_PROJECT_IMAGE_EXTRACTION_ENABLED", True)
        if enable_image_extraction is None
        else enable_image_extraction
    )
    email_attachment_recursion_enabled = (
        _env_bool("GBRAIN_EMAIL_ATTACHMENT_RECURSION_ENABLED", True)
        if enable_email_attachment_recursion is None
        else enable_email_attachment_recursion
    )

    results: list[ProjectCompiledSource] = []
    processed: set[Path] = set()
    while True:
        source_files = [path for path in _iter_project_source_files(paths["root"], paths) if path.resolve() not in processed]
        if not source_files:
            break
        for source_path in source_files:
            processed.add(source_path.resolve())
            results.append(_compile_project_source(
            source_path,
            workspace,
            paths,
            started_at,
            pdf_extractor=pdf_extractor,
            meeting_extractor=meeting_extractor,
            media_transcriber=media_transcriber,
            email_extractor=email_extractor,
            image_extractor=image_extractor,
            enable_pdf_structured_extraction=pdf_enabled,
            enable_media_transcription=media_transcription_enabled,
            enable_email_extraction=email_extraction_enabled,
            enable_image_extraction=image_extraction_enabled,
            enable_email_attachment_recursion=email_attachment_recursion_enabled,
            ))
    manifest = {
        "schema_version": 1,
        "source_id": source_id,
        "workspace_id": getattr(workspace, "id", None),
        "workspace_name": getattr(workspace, "name", ""),
        "workspace_slug": getattr(workspace, "slug", ""),
        "brand": getattr(workspace, "brand", ""),
        "source_root": str(paths["root"].resolve()),
        "derived_path": str(paths["derived"].resolve()),
        "started_at": started_at,
        "finished_at": _utc_now(),
        "environment_ok": environment["ok"],
        "items": [_result_to_project_manifest_item(result, paths["root"], paths["derived"]) for result in results],
        "summary": {
            "total": len(results),
            "compiled": sum(1 for result in results if result.status == "compiled"),
            "pending_extractor_capability": sum(
                1 for result in results if result.status == "pending_extractor_capability"
            ),
            "pending_transcription": sum(1 for result in results if result.status == "pending_transcription"),
            "skipped": sum(1 for result in results if result.status == "skipped"),
            "failed": sum(1 for result in results if result.status == "failed"),
        },
    }
    manifest_path = paths["manifests"] / PROJECT_INGEST_MANIFEST_NAME
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    git_status = _commit_derived_changes(paths["derived"], manifest["summary"], settings.local_git_enabled)
    manifest["local_git"] = git_status
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def _compile_project_source(
    source_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    *,
    pdf_extractor: PDFExtractor | None = None,
    meeting_extractor: MeetingExtractor | None = None,
    media_transcriber: MediaTranscriber | None = None,
    email_extractor: EmailExtractor | None = None,
    image_extractor: ImageExtractor | None = None,
    enable_pdf_structured_extraction: bool = False,
    enable_media_transcription: bool = False,
    enable_email_extraction: bool = False,
    enable_image_extraction: bool = False,
    enable_email_attachment_recursion: bool = False,
) -> ProjectCompiledSource:
    suffix = source_path.suffix.lower()
    try:
        classification = classify_source_file(source_path, source_scope="project")
        classifier_metadata = classification.to_manifest_metadata()
        source_hash = _sha256_file(source_path)
        target_path = _project_target_path(source_path, paths, _project_category(source_path, paths["root"]))
        if suffix in TEXT_EXTENSIONS:
            _compile_project_text_source(
                source_path,
                target_path,
                workspace,
                paths,
                ingested_at,
                source_hash,
                classification,
            )
            return ProjectCompiledSource(
                source_path,
                "compiled",
                target_path,
                source_sha256=source_hash,
                metadata=classifier_metadata,
            )
        if suffix in DOCX_EXTENSIONS:
            _compile_project_docx_source(
                source_path,
                target_path,
                workspace,
                paths,
                ingested_at,
                source_hash,
                classification,
            )
            return ProjectCompiledSource(
                source_path,
                "compiled",
                target_path,
                source_sha256=source_hash,
                metadata=classifier_metadata,
            )
        if suffix in PDF_EXTENSIONS:
            if classification.extraction_complexity == "simple_text":
                _compile_project_pdf_text_source(
                    source_path,
                    target_path,
                    workspace,
                    paths,
                    ingested_at,
                    source_hash,
                    classification,
                )
                return ProjectCompiledSource(
                    source_path,
                    "compiled",
                    target_path,
                    source_sha256=source_hash,
                    metadata={**classifier_metadata, "extraction_status": "pdf_text_extracted"},
                )
            if not enable_pdf_structured_extraction:
                if target_path.exists():
                    target_path.unlink()
                return ProjectCompiledSource(
                    source_path,
                    "pending_extractor_capability",
                    error=PDF_STRUCTURED_EXTRACTION_REQUIRED,
                    source_sha256=source_hash,
                    metadata={
                        **classifier_metadata,
                        "extraction_status": "pending_extractor_capability",
                    },
                )
            result = (pdf_extractor or extract_pdf_structured_markdown)(source_path)
            approved_target_path = target_path
            _compile_project_pdf_structured_source(
                source_path,
                approved_target_path,
                workspace,
                paths,
                ingested_at,
                source_hash,
                result,
                classification,
            )
            return ProjectCompiledSource(
                source_path,
                "compiled",
                approved_target_path,
                source_sha256=source_hash,
                metadata={
                    **classifier_metadata,
                    "extraction_status": result.extraction_status,
                    "review_status": "approved",
                    "source_scope_review_policy": "project_no_admin_review",
                    "extractor_review_status": result.review_status,
                    "extractor": result.extractor,
                    "language_policy": result.language_policy,
                    "page_count": result.page_count,
                    "pages_analyzed": result.pages_analyzed,
                    "model_profile": result.model_profile,
                    "provider": result.provider,
                    "model": result.model,
                    "token_usage": result.token_usage,
                    "warnings": list(result.warnings),
                    "vision_pages": list(result.vision_pages),
                    "vision_image_count": result.vision_image_count,
                },
            )
        if suffix in MEDIA_EXTENSIONS:
            transcript_path = find_transcript_sidecar(source_path)
            transcription_result: MediaTranscriptionResult | None = None
            generated_transcript_path: Path | None = None
            if transcript_path is None and enable_media_transcription:
                transcription_result = (media_transcriber or transcribe_media_to_markdown)(source_path)
                generated_transcript_path = _auto_transcript_path(source_path)
                generated_transcript_path.write_text(transcription_result.transcript_text, encoding="utf-8")
                transcript_path = generated_transcript_path
            if transcript_path is not None:
                result = (meeting_extractor or extract_meeting_structured_markdown)(
                    title=source_path.stem,
                    transcript_path=transcript_path,
                    source_media_path=source_path,
                    source_label=_relative_posix(source_path, paths["root"]),
                )
                approved_target_path = target_path
                _compile_project_meeting_structured_source(
                    source_path,
                    transcript_path,
                    approved_target_path,
                    workspace,
                    paths,
                    ingested_at,
                    source_hash,
                    result,
                    classification,
                    transcription_result=transcription_result,
                    generated_transcript_path=generated_transcript_path,
                )
                return ProjectCompiledSource(
                    source_path,
                    "compiled",
                    approved_target_path,
                    source_sha256=source_hash,
                    metadata={
                        **classifier_metadata,
                        "extraction_status": result.extraction_status,
                        "review_status": "approved",
                        "source_scope_review_policy": "project_no_admin_review",
                        "extractor_review_status": result.review_status,
                        "extractor": result.extractor,
                        "language_policy": result.language_policy,
                        "transcription_status": (
                            transcription_result.transcription_status
                            if transcription_result is not None
                            else result.transcription_status
                        ),
                        "transcript_file": _relative_posix(transcript_path, paths["root"]),
                        "transcription_extractor": (
                            transcription_result.extractor if transcription_result is not None else None
                        ),
                        "transcription_model_profile": (
                            transcription_result.model_profile if transcription_result is not None else None
                        ),
                        "transcription_provider": (
                            transcription_result.provider if transcription_result is not None else None
                        ),
                        "transcription_model": transcription_result.model if transcription_result is not None else None,
                        "transcription_token_usage": (
                            transcription_result.token_usage if transcription_result is not None else None
                        ),
                        "transcription_segment_count": (
                            transcription_result.segment_count if transcription_result is not None else None
                        ),
                        "transcript_refinement_status": (
                            transcription_result.refinement_status if transcription_result is not None else None
                        ),
                        "transcript_refinement_model_profile": (
                            transcription_result.refinement_model_profile if transcription_result is not None else None
                        ),
                        "transcript_refinement_provider": (
                            transcription_result.refinement_provider if transcription_result is not None else None
                        ),
                        "transcript_refinement_model": (
                            transcription_result.refinement_model if transcription_result is not None else None
                        ),
                        "transcript_refinement_token_usage": (
                            transcription_result.refinement_token_usage if transcription_result is not None else None
                        ),
                        "transcript_terminology": (
                            list(transcription_result.terminology) if transcription_result is not None else None
                        ),
                        "generated_transcript_file": (
                            _relative_posix(generated_transcript_path, paths["root"])
                            if generated_transcript_path is not None
                            else None
                        ),
                        "segment_count": result.segment_count,
                        "action_item_count": result.action_item_count,
                        "decision_count": result.decision_count,
                        "risk_count": result.risk_count,
                        "warnings": list(result.warnings),
                    },
                )
            return ProjectCompiledSource(
                source_path,
                "pending_transcription",
                error=MEDIA_TRANSCRIPTION_REQUIRED,
                source_sha256=source_hash,
                metadata={
                    **classifier_metadata,
                    "extraction_status": "pending_meeting_transcription",
                    "transcription_status": "pending_transcription",
                },
            )
        if classification.file_kind == "email":
            if suffix != ".eml" or not enable_email_extraction:
                return ProjectCompiledSource(
                    source_path,
                    "pending_extractor_capability",
                    error=classification.classifier_reason,
                    source_sha256=source_hash,
                    metadata={
                        **classifier_metadata,
                        "extraction_status": "pending_extractor_capability",
                    },
                )
            result = (email_extractor or extract_email_structured_markdown)(source_path)
            extracted_attachments = ()
            if enable_email_attachment_recursion:
                extracted_attachments = extract_email_attachments(source_path, _email_attachment_dir(source_path))
            _compile_project_email_structured_source(
                source_path,
                target_path,
                workspace,
                paths,
                ingested_at,
                source_hash,
                result,
                classification,
            )
            return ProjectCompiledSource(
                source_path,
                "compiled",
                target_path,
                source_sha256=source_hash,
                metadata={
                    **classifier_metadata,
                    "extraction_status": result.extraction_status,
                    "review_status": "approved",
                    "source_scope_review_policy": "project_no_admin_review",
                    "extractor": result.extractor,
                    "language_policy": result.language_policy,
                    "email_subject": result.subject,
                    "email_sender": result.sender,
                    "email_recipients": list(result.recipients),
                    "email_message_date": result.message_date,
                    "email_attachments": list(result.attachment_names),
                    "model_profile": result.model_profile,
                    "provider": result.provider,
                    "model": result.model,
                    "token_usage": result.token_usage,
                    "email_extracted_attachment_files": [
                        _relative_posix(attachment.path, paths["root"]) for attachment in extracted_attachments
                    ],
                    "warnings": list(result.warnings),
                },
            )
        if classification.file_kind == "image":
            if not enable_image_extraction:
                return ProjectCompiledSource(
                    source_path,
                    "pending_extractor_capability",
                    error=classification.classifier_reason,
                    source_sha256=source_hash,
                    metadata={
                        **classifier_metadata,
                        "extraction_status": "pending_extractor_capability",
                    },
                )
            result = (image_extractor or extract_image_structured_markdown)(source_path)
            _compile_project_image_structured_source(
                source_path,
                target_path,
                workspace,
                paths,
                ingested_at,
                source_hash,
                result,
                classification,
            )
            return ProjectCompiledSource(
                source_path,
                "compiled",
                target_path,
                source_sha256=source_hash,
                metadata={
                    **classifier_metadata,
                    "extraction_status": result.extraction_status,
                    "review_status": "approved",
                    "source_scope_review_policy": "project_no_admin_review",
                    "extractor": result.extractor,
                    "language_policy": result.language_policy,
                    "image_kind": result.image_kind,
                    "model_profile": result.model_profile,
                    "provider": result.provider,
                    "model": result.model,
                    "token_usage": result.token_usage,
                    "warnings": list(result.warnings),
                },
            )
        if classification.file_kind == "archive":
            return ProjectCompiledSource(
                source_path,
                "pending_extractor_capability",
                error=classification.classifier_reason,
                source_sha256=source_hash,
                metadata={
                    **classifier_metadata,
                    "extraction_status": "pending_extractor_capability",
                },
            )
        return ProjectCompiledSource(
            source_path,
            "skipped",
            error=f"unsupported file extension: {source_path.suffix}",
            source_sha256=source_hash,
            metadata=classifier_metadata,
        )
    except Exception as exc:  # pragma: no cover - defensive manifest path
        return ProjectCompiledSource(source_path, "failed", error=str(exc))


def _compile_project_text_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    classification: ExtractorClassification,
) -> None:
    text = source_path.read_text(encoding="utf-8-sig")
    frontmatter, body = _split_frontmatter(text)
    title = frontmatter.get("title") or source_path.stem
    merged = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        **frontmatter,
        "title": title,
        "type": frontmatter.get("type", "project_document"),
        "content_kind": frontmatter.get("content_kind", "project_text_source"),
        "extraction_status": frontmatter.get("extraction_status", "native_text"),
    }
    _write_markdown(target_path, merged, body.strip() + "\n")


def _compile_project_docx_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    classification: ExtractorClassification,
) -> None:
    from docx import Document

    document = Document(str(source_path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            blocks.append(value)
    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", "<br>") for cell in row.cells] for row in table.rows]
        if rows:
            blocks.append(_markdown_table(rows))

    body = "\n\n".join(blocks).strip()
    category = _project_category(source_path, paths["root"])
    content_kind = "meeting_transcript" if category == "meetings" else "project_docx_text_extracted"
    doc_type = "meeting" if category == "meetings" else "project_document"
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": source_path.stem,
        "type": doc_type,
        "content_kind": content_kind,
        "extraction_status": "docx_text_extracted",
    }
    markdown = (
        f"# {source_path.stem}\n\n"
        "> This page was extracted from a Project_R managed project DOCX source. "
        "Review decisions, action items, and commercial facts before promoting them to company knowledge.\n\n"
        "## Extracted Text\n\n"
        f"{body}\n"
    )
    _write_markdown(target_path, frontmatter, markdown)


def _compile_project_pdf_text_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    classification: ExtractorClassification,
) -> None:
    from pypdf import PdfReader

    reader = PdfReader(str(source_path))
    page_texts = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            page_texts.append(f"## Page {index}\n\n{text}")
    body = "\n\n".join(page_texts).strip()
    if not body:
        raise ValueError("PDF text extraction produced no readable content")
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": source_path.stem,
        "type": "project_document",
        "content_kind": "project_pdf_text_extracted",
        "extraction_status": "pdf_text_extracted",
        "review_status": "approved",
        "source_scope_review_policy": "project_no_admin_review",
        "page_count": len(reader.pages),
    }
    markdown = (
        f"# {source_path.stem}\n\n"
        "> This page was extracted from a Project_R managed project PDF with selectable text. "
        "If the PDF contains important layout, table, drawing, stamp, or scan information, rerun it through the complex PDF extractor.\n\n"
        f"{body}\n"
    )
    _write_markdown(target_path, frontmatter, markdown)


def _compile_project_pdf_structured_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    extraction: PDFStructuredExtractionResult,
    classification: ExtractorClassification,
) -> None:
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": source_path.stem,
        "type": "project_document",
        "content_kind": "project_pdf_structured_extract",
        "extraction_status": extraction.extraction_status,
        "review_status": "approved",
        "source_scope_review_policy": "project_no_admin_review",
        "extractor_review_status": extraction.review_status,
        "extractor": extraction.extractor,
        "language_policy": extraction.language_policy,
        "page_count": extraction.page_count,
        "pages_analyzed": extraction.pages_analyzed,
        "model_profile": extraction.model_profile,
        "provider": extraction.provider,
        "model": extraction.model,
        "vision_pages": list(extraction.vision_pages),
        "vision_image_count": extraction.vision_image_count,
    }
    if extraction.warnings:
        frontmatter["extraction_warnings"] = list(extraction.warnings)
    _write_markdown(target_path, frontmatter, extraction.markdown)


def _compile_project_meeting_structured_source(
    source_path: Path,
    transcript_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    extraction: MeetingStructuredExtractionResult,
    classification: ExtractorClassification,
    transcription_result: MediaTranscriptionResult | None = None,
    generated_transcript_path: Path | None = None,
) -> None:
    transcription_status = (
        transcription_result.transcription_status if transcription_result is not None else extraction.transcription_status
    )
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": source_path.stem,
        "type": "meeting",
        "content_kind": "meeting_structured_extract",
        "authority_level": "project_source_record",
        "project_r_transcript_file": _relative_posix(transcript_path, paths["root"]),
        "extraction_status": extraction.extraction_status,
        "review_status": "approved",
        "source_scope_review_policy": "project_no_admin_review",
        "extractor_review_status": extraction.review_status,
        "extractor": extraction.extractor,
        "language_policy": extraction.language_policy,
        "transcription_status": transcription_status,
        "transcription_extractor": transcription_result.extractor if transcription_result is not None else None,
        "transcription_model_profile": transcription_result.model_profile if transcription_result is not None else None,
        "transcription_provider": transcription_result.provider if transcription_result is not None else None,
        "transcription_model": transcription_result.model if transcription_result is not None else None,
        "transcription_token_usage": transcription_result.token_usage if transcription_result is not None else None,
        "generated_transcript_file": (
            _relative_posix(generated_transcript_path, paths["root"]) if generated_transcript_path is not None else None
        ),
        "segment_count": extraction.segment_count,
        "action_item_count": extraction.action_item_count,
        "decision_count": extraction.decision_count,
        "risk_count": extraction.risk_count,
        "tags": ["项目会议", "音视频转写", str(getattr(workspace, "brand", "") or ""), str(getattr(workspace, "slug", "") or "")],
    }
    if extraction.warnings:
        frontmatter["extraction_warnings"] = list(extraction.warnings)
    _write_markdown(target_path, frontmatter, extraction.markdown)


def _compile_project_email_structured_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    extraction: EmailStructuredExtractionResult,
    classification: ExtractorClassification,
) -> None:
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": extraction.subject or source_path.stem,
        "type": "email",
        "content_kind": "email_thread_structured_extract",
        "authority_level": "project_source_record",
        "extraction_status": extraction.extraction_status,
        "review_status": "approved",
        "source_scope_review_policy": "project_no_admin_review",
        "extractor": extraction.extractor,
        "language_policy": extraction.language_policy,
        "email_subject": extraction.subject,
        "email_sender": extraction.sender,
        "email_recipients": list(extraction.recipients),
        "email_message_date": extraction.message_date,
        "email_attachments": list(extraction.attachment_names),
        "model_profile": extraction.model_profile,
        "provider": extraction.provider,
        "model": extraction.model,
        "token_usage": extraction.token_usage,
        "tags": ["项目邮件", "邮件线程", str(getattr(workspace, "brand", "") or ""), str(getattr(workspace, "slug", "") or "")],
    }
    if extraction.warnings:
        frontmatter["extraction_warnings"] = list(extraction.warnings)
    _write_markdown(target_path, frontmatter, extraction.markdown)


def _compile_project_image_structured_source(
    source_path: Path,
    target_path: Path,
    workspace: Any,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
    extraction: ImageStructuredExtractionResult,
    classification: ExtractorClassification,
) -> None:
    frontmatter = {
        **_project_frontmatter(workspace, source_path, paths, ingested_at, source_hash),
        **classification.to_manifest_metadata(),
        "title": source_path.stem,
        "type": "image",
        "content_kind": "image_structured_extract",
        "authority_level": "project_source_record",
        "extraction_status": extraction.extraction_status,
        "review_status": "approved",
        "source_scope_review_policy": "project_no_admin_review",
        "extractor": extraction.extractor,
        "language_policy": extraction.language_policy,
        "image_kind": extraction.image_kind,
        "model_profile": extraction.model_profile,
        "provider": extraction.provider,
        "model": extraction.model,
        "token_usage": extraction.token_usage,
        "tags": ["项目图片", "截图", str(getattr(workspace, "brand", "") or ""), str(getattr(workspace, "slug", "") or "")],
    }
    if extraction.warnings:
        frontmatter["extraction_warnings"] = list(extraction.warnings)
    _write_markdown(target_path, frontmatter, extraction.markdown)


def _project_frontmatter(
    workspace: Any,
    source_path: Path,
    paths: dict[str, Path],
    ingested_at: str,
    source_hash: str,
) -> dict[str, Any]:
    return {
        "source_domain": "project_workspace",
        "authority_level": "project_source_record",
        "project_r_workspace_id": getattr(workspace, "id", None),
        "project_r_workspace_name": getattr(workspace, "name", ""),
        "project_r_workspace_slug": getattr(workspace, "slug", ""),
        "project_r_workspace_brand": getattr(workspace, "brand", ""),
        "project_r_source_file": _relative_posix(source_path, paths["root"]),
        "project_r_source_sha256": source_hash,
        "project_r_ingested_at": ingested_at,
        "tags": ["项目资料", str(getattr(workspace, "brand", "") or ""), str(getattr(workspace, "slug", "") or "")],
    }


def _project_target_path(source_path: Path, paths: dict[str, Path], category: str) -> Path:
    root = paths["root"]
    rel = source_path.resolve().relative_to(root.resolve())
    rest_parts = list(rel.parts[1:-1]) if rel.parts and rel.parts[0] in PROJECT_DIR_CATEGORY_MAP else list(rel.parts[:-1])
    target_dir = paths["derived"] / category
    for part in rest_parts:
        target_dir = target_dir / _safe_filename(part)
    return target_dir / f"{_safe_filename(source_path.stem)}.md"


def _auto_transcript_path(source_path: Path) -> Path:
    return source_path.with_name(f"{source_path.stem}.auto.transcript.md")


def _email_attachment_dir(source_path: Path) -> Path:
    return source_path.with_name(f"{source_path.stem}{EMAIL_ATTACHMENT_DIR_SUFFIX}")


def _project_category(source_path: Path, root: Path) -> str:
    rel = source_path.resolve().relative_to(root.resolve())
    first = rel.parts[0] if rel.parts else ""
    return PROJECT_DIR_CATEGORY_MAP.get(first, "documents")


def _iter_project_source_files(root: Path, paths: dict[str, Path]) -> list[Path]:
    sidecar_dirs = _pdf_image_sidecar_dirs(root)
    transcript_sidecars = find_transcript_sidecars_for_media_files(root, MEDIA_EXTENSIONS)
    files: list[Path] = []
    for source_path in sorted(root.rglob("*")):
        if not source_path.is_file():
            continue
        if source_path.resolve() in transcript_sidecars:
            continue
        if _is_inside_runtime_dir(source_path, paths, sidecar_dirs):
            continue
        files.append(source_path)
    return files


def _is_inside_runtime_dir(path: Path, paths: dict[str, Path], sidecar_dirs: list[Path]) -> bool:
    resolved = path.resolve()
    runtime_roots = [paths["derived"], paths["manifests"], paths["root"] / ".trash", paths["root"] / ".git"]
    runtime_roots.extend(sidecar_dirs)
    for runtime_root in runtime_roots:
        try:
            resolved.relative_to(runtime_root.resolve())
            return True
        except ValueError:
            continue
    return any(part in PROJECT_RUNTIME_DIRS for part in resolved.parts)


def _result_to_project_manifest_item(result: ProjectCompiledSource, root: Path, derived_path: Path) -> dict[str, Any]:
    item: dict[str, Any] = {
        "source_file": _relative_posix(result.source_path, root),
        "status": result.status,
        "source_sha256": result.source_sha256,
    }
    if result.target_path is not None:
        item["target_file"] = _relative_posix(result.target_path, derived_path)
    if result.error:
        item["error"] = result.error
    item.update({key: value for key, value in result.metadata.items() if value not in (None, [], {})})
    return item


def _settings_for_paths(paths: dict[str, Path]) -> GBrainSettings:
    return GBrainSettings(
        enabled=False,
        base_url="",
        home_path=paths["root"],
        raw_path=paths["root"],
        derived_path=paths["derived"],
        manifests_path=paths["manifests"],
        local_git_enabled=False,
    )


def _utc_now() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()
