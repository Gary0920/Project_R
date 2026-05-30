from html import escape
from pathlib import Path
import re
import zipfile


def _text_run(text: str, *, bold: bool = False, size: int | None = None) -> str:
    props = ""
    if bold or size:
        parts = []
        if bold:
            parts.append("<w:b/>")
        if size:
            parts.append(f'<w:sz w:val="{size}"/>')
        props = f"<w:rPr>{''.join(parts)}</w:rPr>"
    return f'<w:r>{props}<w:t xml:space="preserve">{escape(text)}</w:t></w:r>'


def _paragraph(text: str, *, bold: bool = False, size: int | None = None) -> str:
    return f"<w:p>{_text_run(text, bold=bold, size=size)}</w:p>"


def _clean_document_content(content: str) -> str:
    text = content.strip()
    fenced = re.findall(r"```(?:\w+)?\s*\n(.*?)```", text, flags=re.DOTALL)
    if fenced:
        text = max(fenced, key=len).strip()
    text = re.sub(r"^\s*[-*]\s*\[\[[^\]]+\]\]\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^#+\s*本次使用的来源文件[:：]?\s*.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\*\*?本次回答使用的来源文件[:：]\*\*?.*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _paragraph_xml(content: str) -> str:
    lines = _clean_document_content(content).splitlines() or [""]
    paragraphs = []
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            paragraphs.append(_paragraph(""))
            continue
        if re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$", line):
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            size = 32 if level == 1 else 28 if level == 2 else 24
            paragraphs.append(_paragraph(_strip_inline_markdown(heading.group(2)), bold=True, size=size))
            continue
        list_item = re.match(r"^[-*]\s+(.+)$", line)
        numbered_item = re.match(r"^\d+[.)、]\s+(.+)$", line)
        if list_item or numbered_item:
            text = (list_item or numbered_item).group(1)
            paragraphs.append(_paragraph(f"• {_strip_inline_markdown(text)}"))
            continue
        if "|" in line and line.count("|") >= 2:
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            paragraphs.append(_paragraph("    ".join(_strip_inline_markdown(cell) for cell in cells)))
            continue
        paragraphs.append(_paragraph(_strip_inline_markdown(line)))
    return "".join(paragraphs)


def _is_markdown_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$", line.strip()))


def _is_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return "|" in stripped and stripped.count("|") >= 2 and not _is_markdown_table_separator(stripped)


def _split_table_row(line: str) -> list[str]:
    return [_strip_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _add_docx_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index in range(width):
        header_cells[index].text = rows[0][index] if index < len(rows[0]) else ""
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index in range(width):
            cells[index].text = row[index] if index < len(row) else ""


def _render_with_python_docx(title: str, content: str, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(11)

    document.add_heading(title, level=1)
    lines = _clean_document_content(content).splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(_strip_inline_markdown(heading.group(2)), level=level)
            index += 1
            continue
        if _is_markdown_table_row(line):
            rows: list[list[str]] = []
            while index < len(lines):
                table_line = lines[index].strip()
                if _is_markdown_table_separator(table_line):
                    index += 1
                    continue
                if not _is_markdown_table_row(table_line):
                    break
                rows.append(_split_table_row(table_line))
                index += 1
            _add_docx_table(document, rows)
            continue
        list_item = re.match(r"^[-*]\s+(.+)$", line)
        numbered_item = re.match(r"^\d+[.)、]\s+(.+)$", line)
        if list_item:
            document.add_paragraph(_strip_inline_markdown(list_item.group(1)), style="List Bullet")
            index += 1
            continue
        if numbered_item:
            document.add_paragraph(_strip_inline_markdown(numbered_item.group(1)), style="List Number")
            index += 1
            continue
        document.add_paragraph(_strip_inline_markdown(line))
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _render_minimal_docx(title: str, content: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"\s+", " ", title).strip() or "Project_R 文档"
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {_paragraph(safe_title, bold=True, size=34)}
    {_paragraph_xml(content)}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>""",
        )
        docx.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>""",
        )
        docx.writestr("word/document.xml", document_xml)
    return output_path


def render_docx(title: str, content: str, output_path: Path) -> Path:
    """
    Render a Word document from Markdown-like assistant output.
    Uses python-docx when installed, with a small standard-library fallback for
    environments that have not installed the upgraded dependency yet.
    """
    safe_title = re.sub(r"\s+", " ", title).strip() or "Project_R 文档"
    try:
        return _render_with_python_docx(safe_title, content, output_path)
    except ImportError:
        return _render_minimal_docx(safe_title, content, output_path)
