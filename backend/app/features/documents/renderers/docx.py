from __future__ import annotations

from html import escape
from pathlib import Path
import re
import zipfile

from app.features.documents.content_parser import DocumentBlock, InlineSpan, parse_document


def render_docx(title: str, content: str, output_path: Path) -> Path:
    safe_title = re.sub(r"\s+", " ", title).strip() or "Project_R 文档"
    try:
        return _render_with_python_docx(safe_title, content, output_path)
    except ImportError:
        return _render_minimal_docx(safe_title, content, output_path)


def _render_with_python_docx(title: str, content: str, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches, Pt

    parsed = parse_document(title, content)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(11)

    if not _starts_with_title(parsed.blocks, parsed.title):
        document.add_heading(parsed.title, level=1)
    for block in parsed.blocks:
        if block.type == "heading":
            document.add_heading(block.text, level=min(max(block.level, 1), 4))
        elif block.type == "bullet":
            _add_paragraph_with_spans(document, block.spans, style="List Bullet")
        elif block.type == "numbered":
            _add_paragraph_with_spans(document, block.spans, style="List Number")
        elif block.type == "quote":
            _add_paragraph_with_spans(document, block.spans, prefix="引用：")
        elif block.type == "table":
            _add_docx_table(document, block.rows)
        elif block.type == "rule":
            document.add_paragraph("—" * 16)
        elif block.text:
            _add_paragraph_with_spans(document, block.spans)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _starts_with_title(blocks: list[DocumentBlock], title: str) -> bool:
    return bool(blocks and blocks[0].type == "heading" and blocks[0].text.strip() == title.strip())


def _add_paragraph_with_spans(document, spans: list[InlineSpan], *, style: str | None = None, prefix: str = "") -> None:
    paragraph = document.add_paragraph(style=style)
    if prefix:
        run = paragraph.add_run(prefix)
        run.bold = True
    for span in spans:
        run = paragraph.add_run(span.text)
        run.bold = span.bold
        run.italic = span.code


def _add_docx_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index in range(width):
        header_cells[index].text = rows[0][index] if index < len(rows[0]) else ""
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index in range(width):
            cells[index].text = row[index] if index < len(row) else ""


def _render_minimal_docx(title: str, content: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    parsed = parse_document(title, content)
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {_paragraph(parsed.title, bold=True, size=34)}
    {''.join(_block_xml(block) for block in parsed.blocks)}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>"""
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        docx.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>""",
        )
        docx.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>""",
        )
        docx.writestr("word/document.xml", document_xml)
    return output_path


def _block_xml(block: DocumentBlock) -> str:
    if block.type == "table":
        return "".join(_paragraph("    ".join(row)) for row in block.rows)
    if block.type == "heading":
        size = 32 if block.level == 1 else 28 if block.level == 2 else 24
        return _paragraph(block.text, bold=True, size=size)
    if block.type in {"bullet", "numbered"}:
        return _paragraph(f"• {block.text}")
    return _paragraph(block.text)


def _paragraph(text: str, *, bold: bool = False, size: int | None = None) -> str:
    return f"<w:p>{_text_run(text, bold=bold, size=size)}</w:p>"


def _text_run(text: str, *, bold: bool = False, size: int | None = None) -> str:
    props = ""
    if bold or size:
        parts = []
        if bold:
            parts.append("<w:b/>")
        if size:
            parts.append(f'<w:sz w:val="{size}"/>')
        props = f"<w:rPr>{''.join(parts)}</w:rPr>"
    return f'<w:r>{props}<w:t xml:space="preserve">{escape(text)}</w:t></w:r>'
