from __future__ import annotations

from dataclasses import dataclass
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

from fastapi import HTTPException
from sqlalchemy.orm import Session

from app.features.notifications.service import notify_user
from models.workspace import Workspace, WorkspaceFile


@dataclass(frozen=True)
class VersionedLatestMarkdown:
    version_path: Path
    version_rel: str
    latest_path: Path
    latest_rel: str


@dataclass(frozen=True)
class GeneratedMeetingMarkdowns:
    minutes_version_path: Path
    minutes_version_rel: str
    minutes_latest_path: Path
    minutes_latest_rel: str
    actions_version_path: Path
    actions_version_rel: str
    actions_latest_path: Path
    actions_latest_rel: str


def extract_text_from_docx(file_bytes: bytes, filename: str = "") -> str:
    import io

    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX 解析组件未安装")
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"DOCX 文件解析失败：{exc}")


def workspace_file_uploader(db: Session, workspace_id: int, rel_path: str) -> int | None:
    meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    return meta.uploaded_by if meta else None


def write_versioned_latest_markdown(
    *,
    root: Path,
    target_dir: Path,
    version_filename: str,
    latest_filename: str,
    content: str,
    resolve_conflict_path: Callable[[Path, str, str], Path | None],
    error_detail: str,
) -> VersionedLatestMarkdown:
    version_path = resolve_conflict_path(target_dir, version_filename, "keep_both")
    if version_path is None:
        raise HTTPException(status_code=500, detail=error_detail)
    version_path.write_text(content, encoding="utf-8")
    version_rel = version_path.relative_to(root).as_posix()

    latest_path = target_dir / latest_filename
    latest_path.write_text(content, encoding="utf-8")
    latest_rel = latest_path.relative_to(root).as_posix()

    return VersionedLatestMarkdown(
        version_path=version_path,
        version_rel=version_rel,
        latest_path=latest_path,
        latest_rel=latest_rel,
    )


def write_numbered_latest_markdown(
    *,
    root: Path,
    target_dir: Path,
    prefix: str,
    latest_filename: str,
    content: str,
    next_version_number: Callable[[Path, str], int],
) -> VersionedLatestMarkdown:
    version_number = next_version_number(target_dir, prefix)
    version_filename = f"{prefix}-v{version_number}.md"
    version_path = target_dir / version_filename
    version_path.write_text(content, encoding="utf-8")
    version_rel = version_path.relative_to(root).as_posix()

    latest_path = target_dir / latest_filename
    latest_path.write_text(content, encoding="utf-8")
    latest_rel = latest_path.relative_to(root).as_posix()

    return VersionedLatestMarkdown(
        version_path=version_path,
        version_rel=version_rel,
        latest_path=latest_path,
        latest_rel=latest_rel,
    )


def write_generated_meeting_markdowns(
    *,
    root: Path,
    minutes_dir: Path,
    actions_dir: Path,
    minutes_version: int,
    actions_version: int,
    minutes_md: str,
    actions_md: str,
) -> GeneratedMeetingMarkdowns:
    minutes_version_path = minutes_dir / f"minutes-v{minutes_version}.md"
    minutes_version_path.write_text(minutes_md, encoding="utf-8")
    minutes_latest_path = minutes_dir / "minutes-latest.md"
    minutes_latest_path.write_text(minutes_md, encoding="utf-8")

    actions_version_path = actions_dir / f"actions-v{actions_version}.md"
    actions_version_path.write_text(actions_md, encoding="utf-8")
    actions_latest_path = actions_dir / "actions-latest.md"
    actions_latest_path.write_text(actions_md, encoding="utf-8")

    return GeneratedMeetingMarkdowns(
        minutes_version_path=minutes_version_path,
        minutes_version_rel=minutes_version_path.relative_to(root).as_posix(),
        minutes_latest_path=minutes_latest_path,
        minutes_latest_rel=minutes_latest_path.relative_to(root).as_posix(),
        actions_version_path=actions_version_path,
        actions_version_rel=actions_version_path.relative_to(root).as_posix(),
        actions_latest_path=actions_latest_path,
        actions_latest_rel=actions_latest_path.relative_to(root).as_posix(),
    )


def upsert_generated_meeting_file_metadata(
    *,
    workspace_id: int,
    user_id: int,
    generated_files: GeneratedMeetingMarkdowns,
    minutes_md: str,
    actions_md: str,
    upsert_workspace_file: Callable[[int, int, int, str, str, str, int, Path], WorkspaceFile],
    rag_status: str | None = None,
) -> None:
    records = [
        (
            generated_files.minutes_version_rel,
            generated_files.minutes_version_path.name,
            len(minutes_md.encode("utf-8")),
            generated_files.minutes_version_path,
        ),
        (
            generated_files.minutes_latest_rel,
            "minutes-latest.md",
            len(minutes_md.encode("utf-8")),
            generated_files.minutes_latest_path,
        ),
        (
            generated_files.actions_version_rel,
            generated_files.actions_version_path.name,
            len(actions_md.encode("utf-8")),
            generated_files.actions_version_path,
        ),
        (
            generated_files.actions_latest_rel,
            "actions-latest.md",
            len(actions_md.encode("utf-8")),
            generated_files.actions_latest_path,
        ),
    ]
    for rel_path, filename, size, path in records:
        meta = upsert_workspace_file(
            workspace_id,
            user_id,
            rel_path,
            filename,
            "text/markdown",
            size,
            path,
        )
        if rag_status is not None:
            meta.rag_status = rag_status


def notify_meeting_run_finished(
    db: Session,
    *,
    workspace: Workspace,
    actor_user_id: int,
    folder_path: str,
    title: str,
    status: str,
    detail: str,
) -> None:
    severity = "success" if status == "completed" else "warning" if status == "partial" else "critical"
    notify_user(
        db,
        actor_user_id,
        category="workspace",
        severity=severity,
        title=title,
        content=f"{workspace.name}：{detail}",
        action_status="none" if status == "completed" else "pending",
        action_kind="open_workspace",
        action_payload={"workspace_id": workspace.id, "path": folder_path},
        event_key=f"workspace:{workspace.id}:meeting:{folder_path}:{title}:{datetime.now(timezone.utc).timestamp()}",
    )


def filename_match_tokens(filename: str) -> set[str]:
    stem = Path(filename or "").stem.lower()
    raw_tokens = re.findall(r"[a-z0-9]+|[\u4e00-\u9fff]{2,}", stem)
    ignored = {"audio", "video", "meeting", "summary", "transcript", "纪要", "总结", "转录", "会议"}
    return {token for token in raw_tokens if len(token) >= 2 and token not in ignored}


def read_auxiliary_summaries(folder_dir: Path, source_filename: str = "", max_chars: int = 20000) -> str:
    summary_dir = folder_dir / "03-辅助总结"
    if not summary_dir.exists() or not summary_dir.is_dir():
        return ""

    source_tokens = filename_match_tokens(source_filename)
    candidates = [
        child for child in sorted(summary_dir.iterdir(), key=lambda item: item.name.lower())
        if child.is_file() and not child.name.startswith("~$") and child.suffix.lower() in {".md", ".txt", ".docx"}
    ]
    if source_tokens:
        matched = [
            child for child in candidates
            if source_tokens.intersection(filename_match_tokens(child.name))
        ]
        candidates = matched

    sections: list[str] = []
    total = 0
    for child in candidates:
        suffix = child.suffix.lower()
        try:
            if suffix == ".docx":
                text = extract_text_from_docx(child.read_bytes(), child.name)
            else:
                text = child.read_text(encoding="utf-8")
        except Exception as exc:
            sections.append(f"### {child.name}\n\n> 辅助总结读取失败：{exc}")
            continue
        text = text.strip()
        if not text:
            continue
        remaining = max_chars - total
        if remaining <= 0:
            break
        clipped = text[:remaining]
        total += len(clipped)
        suffix_note = "\n\n> 已截断，仅保留前部内容。" if len(text) > len(clipped) else ""
        sections.append(f"### {child.name}\n\n{clipped}{suffix_note}")
    return "\n\n".join(sections)
