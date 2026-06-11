from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


SKILL_NAME = "docx-text-preprocess"
SKILL_VERSION = "1.0.0"
PROMPT_VERSION = "rules-docx-text-v1"
LANGUAGE_POLICY = "bilingual_zh_en_aligned"


@dataclass(frozen=True)
class DocxTextPreprocessResult:
    frontmatter: dict[str, Any]
    markdown: str
    metadata: dict[str, Any] = field(default_factory=dict)


def preprocess_docx_text(
    *,
    source_path: Path,
    source_scope: str,
    source_id: str,
    source_file: str,
    source_sha256: str,
    created_at: str,
    title: str | None = None,
    content_kind: str = "docx_text_source",
    document_type: str = "document",
    extra_frontmatter: dict[str, Any] | None = None,
) -> DocxTextPreprocessResult:
    from docx import Document

    document = Document(str(source_path))
    paragraphs = [_normalize_text(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [value for value in paragraphs if value]
    tables = [_table_to_markdown(table) for table in document.tables]
    tables = [value for value in tables if value]

    resolved_title = title or source_path.stem
    body_blocks = paragraphs + tables
    body = "\n\n".join(body_blocks).strip() or "_No readable DOCX text content extracted._"
    timeline = _timeline_signals(paragraphs)
    entity_values = _entities_mentioned(resolved_title, extra_frontmatter or {})
    metadata = {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "timeline_signal_count": len(timeline),
    }

    frontmatter = {
        "title": resolved_title,
        "source_scope": source_scope,
        "source_id": source_id,
        "content_kind": content_kind,
        "source_file": source_file,
        "source_file_sha256": source_sha256,
        "source_file_type": "docx",
        "preprocess_skill": SKILL_NAME,
        "preprocess_version": SKILL_VERSION,
        "preprocess_status": "succeeded",
        "model_profile": "none",
        "prompt_version": PROMPT_VERSION,
        "language_policy": LANGUAGE_POLICY,
        "created_at": created_at,
        "docx_paragraph_count": len(paragraphs),
        "docx_table_count": len(tables),
    }
    if extra_frontmatter:
        frontmatter.update({key: value for key, value in extra_frontmatter.items() if key not in frontmatter})

    markdown = "\n\n".join(
        [
            f"# {resolved_title}",
            "## Source Summary",
            _source_summary(source_scope, source_file, document_type),
            "## Extracted Facts",
            body,
            "## Entities Mentioned",
            _list_or_empty(entity_values, "_No explicit entities found by deterministic DOCX extraction._"),
            "## Events / Timeline Signals",
            _list_or_empty(timeline, "_No explicit timeline signals found by deterministic DOCX extraction._"),
            "## Original Evidence",
            _original_evidence(source_file, source_sha256, metadata),
            "## Preprocess Notes",
            _preprocess_notes(metadata),
        ]
    ).strip() + "\n"
    return DocxTextPreprocessResult(frontmatter=frontmatter, markdown=markdown, metadata=metadata)


def _normalize_text(value: str) -> str:
    return re.sub(r"[ \t]+", " ", value.replace("\r\n", "\n").replace("\r", "\n")).strip()


def _table_to_markdown(table: Any) -> str:
    rows = [[_normalize_text(cell.text).replace("\n", "<br>") for cell in row.cells] for row in table.rows]
    rows = [row for row in rows if any(cell for cell in row)]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]
    return "\n".join(_markdown_row(row) for row in [header, separator, *body])


def _markdown_row(row: list[str]) -> str:
    return "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"


def _source_summary(source_scope: str, source_file: str, document_type: str) -> str:
    return "\n".join(
        [
            f"- Scope: `{source_scope}`",
            f"- Original file: `{source_file}`",
            f"- Document type: `{document_type}`",
            "- Processing: DOCX paragraph and table extraction only; no model rewriting was applied.",
        ]
    )


def _entities_mentioned(title: str, extra_frontmatter: dict[str, Any]) -> list[str]:
    values = [title]
    for key in ("type", "content_kind", "project_r_workspace_name", "project_r_workspace_brand"):
        if extra_frontmatter.get(key):
            values.append(str(extra_frontmatter[key]))
    return _dedupe(values)


def _timeline_signals(paragraphs: list[str]) -> list[str]:
    signals: list[str] = []
    for paragraph in paragraphs:
        for match in re.finditer(r"\b(?:20\d{2})[-/](?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])\b", paragraph):
            signals.append(match.group(0))
    return _dedupe(signals)


def _original_evidence(source_file: str, source_sha256: str, metadata: dict[str, Any]) -> str:
    return "\n".join(
        [
            f"- Source file: `{source_file}`",
            f"- Source SHA256: `{source_sha256}`",
            f"- Extracted non-empty paragraphs: {metadata['paragraph_count']}",
            f"- Extracted tables: {metadata['table_count']}",
        ]
    )


def _preprocess_notes(metadata: dict[str, Any]) -> str:
    return "\n".join(
        [
            "- Source DOCX file was not modified.",
            "- Text and tables were extracted deterministically with python-docx.",
            f"- Timeline signals found by date pattern: {metadata['timeline_signal_count']}",
        ]
    )


def _list_or_empty(values: list[str], empty: str) -> str:
    if not values:
        return empty
    return "\n".join(f"- {value}" for value in values)


def _dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        normalized = value.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result
