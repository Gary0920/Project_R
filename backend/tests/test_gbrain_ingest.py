import json
import tempfile
import unittest
from pathlib import Path

import yaml

from core.gbrain import GBrainSettings
from core.gbrain_ingest import approve_pending_review_markdown, compile_company_wiki_sources
from core.pdf_structured_extraction import PDFStructuredExtractionResult


class GBrainIngestTests(unittest.TestCase):
    def _settings_for_root(self, root: Path) -> GBrainSettings:
        return GBrainSettings(
            enabled=False,
            base_url="",
            home_path=root,
            raw_path=root / "raw",
            derived_path=root / "derived",
            manifests_path=root / "manifests",
            local_git_enabled=False,
        )

    def test_compiles_markdown_with_existing_frontmatter_and_provenance(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            source = settings.raw_path / "书面化原则.md"
            source.write_text(
                "---\n"
                "title: 书面化原则\n"
                "type: rule\n"
                "tags:\n"
                "  - 规则\n"
                "---\n\n"
                "# 书面化原则\n\n"
                "重要事项需要留痕。\n",
                encoding="utf-8",
            )

            manifest = compile_company_wiki_sources(settings)

            target = settings.derived_path / "rules" / "书面化原则.md"
            self.assertEqual(manifest["summary"]["compiled"], 1)
            self.assertTrue(target.exists())
            frontmatter, body = self._read_frontmatter(target)
            self.assertEqual(frontmatter["title"], "书面化原则")
            self.assertEqual(frontmatter["type"], "rule")
            self.assertEqual(frontmatter["project_r_source_file"], "书面化原则.md")
            self.assertEqual(frontmatter["extraction_status"], "native_text")
            self.assertIn("重要事项需要留痕", body)

    def test_compiles_docx_to_meeting_markdown(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            source = settings.raw_path / "项目会议.docx"
            document = Document()
            document.add_paragraph("会议主题：样板项目复盘")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "事项"
            table.cell(0, 1).text = "负责人"
            table.cell(1, 0).text = "确认报价"
            table.cell(1, 1).text = "Gary"
            document.save(str(source))

            manifest = compile_company_wiki_sources(settings)

            target = settings.derived_path / "meetings" / "项目会议.md"
            self.assertEqual(manifest["summary"]["compiled"], 1)
            frontmatter, body = self._read_frontmatter(target)
            self.assertEqual(frontmatter["type"], "meeting")
            self.assertEqual(frontmatter["content_kind"], "meeting_transcript")
            self.assertIn("会议主题：样板项目复盘", body)
            self.assertIn("| 事项 | 负责人 |", body)

    def test_skips_audio_and_writes_manifest(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            (settings.raw_path / "meeting.mp3").write_bytes(b"not a real audio file")

            manifest = compile_company_wiki_sources(settings)

            manifest_path = settings.manifests_path / "company-wiki-ingest-manifest.json"
            saved_manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["summary"]["skipped"], 1)
            self.assertEqual(saved_manifest["summary"]["skipped"], 1)
            self.assertIn("transcription is pending", saved_manifest["items"][0]["error"])
            self.assertEqual(saved_manifest["items"][0]["transcription_status"], "pending_transcription")

    def test_audio_with_transcript_sidecar_goes_to_pending_review_meeting(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            media = settings.raw_path / "客户会议.mp3"
            transcript = settings.raw_path / "客户会议.transcript.md"
            media.write_bytes(b"not a real audio file")
            transcript.write_text(
                "[00:00] Gary: 今天确认 VMU 样板需要下周完成。\n"
                "[00:25] Amy: 行动项：Amy 负责跟进客户确认。\n"
                "[00:42] Gary: 决定使用黑色窗框。\n"
                "[01:10] Amy: 风险是供应商交期不确定。\n",
                encoding="utf-8",
            )

            manifest = compile_company_wiki_sources(settings)

            target = settings.derived_path / ".pending_review" / "meetings" / "客户会议.md"
            self.assertEqual(manifest["summary"]["total"], 1)
            self.assertEqual(manifest["summary"]["compiled"], 1)
            self.assertTrue(target.exists())
            frontmatter, body = self._read_frontmatter(target)
            self.assertEqual(frontmatter["content_kind"], "meeting_structured_extract")
            self.assertEqual(frontmatter["review_status"], "pending_review")
            self.assertEqual(frontmatter["language_policy"], "bilingual_zh_en_aligned")
            self.assertEqual(frontmatter["transcription_status"], "transcript_sidecar_provided")
            self.assertEqual(frontmatter["project_r_transcript_file"], "客户会议.transcript.md")
            self.assertIn("Action Items", body)
            self.assertIn("Amy 负责跟进客户确认", body)
            item = manifest["items"][0]
            self.assertEqual(item["target_file"], ".pending_review/meetings/客户会议.md")
            self.assertEqual(item["approved_target_file"], "meetings/客户会议.md")
            self.assertEqual(item["transcript_file"], "客户会议.transcript.md")
            self.assertGreaterEqual(item["action_item_count"], 1)

    def test_skips_pdf_until_structured_extraction_is_available(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            (settings.raw_path / "standard.pdf").write_bytes(b"%PDF-1.4\nnot a real pdf")

            manifest = compile_company_wiki_sources(settings)

            self.assertEqual(manifest["summary"]["compiled"], 0)
            self.assertEqual(manifest["summary"]["skipped"], 1)
            self.assertFalse((settings.derived_path / "standards" / "standard.md").exists())
            self.assertIn("structured extraction", manifest["items"][0]["error"])

    def test_compiles_pdf_with_structured_extractor_when_enabled(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            source = settings.raw_path / "standard.pdf"
            source.write_bytes(b"%PDF-1.4\nnot a real pdf")
            stale_approved = settings.derived_path / "standards" / "standard.md"
            stale_approved.parent.mkdir(parents=True)
            stale_approved.write_text(
                "---\nreview_status: approved\n---\n\n# stale approved version\n",
                encoding="utf-8",
            )

            def fake_pdf_extractor(path: Path) -> PDFStructuredExtractionResult:
                self.assertEqual(path, source)
                return PDFStructuredExtractionResult(
                    markdown=(
                        "# standard\n\n"
                        "## 审核状态 / Review Status\n\n"
                        "- 中文：pending_review\n"
                        "  English: pending_review\n\n"
                        "## 核心结论 / Key Conclusions\n\n"
                        "- 中文：示例结构化结论 (p. 1)\n"
                        "  English: Example structured conclusion (p. 1)\n\n"
                        "## 待审核问题 / Review Questions\n\n"
                        "- 中文：需要复核表格。\n"
                        "  English: The table requires review.\n"
                    ),
                    page_count=2,
                    pages_analyzed=2,
                    model_profile="test-profile",
                    provider="test-provider",
                    model="test-model",
                    token_usage={"input_tokens": 10, "output_tokens": 5},
                    warnings=("table layout requires review",),
                )

            manifest = compile_company_wiki_sources(
                settings,
                pdf_extractor=fake_pdf_extractor,
                enable_pdf_structured_extraction=True,
            )

            target = settings.derived_path / ".pending_review" / "standards" / "standard.md"
            self.assertEqual(manifest["summary"]["compiled"], 1)
            self.assertTrue(target.exists())
            self.assertFalse(stale_approved.exists())
            frontmatter, body = self._read_frontmatter(target)
            self.assertEqual(frontmatter["content_kind"], "external_standard_structured_extract")
            self.assertEqual(frontmatter["extraction_status"], "pdf_structured_mvp_pending_review")
            self.assertEqual(frontmatter["review_status"], "pending_review")
            self.assertEqual(frontmatter["language_policy"], "bilingual_zh_en_aligned")
            self.assertEqual(frontmatter["model_profile"], "test-profile")
            self.assertIn("示例结构化结论", body)
            self.assertIn("Example structured conclusion", body)
            item = manifest["items"][0]
            self.assertEqual(item["target_file"], ".pending_review/standards/standard.md")
            self.assertEqual(item["approved_target_file"], "standards/standard.md")
            self.assertEqual(item["extraction_status"], "pdf_structured_mvp_pending_review")
            self.assertEqual(item["review_status"], "pending_review")
            self.assertEqual(item["language_policy"], "bilingual_zh_en_aligned")
            self.assertEqual(item["page_count"], 2)
            self.assertEqual(item["warnings"], ["table layout requires review"])

    def test_pdf_image_sidecar_folder_is_not_scanned_as_raw_sources(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            settings.raw_path.mkdir(parents=True)
            (settings.raw_path / "AS 2047.pdf").write_bytes(b"%PDF-1.4\nnot a real pdf")
            sidecar = settings.raw_path / "AS-2047"
            sidecar.mkdir()
            (sidecar / "p001.png").write_bytes(b"not a real png")

            manifest = compile_company_wiki_sources(settings)

            self.assertEqual(manifest["summary"]["total"], 1)
            self.assertEqual(manifest["items"][0]["source_file"], "AS 2047.pdf")

    def test_approve_pending_review_markdown_promotes_to_syncable_path(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            settings = self._settings_for_root(root)
            pending = settings.derived_path / ".pending_review" / "standards" / "standard.md"
            pending.parent.mkdir(parents=True)
            pending.write_text(
                "---\n"
                "title: standard\n"
                "review_status: pending_review\n"
                "---\n\n"
                "# standard\n\n"
                "待审核内容\n",
                encoding="utf-8",
            )

            result = approve_pending_review_markdown(settings, ".pending_review/standards/standard.md", reviewer_id=1)

            approved = settings.derived_path / "standards" / "standard.md"
            self.assertEqual(result["approved_file"], "standards/standard.md")
            self.assertFalse(pending.exists())
            frontmatter, body = self._read_frontmatter(approved)
            self.assertEqual(frontmatter["review_status"], "approved")
            self.assertEqual(frontmatter["project_r_reviewer_id"], 1)
            self.assertIn("待审核内容", body)

    def _read_frontmatter(self, path: Path):
        text = path.read_text(encoding="utf-8")
        _, raw_frontmatter, body = text.split("---", 2)
        return yaml.safe_load(raw_frontmatter), body


if __name__ == "__main__":
    unittest.main()
