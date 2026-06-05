import tempfile
import unittest
from pathlib import Path


class DocxTextPreprocessTests(unittest.TestCase):
    def test_preprocess_docx_text_outputs_gbrain_ready_template(self):
        from docx import Document

        from core.docx_text_preprocess import preprocess_docx_text

        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "项目会议.docx"
            document = Document()
            document.add_paragraph("会议主题：样板项目复盘")
            document.add_paragraph("日期：2026-06-05")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "事项"
            table.cell(0, 1).text = "负责人"
            table.cell(1, 0).text = "确认报价"
            table.cell(1, 1).text = "Gary"
            document.save(str(source))
            original_size = source.stat().st_size

            result = preprocess_docx_text(
                source_path=source,
                source_scope="project",
                source_id="project-test-1",
                source_file="03-会议纪要/项目会议.docx",
                source_sha256="abc",
                created_at="2026-06-05T00:00:00+00:00",
                content_kind="meeting_transcript",
                document_type="meeting",
            )

            self.assertEqual(source.stat().st_size, original_size)
            self.assertEqual(result.frontmatter["preprocess_skill"], "docx-text-preprocess")
            self.assertEqual(result.frontmatter["source_file_type"], "docx")
            self.assertEqual(result.frontmatter["docx_paragraph_count"], 2)
            self.assertEqual(result.frontmatter["docx_table_count"], 1)
            self.assertIn("## Source Summary", result.markdown)
            self.assertIn("## Extracted Facts", result.markdown)
            self.assertIn("## Entities Mentioned", result.markdown)
            self.assertIn("## Events / Timeline Signals", result.markdown)
            self.assertIn("## Original Evidence", result.markdown)
            self.assertIn("会议主题：样板项目复盘", result.markdown)
            self.assertIn("| 事项 | 负责人 |", result.markdown)
            self.assertIn("2026-06-05", result.markdown)


if __name__ == "__main__":
    unittest.main()
