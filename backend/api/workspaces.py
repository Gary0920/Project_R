import json
import re
import base64
import binascii
import mimetypes
from datetime import datetime, timezone
from pathlib import Path
import shutil
import subprocess
import uuid

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy import or_, select
from sqlalchemy.orm import Session

from app.features.agents.schemas import AgentRunResponse
from api.auth import get_current_user
from app.shared.time.schemas import UTCDateTimeModel
from app.shared.time.utils import serialize_datetime_utc
from app.features.agents.events import add_agent_event, create_agent_run, finish_agent_run, serialize_agent_run
from core.gbrain import (
    GBrainAdapter,
    customer_source_id_for_workspace,
    customer_source_paths_for_workspace,
    project_source_id_for_workspace,
    project_source_paths_for_workspace,
)
from app.features.knowledge.gbrain.customer_sources import CUSTOMER_WORKSPACE_INGEST_MANIFEST_NAME, compile_customer_workspace_sources
from app.features.knowledge.gbrain.graph import (
    apply_entity_merge_candidate_action,
    build_entity_merge_candidate_preview,
    build_entity_merge_candidates,
    build_source_graph,
)
from app.features.knowledge.gbrain.project_ingest import PROJECT_INGEST_MANIFEST_NAME, compile_project_workspace_sources
from app.features.notifications.service import (
    notify_user,
    notify_workspace_bulk_delete_risk,
    notify_workspace_joined,
)
from app.features.workspaces.files.signature import (
    file_signature as _file_signature,
    record_file_signature as _record_file_signature,
)
from app.features.workspaces.ingest.projection import (
    finalize_workspace_ingest_projection as _finalize_workspace_ingest_projection,
    update_workspace_file_rag_statuses_from_manifest as _update_workspace_file_rag_statuses_from_manifest,
)
from app.features.workspaces.ingest.executor import execute_workspace_ingest_core
from app.features.workspaces.ingest.agent_runs import (
    add_workspace_ingest_result_event,
    add_workspace_ingest_started_event,
    create_queued_workspace_ingest_agent_run,
    fail_workspace_ingest_agent_run,
    finish_workspace_ingest_agent_run,
    get_or_create_workspace_ingest_agent_run as _get_or_create_workspace_ingest_agent_run,
    serialize_workspace_ingest_agent_run as _serialize_workspace_ingest_agent_run,
    write_immediate_workspace_ingest_agent_run as _write_immediate_workspace_ingest_agent_run,
)
from app.features.workspaces.ingest.audit import workspace_ingest_audit_fields
from app.features.workspaces.ingest.jobs import (
    mark_workspace_ingest_job_completed,
    mark_workspace_ingest_job_failed,
    mark_workspace_ingest_job_queued,
    mark_workspace_ingest_job_running,
    workspace_ingest_job_run_id as _workspace_ingest_job_run_id,
    workspace_ingest_request_detail as _workspace_ingest_request_detail,
    workspace_ingest_request_from_job as _workspace_ingest_request_from_job,
    workspace_ingest_request_label as _workspace_ingest_request_label,
    workspace_ingest_run_id_from_job as _workspace_ingest_run_id_from_job,
)
from app.features.workspaces.ingest.notifications import (
    notify_workspace_ingest_failed,
    notify_workspace_ingest_finished as _notify_workspace_ingest_finished,
    notify_workspace_ingest_queued,
)
from app.features.workspaces.ingest.run import (
    derive_workspace_ingest_run_status as _derive_workspace_ingest_run_status,
    finalize_workspace_ingest_manifest as _finalize_workspace_ingest_manifest,
    overall_workspace_ingest_rag_status as _overall_project_ingest_status,
    workspace_ingest_manifest_counts as _workspace_ingest_manifest_counts,
    workspace_ingest_result_payload as _workspace_ingest_result_payload,
    workspace_ingest_run_payload as _workspace_ingest_run_payload,
    workspace_ingest_run_status_label as _workspace_ingest_run_status_label,
    workspace_ingest_status_event as _workspace_ingest_status_event,
    workspace_ingest_summary_text as _workspace_ingest_summary_text,
)
from app.features.workspaces.files.service import (
    DEFAULT_UNFILED_DIR,
    DEFAULT_PROJECT_WORKSPACE_TEMPLATE_DIRS,
    DEFAULT_WORKSPACE_DIRS,
    MAX_WORKSPACE_ADMIN_UPLOAD_BYTES,
    MAX_WORKSPACE_ADMIN_UPLOAD_MB,
    MAX_WORKSPACE_UPLOAD_BYTES,
    MAX_WORKSPACE_UPLOAD_MB,
    MEETING_SUBDIRS,
    is_template_root as _is_template_root,
    make_meeting_folder_name,
    meeting_folder_collision_free,
    meeting_parent_path,
    member_can_mutate_file as _member_can_mutate_file,
    member_can_restore_file as _member_can_restore_file,
    resolve_conflict_path as _resolve_conflict_path,
    resolve_workspace_child as _resolve_workspace_child,
    safe_name as _safe_name,
    safe_relative_path as _safe_relative_path,
    TRASH_DIRNAME,
    trash_target as _trash_target,
    upload_limit_for,
)
from models import SessionLocal, get_db
from models.audit_log import AuditLog
from models.attachment import SessionAttachment
from models.workspace import Workspace, WorkspaceMember, WorkspaceGroupAccess, WorkspaceFile
from models.workspace_ingest_job import WorkspaceIngestJob
from models.user import User

router = APIRouter(prefix="/workspaces", tags=["workspaces"])

BASE_DIR = Path(__file__).resolve().parent.parent
WORKSPACES_ROOT = BASE_DIR / "workspace_data"
PROJECT_ROOT_NAME = "project"
USER_ROOT_NAME = "user"
CUSTOMER_ROOT_NAME = "customer"
PROJECT_BRANDS = ("AURA", "BFI", "SPECWISE", "SYNOVA")
CUSTOMER_BRAND = "CUSTOMER"
CRM_WORKSPACE_NAME = "CRM"
CRM_WORKSPACE_SLUG = "CRM"
CRM_RAW_DIR = "raw"


def _normalize_group_name(value: str | None) -> str:
    return (value or "").strip()


class CreateWorkspaceRequest(BaseModel):
    name: str
    description: str = ""
    brand: str = "BFI"
    workspace_kind: str = "project"


class UpdateWorkspaceRequest(BaseModel):
    name: str | None = None
    description: str | None = None
    is_hidden: bool | None = None


class WorkspaceResponse(UTCDateTimeModel):
    id: int
    name: str
    slug: str
    description: str
    created_by: int
    member_count: int = 0
    brand: str = "BFI"
    workspace_kind: str = "project"
    is_default: bool = False
    is_archived: bool
    is_hidden: bool = False
    can_rename: bool = True
    can_delete: bool = True
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class WorkspaceDetailResponse(WorkspaceResponse):
    storage_path: str
    members: list["MemberResponse"]
    access_groups: list[str] = Field(default_factory=list)


class WorkspaceFileItemResponse(UTCDateTimeModel):
    id: int | None = None
    name: str
    path: str
    type: str
    size: int | None = None
    updated_at: datetime | None = None
    uploaded_by: int | None = None
    uploader_name: str | None = None
    deleted_at: datetime | None = None
    deleted_by: int | None = None
    rag_status: str | None = None
    can_delete: bool = False
    can_restore: bool = False
    children: list["WorkspaceFileItemResponse"] = Field(default_factory=list)


class WorkspaceFilesResponse(BaseModel):
    workspace_id: int
    root_name: str
    items: list[WorkspaceFileItemResponse]


class WorkspaceKnowledgeGraphResponse(BaseModel):
    ok: bool
    workspace_id: int
    workspace_name: str
    workspace_kind: str
    source_id: str
    source_scope: str
    intelligence_kind: str
    derived_path: str | None = None
    focus: str | None = None
    entity_type: str | None = None
    nodes: list[dict]
    edges: list[dict]
    events: list[dict]
    profile_cards: list[dict] = Field(default_factory=list)
    stats: dict | None = None
    warnings: list[str] = Field(default_factory=list)


class WorkspaceEntityMergeCandidatesResponse(BaseModel):
    ok: bool
    workspace_id: int
    workspace_name: str
    workspace_kind: str
    source_id: str
    source_scope: str
    derived_path: str | None = None
    focus: str | None = None
    candidates: list[dict]
    stats: dict | None = None
    warnings: list[str] = Field(default_factory=list)


class WorkspaceEntityMergeActionRequest(BaseModel):
    candidate_id: str
    action: str


class UploadWorkspaceFileRequest(BaseModel):
    directory: str = ""
    filename: str
    content_base64: str
    content_type: str = "application/octet-stream"
    conflict_strategy: str = "keep_both"


class CreateWorkspaceFolderRequest(BaseModel):
    parent_path: str = ""
    name: str


MEETING_TYPES = [
    "项目统筹会",
    "客户沟通会",
    "技术交底",
    "现场协调",
    "内部复盘",
    "培训分享",
    "其他",
]
MEETING_TYPE_META_FILENAME = ".meeting-meta.json"


class CreateMeetingFolderRequest(BaseModel):
    topic: str
    meeting_time: str | None = None  # ISO-8601 datetime string, optional
    meeting_type: str = "其他"  # one of MEETING_TYPES


class SaveMeetingTranscriptRequest(BaseModel):
    folder_path: str  # relative path of the meeting folder inside the workspace
    content: str
    input_type: str = "paste"  # paste / txt / md / docx
    original_filename: str = ""


class MeetingFolderResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    created_dirs: list[str]
    created_files: list[str]
    gbrain_ingest: bool = False
    agent_run: AgentRunResponse | None = None


class SaveMeetingTranscriptResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    transcript_v1_path: str
    transcript_latest_path: str
    gbrain_ingest: bool = False
    agent_run: AgentRunResponse | None = None


class RenameWorkspacePathRequest(BaseModel):
    path: str
    new_name: str


class MoveWorkspacePathRequest(BaseModel):
    path: str
    target_directory: str = ""
    conflict_strategy: str = "keep_both"


class CopyWorkspacePathRequest(BaseModel):
    path: str
    target_directory: str = ""
    conflict_strategy: str = "keep_both"


class SaveAttachmentToWorkspaceRequest(BaseModel):
    session_id: int
    attachment_id: int
    conflict_strategy: str = "keep_both"


class WorkspaceFileMutationResponse(BaseModel):
    ok: bool
    path: str
    file_id: int | None = None
    rag_status: str | None = None
    agent_run: AgentRunResponse | None = None


class WorkspaceMultiUploadResponse(BaseModel):
    ok: bool
    files: list[WorkspaceFileMutationResponse]
    agent_run: AgentRunResponse | None = None


class WorkspaceTrashClearResponse(BaseModel):
    ok: bool
    deleted_files: int
    agent_run: AgentRunResponse | None = None


class RestoreWorkspaceFileRequest(BaseModel):
    file_id: int


class WorkspaceKnowledgeIngestRequest(BaseModel):
    path: str = ""
    recursive: bool = True


class WorkspaceKnowledgeRefreshResponse(BaseModel):
    ok: bool
    workspace_id: int
    indexed_files: int
    rag_status: str
    compiled_files: int = 0
    pending_extractor_capability_files: int = 0
    pending_transcription_files: int = 0
    skipped_files: int = 0
    failed_files: int = 0
    pending_reviews_created: int = 0
    ingest_path: str = ""
    ingest_recursive: bool = True
    gbrain_source_id: str | None = None
    gbrain_status: str | None = None
    gbrain_sync_status: str | None = None
    gbrain_think_status: str | None = None
    gbrain_error: str | None = None
    run_status: str | None = None
    run_id: str | None = None
    manifest: dict | None = None
    agent_run: AgentRunResponse | None = None


class WorkspaceKnowledgeIngestJobResponse(UTCDateTimeModel):
    id: int
    workspace_id: int
    requested_by: int
    status: str
    result: dict = Field(default_factory=dict)
    error_message: str = ""
    created_at: datetime
    started_at: datetime | None = None
    finished_at: datetime | None = None
    agent_run: AgentRunResponse | None = None


class MemberResponse(UTCDateTimeModel):
    user_id: int
    username: str
    nickname: str
    role: str
    joined_at: datetime


class UpsertWorkspaceMemberRequest(BaseModel):
    user_id: int | None = None
    username: str | None = None
    role: str = "member"


class UpdateWorkspaceMemberRoleRequest(BaseModel):
    role: str


class UpsertWorkspaceGroupRequest(BaseModel):
    group_name: str


class WorkspaceGroupResponse(BaseModel):
    group_name: str


class WorkspaceMemberCandidateResponse(BaseModel):
    user_id: int
    username: str
    nickname: str
    work_group: str = ""
    role: str
    is_member: bool = False
    member_role: str | None = None


class WorkspaceGroupCandidateResponse(BaseModel):
    group_name: str
    source: str = "user"
    is_authorized: bool = False


def _slugify(name: str) -> str:
    slug = re.sub(r"[^\w一-鿿-]", "-", name.strip()).strip("-")
    return re.sub(r"-{2,}", "-", slug) or "workspace"


def _safe_username(username: str) -> str:
    return _slugify(username).replace("/", "-") or "user"


def _project_brand_names() -> list[str]:
    names = set(PROJECT_BRANDS)
    for brand, _ in _project_brand_dirs():
        names.add(brand)
    return sorted(names)


def _project_brand_dirs() -> list[tuple[str, Path]]:
    project_root = (WORKSPACES_ROOT / PROJECT_ROOT_NAME).resolve()
    entries: dict[str, Path] = {}
    for brand in PROJECT_BRANDS:
        entries[brand] = (project_root / brand).resolve()
    if project_root.exists():
        for child in project_root.iterdir():
            if child.is_dir() and not child.is_symlink():
                normalized = re.sub(r"[^A-Za-z0-9._-]+", "-", child.name.strip()).strip("-._")
                if normalized:
                    entries[normalized.upper()] = child.resolve()
    return sorted(entries.items(), key=lambda item: item[0])


def _normalize_brand(brand: str) -> str:
    normalized = brand.strip().upper()
    if not re.fullmatch(r"[A-Z0-9][A-Z0-9._-]{0,63}", normalized or ""):
        raise HTTPException(status_code=400, detail="项目品牌不合法")
    if normalized not in _project_brand_names():
        raise HTTPException(status_code=400, detail="项目品牌不合法")
    return normalized


def _normalize_workspace_kind(kind: str | None, brand: str | None = None) -> str:
    normalized = (kind or "").strip().lower()
    if (brand or "").strip().upper() == CUSTOMER_BRAND:
        normalized = "customer"
    if not normalized:
        normalized = "project"
    if normalized not in {"project", "customer"}:
        raise HTTPException(status_code=400, detail="工作区类型不合法")
    return normalized


def _workspace_dirs(workspace: Workspace) -> tuple[str, ...]:
    if workspace.workspace_kind == "project":
        return DEFAULT_PROJECT_WORKSPACE_TEMPLATE_DIRS
    if workspace.workspace_kind == "customer":
        return (CRM_RAW_DIR,)
    return ()


def _is_trash_relative_path(path: Path) -> bool:
    return bool(path.parts) and path.parts[0] == TRASH_DIRNAME


def _ensure_not_trash_path(path: Path) -> None:
    if _is_trash_relative_path(path):
        raise HTTPException(status_code=400, detail="回收站不能作为普通文件夹操作")


def _target_storage_path(workspace: Workspace, owner: User | None = None) -> Path:
    if workspace.workspace_kind == "user":
        return WORKSPACES_ROOT.resolve()
    if workspace.workspace_kind == "customer":
        return (WORKSPACES_ROOT / CUSTOMER_ROOT_NAME / CRM_WORKSPACE_SLUG).resolve()
    brand = _normalize_brand(workspace.brand or "BFI")
    return (WORKSPACES_ROOT / PROJECT_ROOT_NAME / brand / workspace.slug).resolve()


def _ensure_storage_path(workspace: Workspace, *, create_user_scaffold: bool = False) -> str:
    if workspace.workspace_kind == "user":
        return ""
    root = WORKSPACES_ROOT.resolve()
    target = _target_storage_path(workspace)
    path = Path(workspace.storage_path) if workspace.storage_path else target
    resolved = path.resolve()
    legacy_path = (WORKSPACES_ROOT / workspace.slug).resolve()
    if workspace.workspace_kind == "project" and resolved == legacy_path and legacy_path.exists() and not target.exists():
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(legacy_path), str(target))
        path = target
        resolved = target.resolve()
    if not resolved.is_relative_to(root) or (
        workspace.workspace_kind == "project"
        and not resolved.is_relative_to((WORKSPACES_ROOT / PROJECT_ROOT_NAME).resolve())
    ):
        path = target
    if workspace.workspace_kind == "customer" and path.resolve() != target:
        path = target
    path.mkdir(parents=True, exist_ok=True)
    for dirname in _workspace_dirs(workspace):
        (path / dirname).mkdir(parents=True, exist_ok=True)
    (path / TRASH_DIRNAME).mkdir(exist_ok=True)
    return str(path)


def _workspace_file_root(workspace: Workspace) -> Path:
    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不提供后端文件区")
    return Path(_ensure_storage_path(workspace)).resolve()


def _candidate_storage_path(slug: str, brand: str, workspace_kind: str = "project") -> Path:
    if workspace_kind == "customer":
        return (WORKSPACES_ROOT / CUSTOMER_ROOT_NAME / CRM_WORKSPACE_SLUG).resolve()
    return (WORKSPACES_ROOT / PROJECT_ROOT_NAME / brand / slug).resolve()


def _workspace_response(
    db: Session,
    workspace: Workspace,
    member_count: int | None = None,
    user: User | None = None,
) -> WorkspaceResponse:
    count = member_count
    if count is None:
        count = db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == workspace.id).count()
    can_manage = bool(user and _is_workspace_admin(db, user, workspace.id))
    return WorkspaceResponse(
        id=workspace.id,
        name=workspace.name,
        slug=workspace.slug,
        description=workspace.description,
        created_by=workspace.created_by,
        member_count=count,
        brand=workspace.brand,
        workspace_kind=workspace.workspace_kind,
        is_default=workspace.is_default,
        is_archived=workspace.is_archived,
        is_hidden=workspace.is_hidden,
        can_rename=not workspace.is_default and can_manage,
        can_delete=not workspace.is_default and can_manage,
        created_at=workspace.created_at,
        updated_at=workspace.updated_at,
    )


def _serialize_workspace_member(member: WorkspaceMember, user: User) -> MemberResponse:
    return MemberResponse(
        user_id=user.id,
        username=user.username,
        nickname=user.nickname,
        role=member.role,
        joined_at=member.joined_at,
    )


def _validate_workspace_member_role(role: str) -> str:
    normalized = role.strip().lower()
    if normalized not in {"admin", "member"}:
        raise HTTPException(status_code=400, detail="工作区成员角色不合法")
    return normalized


def _workspace_membership(db: Session, user_id: int, workspace_id: int) -> WorkspaceMember | None:
    return (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.workspace_id == workspace_id,
            WorkspaceMember.user_id == user_id,
        )
        .first()
    )


def _workspace_access_groups(db: Session, workspace_id: int) -> list[str]:
    return [
        row.group_name
        for row in db.query(WorkspaceGroupAccess)
        .filter(WorkspaceGroupAccess.workspace_id == workspace_id)
        .order_by(WorkspaceGroupAccess.group_name.asc())
        .all()
    ]


def _has_workspace_group_access(db: Session, user: User, workspace_id: int) -> bool:
    group_name = _normalize_group_name(getattr(user, "work_group", ""))
    if not group_name:
        return False
    return bool(
        db.query(WorkspaceGroupAccess)
        .filter(
            WorkspaceGroupAccess.workspace_id == workspace_id,
            WorkspaceGroupAccess.group_name == group_name,
        )
        .first()
    )


def _can_open_workspace(db: Session, user: User, workspace: Workspace) -> bool:
    if workspace.workspace_kind == "user":
        return _workspace_membership(db, user.id, workspace.id) is not None
    if user.role == "admin":
        return True
    if workspace.workspace_kind == "customer":
        return _workspace_membership(db, user.id, workspace.id) is not None or _has_workspace_group_access(db, user, workspace.id)
    if not workspace.is_hidden:
        return True
    return _workspace_membership(db, user.id, workspace.id) is not None or _has_workspace_group_access(db, user, workspace.id)


def _ensure_can_open_workspace(db: Session, user: User, workspace: Workspace) -> None:
    if not _can_open_workspace(db, user, workspace):
        raise HTTPException(status_code=403, detail="无权访问该工作区")


def _workspace_gbrain_graph_scope(workspace: Workspace) -> dict[str, object]:
    if workspace.workspace_kind == "project":
        paths = project_source_paths_for_workspace(workspace)
        return {
            "source_id": project_source_id_for_workspace(workspace),
            "derived_path": _workspace_gbrain_read_path(paths),
            "gbrain_ready_path": paths["gbrain_ready"],
            "legacy_derived_path": paths.get("legacy_derived"),
            "source_scope": "project",
            "intelligence_kind": "project_event_graph",
        }
    if workspace.workspace_kind == "customer":
        paths = customer_source_paths_for_workspace(workspace)
        return {
            "source_id": customer_source_id_for_workspace(workspace),
            "derived_path": _workspace_gbrain_read_path(paths),
            "gbrain_ready_path": paths["gbrain_ready"],
            "legacy_derived_path": paths.get("legacy_derived"),
            "source_scope": "customer",
            "intelligence_kind": "customer_intelligence",
        }
    raise HTTPException(status_code=400, detail="当前工作区不支持 GBrain 图谱")


def _workspace_gbrain_read_path(paths: dict[str, Path]) -> Path:
    ready = paths["gbrain_ready"]
    legacy = paths.get("legacy_derived")
    if _has_markdown_files(ready) or legacy is None:
        return ready
    if _has_markdown_files(legacy):
        return legacy
    return ready


def _has_markdown_files(path: Path) -> bool:
    if not path.exists():
        return False
    return any(item.is_file() and item.suffix.lower() in {".md", ".markdown"} for item in path.rglob("*"))


def _workspace_profile_cards(graph: dict[str, object]) -> list[dict[str, object]]:
    nodes = [node for node in graph.get("nodes", []) if isinstance(node, dict)]
    edges = [edge for edge in graph.get("edges", []) if isinstance(edge, dict)]
    events = [event for event in graph.get("events", []) if isinstance(event, dict)]
    relation_count: dict[str, int] = {}
    event_count: dict[str, int] = {}
    for edge in edges:
        for key in ("from", "to"):
            node_id = str(edge.get(key) or "")
            if node_id:
                relation_count[node_id] = relation_count.get(node_id, 0) + 1
    for event in events:
        node_id = str(event.get("entity_id") or "")
        if node_id:
            event_count[node_id] = event_count.get(node_id, 0) + 1

    def card_priority(node: dict[str, object]) -> tuple[int, int, str]:
        node_id = str(node.get("id") or "")
        kind = str(node.get("entity_type") or "").lower()
        type_score = 0
        if any(token in kind for token in ("client", "customer", "contact", "person", "company")):
            type_score = 3
        elif "project" in kind:
            type_score = 2
        elif "event" in kind:
            type_score = 1
        return (type_score, relation_count.get(node_id, 0) + event_count.get(node_id, 0), str(node.get("title") or "").lower())

    cards: list[dict[str, object]] = []
    for node in sorted(nodes, key=card_priority, reverse=True)[:8]:
        node_id = str(node.get("id") or "")
        cards.append(
            {
                "id": node_id,
                "title": str(node.get("title") or ""),
                "entity_type": str(node.get("entity_type") or "page"),
                "relation_count": relation_count.get(node_id, 0),
                "event_count": event_count.get(node_id, 0),
                "citation": node.get("citation") if isinstance(node.get("citation"), dict) else None,
            }
        )
    return cards


def _is_workspace_admin(db: Session, user: User, workspace_id: int) -> bool:
    member = _workspace_membership(db, user.id, workspace_id)
    if member and member.role == "admin":
        return True
    if user.role == "admin":
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        return bool(workspace and workspace.workspace_kind != "user")
    return False


def _ensure_workspace_admin(db: Session, user: User, workspace_id: int) -> WorkspaceMember:
    member = _ensure_member(db, user.id, workspace_id)
    if member.role != "admin" and user.role != "admin":
        raise HTTPException(status_code=403, detail="仅工作区管理员可管理成员")
    return member


def _ensure_mutable_membership_workspace(workspace: Workspace) -> None:
    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持邀请成员")


def _resolve_member_target_user(db: Session, req: UpsertWorkspaceMemberRequest) -> User:
    target_user: User | None = None
    if req.user_id is not None:
        target_user = db.query(User).filter(User.id == req.user_id).first()
    elif req.username and req.username.strip():
        target_user = db.query(User).filter(User.username == req.username.strip()).first()
    else:
        raise HTTPException(status_code=400, detail="需要提供用户 ID 或用户名")
    if not target_user:
        raise HTTPException(status_code=404, detail="用户不存在")
    if not target_user.is_active:
        raise HTTPException(status_code=400, detail="用户已被禁用")
    return target_user


def _local_workspace_admin_count(db: Session, workspace_id: int) -> int:
    return (
        db.query(WorkspaceMember)
        .filter(WorkspaceMember.workspace_id == workspace_id, WorkspaceMember.role == "admin")
        .count()
    )


def ensure_default_workspace(db: Session, user: User) -> Workspace:
    existing = (
        db.query(Workspace)
        .join(WorkspaceMember, WorkspaceMember.workspace_id == Workspace.id)
        .filter(
            Workspace.workspace_kind == "user",
            Workspace.is_default == True,
            WorkspaceMember.user_id == user.id,
        )
        .first()
    )
    if existing:
        expected_name = f"{user.username}的工作台"
        if existing.name.endswith(" 的私人空间") or existing.name.endswith("的私人空间"):
            existing.name = expected_name
            existing.description = "用户默认个人工作台"
        if existing.storage_path:
            existing.storage_path = ""
            db.commit()
        elif existing.name == expected_name:
            db.commit()
        return existing

    slug_base = f"user-{_safe_username(user.username)}"
    slug = slug_base
    suffix = 2
    while db.query(Workspace).filter(Workspace.slug == slug).first():
        slug = f"{slug_base}-{suffix}"
        suffix += 1
    workspace = Workspace(
        name=f"{user.username}的工作台",
        slug=slug,
        description="用户默认个人工作台",
        created_by=user.id,
        brand="",
        workspace_kind="user",
        is_default=True,
    )
    db.add(workspace)
    db.flush()
    workspace.storage_path = ""
    db.add(WorkspaceMember(workspace_id=workspace.id, user_id=user.id, role="admin"))
    db.commit()
    db.refresh(workspace)
    return workspace


def _find_existing_project_folder(brand: str, slug: str, name: str | None = None) -> Path | None:
    brand_dir = (WORKSPACES_ROOT / PROJECT_ROOT_NAME / brand).resolve()
    if not brand_dir.exists() or not brand_dir.is_dir():
        return None
    candidates = {slug.casefold()}
    if name:
        candidates.add(name.casefold())
    for child in brand_dir.iterdir():
        if not child.is_dir() or child.is_symlink():
            continue
        if child.name.casefold() in candidates or _slugify(child.name) == slug:
            return child.resolve()
    return None


def _register_existing_project_folder(
    db: Session,
    user: User,
    brand: str,
    project_dir: Path,
    *,
    add_member: bool = False,
) -> Workspace | None:
    if not project_dir.exists() or not project_dir.is_dir() or project_dir.is_symlink():
        return None
    slug = _slugify(project_dir.name)
    if db.query(Workspace).filter(Workspace.slug == slug).first():
        return None
    workspace = Workspace(
        name=project_dir.name,
        slug=slug,
        description="",
        created_by=user.id,
        storage_path=str(project_dir.resolve()),
        brand=brand,
        workspace_kind="project",
        is_default=False,
    )
    db.add(workspace)
    db.flush()
    workspace.storage_path = _ensure_storage_path(workspace)
    if add_member:
        db.add(WorkspaceMember(workspace_id=workspace.id, user_id=user.id, role="admin"))
    db.commit()
    db.refresh(workspace)
    return workspace


def _find_crm_workspace(db: Session) -> Workspace | None:
    exact = (
        db.query(Workspace)
        .filter(
            Workspace.workspace_kind == "customer",
            or_(Workspace.slug == CRM_WORKSPACE_SLUG, Workspace.name == CRM_WORKSPACE_NAME),
        )
        .order_by(Workspace.id.asc())
        .first()
    )
    if exact:
        return exact
    return (
        db.query(Workspace)
        .filter(Workspace.workspace_kind == "customer")
        .order_by(Workspace.id.asc())
        .first()
    )


def _ensure_crm_workspace(db: Session, user: User, *, add_member: bool = False) -> Workspace:
    workspace = _find_crm_workspace(db)
    if workspace is None:
        workspace = Workspace(
            name=CRM_WORKSPACE_NAME,
            slug=CRM_WORKSPACE_SLUG,
            description="全公司 CRM 客户情报工作区",
            created_by=user.id,
            storage_path=str((WORKSPACES_ROOT / CUSTOMER_ROOT_NAME / CRM_WORKSPACE_SLUG).resolve()),
            brand=CUSTOMER_BRAND,
            workspace_kind="customer",
            is_default=False,
            is_hidden=True,
        )
        db.add(workspace)
        db.flush()
    workspace.name = CRM_WORKSPACE_NAME
    workspace.slug = CRM_WORKSPACE_SLUG
    workspace.brand = CUSTOMER_BRAND
    workspace.workspace_kind = "customer"
    workspace.storage_path = _ensure_storage_path(workspace)
    if add_member and not _workspace_membership(db, user.id, workspace.id):
        db.add(WorkspaceMember(workspace_id=workspace.id, user_id=user.id, role="admin"))
    db.commit()
    db.refresh(workspace)
    return workspace


def _sync_project_folders(db: Session, user: User) -> None:
    project_root = (WORKSPACES_ROOT / PROJECT_ROOT_NAME).resolve()
    if not project_root.exists():
        return
    for brand, brand_dir in _project_brand_dirs():
        if not brand_dir.exists() or not brand_dir.is_dir():
            continue
        for project_dir in sorted(brand_dir.iterdir(), key=lambda item: item.name.lower()):
            _register_existing_project_folder(db, user, brand, project_dir)


def _audit_detail(workspace_id: int, path: str = "", file_id: int | None = None, **extra) -> str:
    return json.dumps(
        {"workspace_id": workspace_id, "path": path, "file_id": file_id, **extra},
        ensure_ascii=False,
    )


def _write_workspace_audit(db: Session, user_id: int, action: str, detail: str, success: bool = True) -> None:
    db.add(AuditLog(user_id=user_id, action=action, detail=detail[:1000], success=success))


def _write_workspace_file_agent_run(
    db: Session,
    *,
    user_id: int,
    workspace: Workspace,
    source_type: str,
    title: str,
    path: str,
    status: str = "completed",
    detail: str = "",
    result: dict | None = None,
):
    payload = {
        "workspace_id": workspace.id,
        "workspace_name": workspace.name,
        "path": path,
        **(result or {}),
    }
    run = create_agent_run(
        db,
        user_id=user_id,
        workspace_id=workspace.id,
        source_type=source_type,
        source_id=str(path or workspace.id),
        title=title,
        status="running",
    )
    add_agent_event(
        db,
        run,
        event_type="permission_check",
        title="已校验项目权限",
        detail=workspace.name,
        status="completed",
        payload={"workspace_id": workspace.id},
    )
    add_agent_event(
        db,
        run,
        event_type="tool_call",
        title=title,
        detail=detail or path,
        status=status,
        payload=payload,
    )
    return finish_agent_run(
        db,
        run,
        status=status,
        result=payload,
        error_message="" if status != "failed" else detail,
    )


def _raise_with_audit(
    db: Session,
    user_id: int,
    action: str,
    status_code: int,
    message: str,
    detail: str,
) -> None:
    _write_workspace_audit(db, user_id, action, detail, success=False)
    db.commit()
    raise HTTPException(status_code=status_code, detail=message)


def _mark_workspace_rag_pending(db: Session, workspace_id: int) -> None:
    # A3 keeps ingest state at file/run granularity. Historical callers still
    # invoke this after broad file mutations, but bulk-flipping the whole
    # workspace would mark unrelated synced files as dirty.
    return


def _source_status_for_file(meta: WorkspaceFile | None, path: Path | None = None) -> str:
    if not meta:
        return "not_indexed"
    status = meta.rag_status or "new"
    if meta.deleted_at is not None:
        return "source_deleted"
    if path is None:
        return status
    if not path.exists():
        return "source_deleted" if status in {"indexed", "synced", "gbrain_ready"} else status
    if status in {"indexed", "synced", "gbrain_ready"} and meta.source_hash:
        try:
            signature = _file_signature(path)
        except OSError:
            return status
        if signature["source_hash"] != meta.source_hash:
            return "source_changed"
    return status


def _upload_limit_for(user: User, member: WorkspaceMember, workspace: Workspace | None = None) -> tuple[int, str]:
    if workspace and workspace.workspace_kind == "user":
        return MAX_WORKSPACE_UPLOAD_BYTES, f"个人工作台文件超过 {MAX_WORKSPACE_UPLOAD_MB}MB"
    return upload_limit_for(
        user,
        member,
        user_limit_bytes=MAX_WORKSPACE_UPLOAD_BYTES,
        user_limit_mb=MAX_WORKSPACE_UPLOAD_MB,
        admin_limit_bytes=MAX_WORKSPACE_ADMIN_UPLOAD_BYTES,
        admin_limit_mb=MAX_WORKSPACE_ADMIN_UPLOAD_MB,
    )


def _sync_file_descendant_paths(db: Session, workspace_id: int, old_prefix: str, new_prefix: str) -> None:
    metas = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path.like(f"{old_prefix}/%"),
        )
        .all()
    )
    for meta in metas:
        meta.relative_path = f"{new_prefix}/{meta.relative_path[len(old_prefix) + 1:]}"
        meta.rag_status = "new"
        meta.updated_at = datetime.now(timezone.utc)


def _create_copied_file_metadata(
    db: Session,
    *,
    workspace_id: int,
    source_rel_path: str,
    target_file: Path,
    root: Path,
    user_id: int,
) -> WorkspaceFile:
    target_rel_path = target_file.relative_to(root).as_posix()
    source_meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == source_rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    existing_meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == target_rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    if existing_meta:
        db.delete(existing_meta)
        db.flush()
    content_type = source_meta.content_type if source_meta else mimetypes.guess_type(target_file.name)[0] or "application/octet-stream"
    now = datetime.now(timezone.utc)
    meta = WorkspaceFile(
        workspace_id=workspace_id,
        uploaded_by=user_id,
        relative_path=target_rel_path,
        original_name=target_file.name,
        content_type=content_type[:128],
        size=target_file.stat().st_size,
        rag_status="new",
        updated_at=now,
    )
    _record_file_signature(meta, target_file)
    db.add(meta)
    db.flush()
    return meta


def _copy_descendant_file_metadata(
    db: Session,
    *,
    workspace_id: int,
    source_dir: Path,
    target_dir: Path,
    root: Path,
    user_id: int,
) -> None:
    for copied_path in sorted(target_dir.rglob("*")):
        if not copied_path.is_file():
            continue
        source_rel = (source_dir / copied_path.relative_to(target_dir)).relative_to(root).as_posix()
        _create_copied_file_metadata(
            db,
            workspace_id=workspace_id,
            source_rel_path=source_rel,
            target_file=copied_path,
            root=root,
            user_id=user_id,
        )


def _display_user_names(db: Session, user_ids: set[int]) -> dict[int, str]:
    if not user_ids:
        return {}
    users = db.query(User).filter(User.id.in_(user_ids)).all()
    return {item.id: item.nickname or item.username for item in users}


def _build_file_tree(
    root: Path,
    path: Path,
    metadata_by_path: dict[str, WorkspaceFile],
    uploader_names: dict[int, str],
    member: WorkspaceMember,
    user_id: int,
    user_role: str,
    depth: int = 0,
    max_depth: int = 3,
) -> list[WorkspaceFileItemResponse]:
    if depth >= max_depth or not path.exists():
        return []

    items: list[WorkspaceFileItemResponse] = []
    for child in sorted(path.iterdir(), key=lambda item: (not item.is_dir(), item.name.lower())):
        if child.is_symlink():
            continue
        if child.name.startswith(".") and child.name != TRASH_DIRNAME:
            continue
        stat = child.stat()
        item_type = "directory" if child.is_dir() else "file"
        rel_path = child.relative_to(root).as_posix()
        if rel_path == TRASH_DIRNAME:
            items.append(
                WorkspaceFileItemResponse(
                    id=None,
                    name=TRASH_DIRNAME,
                    path=TRASH_DIRNAME,
                    type="directory",
                    size=None,
                    updated_at=datetime.fromtimestamp(stat.st_mtime, timezone.utc),
                    can_delete=False,
                    can_restore=False,
                    children=[],
                )
            )
            continue
        meta = metadata_by_path.get(rel_path)
        items.append(
            WorkspaceFileItemResponse(
                id=meta.id if meta else None,
                name=child.name,
                path=rel_path,
                type=item_type,
                size=None if child.is_dir() else stat.st_size,
                updated_at=datetime.fromtimestamp(stat.st_mtime, timezone.utc),
                uploaded_by=meta.uploaded_by if meta else None,
                uploader_name=uploader_names.get(meta.uploaded_by) if meta else None,
                deleted_at=meta.deleted_at if meta else None,
                deleted_by=meta.deleted_by if meta else None,
                rag_status=_source_status_for_file(meta, child) if item_type == "file" else None,
                can_delete=(
                    item_type == "file" and _member_can_mutate_file(member, user_id, meta, user_role)
                ) or (
                    item_type == "directory" and not _is_template_root(Path(rel_path))
                ),
                can_restore=False,
                children=_build_file_tree(
                    root,
                    child,
                    metadata_by_path,
                    uploader_names,
                    member,
                    user_id,
                    user_role,
                    depth + 1,
                    max_depth,
                ) if child.is_dir() else [],
            )
        )
    return items


def _build_deleted_file_items(
    db: Session,
    workspace_id: int,
    member: WorkspaceMember,
    user_id: int,
    user_role: str,
) -> list[WorkspaceFileItemResponse]:
    files = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.deleted_at.is_not(None))
        .order_by(WorkspaceFile.deleted_at.desc(), WorkspaceFile.id.desc())
        .all()
    )
    names = _display_user_names(db, {item.uploaded_by for item in files} | {item.deleted_by for item in files if item.deleted_by})
    return [
        WorkspaceFileItemResponse(
            id=item.id,
            name=Path(item.relative_path).name,
            path=item.relative_path,
            type="file",
            size=item.size,
            updated_at=item.updated_at,
            uploaded_by=item.uploaded_by,
            uploader_name=names.get(item.uploaded_by),
            deleted_at=item.deleted_at,
            deleted_by=item.deleted_by,
            rag_status=_source_status_for_file(item),
            can_delete=_member_can_restore_file(member, user_id, item, user_role),
            can_restore=True,
            children=[],
        )
        for item in files
    ]


def _upsert_workspace_file(
    db: Session,
    workspace_id: int,
    user_id: int,
    rel_path: str,
    filename: str,
    content_type: str,
    size: int,
    source_path: Path | None = None,
) -> WorkspaceFile:
    existing = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    now = datetime.now(timezone.utc)
    if existing:
        existing.uploaded_by = user_id
        existing.original_name = filename
        existing.content_type = content_type[:128]
        existing.size = size
        existing.rag_status = "new"
        existing.updated_at = now
        existing.trash_path = ""
        if source_path and source_path.exists() and source_path.is_file():
            _record_file_signature(existing, source_path)
        return existing
    meta = WorkspaceFile(
        workspace_id=workspace_id,
        uploaded_by=user_id,
        relative_path=rel_path,
        original_name=filename,
        content_type=content_type[:128],
        size=size,
        rag_status="new",
        updated_at=now,
    )
    if source_path and source_path.exists() and source_path.is_file():
        _record_file_signature(meta, source_path)
    db.add(meta)
    db.flush()
    return meta


@router.post("", response_model=WorkspaceResponse)
def create_workspace(
    req: CreateWorkspaceRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="仅系统管理员可新建工作区")

    workspace_kind = _normalize_workspace_kind(req.workspace_kind, req.brand)
    if workspace_kind == "customer":
        workspace = _ensure_crm_workspace(db, user, add_member=True)
        return _workspace_response(db, workspace, user=user)

    name = req.name.strip()
    if not name:
        raise HTTPException(status_code=400, detail="工作区名称不能为空")
    if len(name) > 128:
        raise HTTPException(status_code=400, detail="工作区名称不能超过 128 个字符")

    brand = _normalize_brand(req.brand)
    slug = _slugify(name)
    existing = db.query(Workspace).filter(Workspace.slug == slug).first()
    if existing:
        raise HTTPException(status_code=409, detail="已存在同名工作区，请选择已有工作区或使用不同名称")
    target_path = _candidate_storage_path(slug, brand, workspace_kind)
    root = WORKSPACES_ROOT.resolve()
    if not target_path.is_relative_to(root):
        raise HTTPException(status_code=400, detail="工作区目录不合法")
    if workspace_kind == "project":
        existing_folder = _find_existing_project_folder(brand, slug, name)
        if existing_folder:
            workspace = _register_existing_project_folder(db, user, brand, existing_folder, add_member=True)
            if workspace:
                return _workspace_response(db, workspace, user=user)
            raise HTTPException(status_code=409, detail="后端已存在同名项目文件夹，请选择已有项目或更换项目名称")
    if target_path.exists():
        raise HTTPException(status_code=409, detail="后端已存在同名工作区文件夹，请选择已有工作区或更换名称")

    workspace = Workspace(
        name=name,
        slug=slug,
        description=req.description.strip(),
        created_by=user.id,
        brand=brand,
        workspace_kind=workspace_kind,
        is_default=False,
        is_hidden=workspace_kind == "customer",
    )
    db.add(workspace)
    db.commit()
    db.refresh(workspace)

    workspace.storage_path = _ensure_storage_path(workspace)
    db.commit()

    db.add(WorkspaceMember(workspace_id=workspace.id, user_id=user.id, role="admin"))
    db.commit()

    return _workspace_response(db, workspace, member_count=1, user=user)


@router.get("", response_model=list[WorkspaceResponse])
def list_workspaces(
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    ensure_default_workspace(db, user)
    _sync_project_folders(db, user)
    _ensure_crm_workspace(db, user)
    if user.role == "admin":
        workspaces = (
            db.query(Workspace)
            .filter(or_(Workspace.workspace_kind != "user", Workspace.created_by == user.id))
            .order_by(Workspace.is_default.desc(), Workspace.updated_at.desc())
            .all()
        )
        workspaces = [w for w in workspaces if w.workspace_kind != "customer" or w.slug == CRM_WORKSPACE_SLUG]
    else:
        member_workspace_ids = select(WorkspaceMember.workspace_id).where(WorkspaceMember.user_id == user.id)
        group_workspace_ids = select(WorkspaceGroupAccess.workspace_id).where(
            WorkspaceGroupAccess.group_name == _normalize_group_name(user.work_group)
        )
        workspaces = (
            db.query(Workspace)
            .filter(
                or_(
                    Workspace.id.in_(member_workspace_ids),
                    Workspace.workspace_kind == "project",
                    Workspace.workspace_kind == "customer",
                    Workspace.id.in_(group_workspace_ids),
                )
            )
            .order_by(Workspace.is_default.desc(), Workspace.updated_at.desc())
            .all()
        )
        workspaces = [w for w in workspaces if _can_open_workspace(db, user, w)]
        workspaces = [w for w in workspaces if w.workspace_kind != "customer" or w.slug == CRM_WORKSPACE_SLUG]
    return [_workspace_response(db, w, user=user) for w in workspaces]


@router.get("/search")
def search_workspaces(
    q: str = Query(default=""),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    brand: str | None = None,
):
    _sync_project_folders(db, user)
    _ensure_crm_workspace(db, user)
    query = q.strip()
    raw_brand = (brand or "").strip().upper()
    search_customer = raw_brand == CUSTOMER_BRAND
    normalized_brand = _normalize_brand(raw_brand) if raw_brand and not search_customer else None
    pattern = f"%{query}%" if query else "%"
    workspace_query = (
        db.query(Workspace)
        .filter(
            Workspace.workspace_kind == "customer" if search_customer else Workspace.workspace_kind.in_(("project", "customer")),
            Workspace.is_archived == False,
            Workspace.name.ilike(pattern) | Workspace.slug.ilike(pattern),
        )
    )
    if normalized_brand:
        workspace_query = workspace_query.filter(Workspace.workspace_kind == "project", Workspace.brand == normalized_brand)
    if search_customer:
        workspace_query = workspace_query.filter(Workspace.slug == CRM_WORKSPACE_SLUG)
    else:
        workspace_query = workspace_query.filter(
            or_(Workspace.workspace_kind == "project", Workspace.slug == CRM_WORKSPACE_SLUG)
        )
    if user.role != "admin":
        member_workspace_ids = select(WorkspaceMember.workspace_id).where(WorkspaceMember.user_id == user.id)
        group_workspace_ids = select(WorkspaceGroupAccess.workspace_id).where(
            WorkspaceGroupAccess.group_name == _normalize_group_name(user.work_group)
        )
        workspace_query = workspace_query.filter(
            or_(
                (Workspace.workspace_kind == "project") & (Workspace.is_hidden == False),
                Workspace.id.in_(member_workspace_ids),
                Workspace.id.in_(group_workspace_ids),
            )
        )
    workspaces = workspace_query.order_by(
        Workspace.brand.asc(),
        Workspace.updated_at.desc(),
        Workspace.name.asc(),
    ).limit(100).all()
    member_ids = {
        row.workspace_id
        for row in db.query(WorkspaceMember.workspace_id)
        .filter(WorkspaceMember.user_id == user.id)
        .all()
    }
    manageable_ids = {w.id for w in workspaces if _is_workspace_admin(db, user, w.id)}
    return [
        {
            "id": w.id,
            "name": w.name,
            "slug": w.slug,
            "description": w.description,
            "brand": w.brand,
            "workspace_kind": w.workspace_kind,
            "is_default": w.is_default,
            "is_hidden": w.is_hidden,
            "member_count": db.query(WorkspaceMember)
            .filter(WorkspaceMember.workspace_id == w.id)
            .count(),
            "is_member": w.id in member_ids,
            "can_open": _can_open_workspace(db, user, w),
            "can_rename": not w.is_default and w.id in manageable_ids,
            "can_delete": not w.is_default and w.id in manageable_ids,
            "is_archived": w.is_archived,
            "created_at": serialize_datetime_utc(w.created_at),
            "updated_at": serialize_datetime_utc(w.updated_at),
            "created_by": w.created_by,
        }
        for w in workspaces
    ]


@router.get("/{workspace_id}", response_model=WorkspaceDetailResponse)
def get_workspace(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")

    if not _can_open_workspace(db, user, workspace):
        raise HTTPException(status_code=403, detail="你无权访问该隐藏项目")
    member = _ensure_member(db, user.id, workspace_id)

    members = (
        db.query(WorkspaceMember, User)
        .join(User, WorkspaceMember.user_id == User.id)
        .filter(WorkspaceMember.workspace_id == workspace_id)
        .all()
    )
    member_list = [
        MemberResponse(
            user_id=member_user.id,
            username=member_user.username,
            nickname=member_user.nickname,
            role=wm.role,
            joined_at=wm.joined_at,
        )
        for wm, member_user in members
    ]

    can_manage = _is_workspace_admin(db, user, workspace.id)
    return WorkspaceDetailResponse(
        id=workspace.id,
        name=workspace.name,
        slug=workspace.slug,
        description=workspace.description,
        created_by=workspace.created_by,
        member_count=len(member_list),
        brand=workspace.brand,
        workspace_kind=workspace.workspace_kind,
        is_default=workspace.is_default,
        is_archived=workspace.is_archived,
        is_hidden=workspace.is_hidden,
        can_rename=not workspace.is_default and can_manage,
        can_delete=not workspace.is_default and can_manage,
        storage_path=workspace.storage_path,
        members=member_list,
        access_groups=_workspace_access_groups(db, workspace.id),
        created_at=workspace.created_at,
        updated_at=workspace.updated_at,
    )


@router.get("/{workspace_id}/files", response_model=WorkspaceFilesResponse)
def list_workspace_files(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    include_deleted: bool = False,
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")

    member = _ensure_member(db, user.id, workspace_id)
    storage_path = _ensure_storage_path(workspace)
    if workspace.storage_path != storage_path:
        workspace.storage_path = storage_path
        db.commit()

    root = Path(storage_path).resolve()
    if include_deleted:
        return WorkspaceFilesResponse(
            workspace_id=workspace.id,
            root_name=workspace.name,
            items=_build_deleted_file_items(db, workspace.id, member, user.id, user.role),
        )
    metas = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.deleted_at.is_(None))
        .all()
    )
    metadata_by_path = {item.relative_path: item for item in metas}
    uploader_names = _display_user_names(db, {item.uploaded_by for item in metas})
    return WorkspaceFilesResponse(
        workspace_id=workspace.id,
        root_name=workspace.name,
        items=_build_file_tree(root, root, metadata_by_path, uploader_names, member, user.id, user.role),
    )


@router.get("/{workspace_id}/files/content")
def get_workspace_file_content(
    workspace_id: int,
    path: str = Query(..., min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(path)
    target = _resolve_workspace_child(root, rel)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    rel_path = target.relative_to(root).as_posix()
    meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    if meta and meta.trash_path:
        raise HTTPException(status_code=404, detail="文件不存在")
    media_type = (meta.content_type if meta else None) or mimetypes.guess_type(target.name)[0] or "application/octet-stream"
    return FileResponse(target, media_type=media_type, filename=target.name)


@router.post("/{workspace_id}/files/upload", response_model=WorkspaceMultiUploadResponse)
async def upload_workspace_files(
    workspace_id: int,
    directory: str = Form(default=""),
    files: list[UploadFile] = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    limit_bytes, limit_message = _upload_limit_for(user, member, workspace)
    root = _workspace_file_root(workspace)
    rel_dir = _safe_relative_path(directory)
    _ensure_not_trash_path(rel_dir)
    target_dir = _resolve_workspace_child(root, rel_dir)
    if not target_dir.exists() or not target_dir.is_dir():
        _raise_with_audit(
            db,
            user.id,
            "workspace_file_upload",
            400,
            "目标文件夹不存在",
            _audit_detail(workspace_id, rel_dir.as_posix(), actor_id=user.id, error="target directory missing"),
        )

    responses: list[WorkspaceFileMutationResponse] = []
    uploaded_paths: list[str] = []
    for upload in files:
        filename = _safe_name(upload.filename or "untitled")
        content = await upload.read()
        if len(content) > limit_bytes:
            _raise_with_audit(
                db,
                user.id,
                "workspace_file_upload",
                400,
                limit_message,
                _audit_detail(workspace_id, (rel_dir / filename).as_posix(), actor_id=user.id, error="file too large"),
            )
        conflict_path = _resolve_conflict_path(target_dir, filename, "keep_both")
        if conflict_path is None:
            responses.append(WorkspaceFileMutationResponse(ok=False, path=(rel_dir / filename).as_posix(), rag_status="skipped"))
            continue
        target_path = _resolve_workspace_child(root, conflict_path.relative_to(root))
        if target_path.exists() and target_path.is_dir():
            _raise_with_audit(
                db,
                user.id,
                "workspace_file_upload",
                400,
                "不能覆盖文件夹",
                _audit_detail(workspace_id, (rel_dir / filename).as_posix(), actor_id=user.id, error="target is directory"),
            )
        target_path.write_bytes(content)
        rel_path = target_path.relative_to(root).as_posix()
        meta = _upsert_workspace_file(
            db,
            workspace_id,
            user.id,
            rel_path,
            filename,
            upload.content_type or "application/octet-stream",
            len(content),
            target_path,
        )
        _write_workspace_audit(
            db,
            user.id,
            "workspace_file_upload",
            _audit_detail(workspace_id, rel_path, meta.id, actor_id=user.id, size=len(content)),
        )
        responses.append(WorkspaceFileMutationResponse(ok=True, path=rel_path, file_id=meta.id, rag_status=meta.rag_status))
        uploaded_paths.append(rel_path)
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_file_upload",
        title="上传项目文件",
        path=rel_dir.as_posix(),
        detail=f"上传 {len(uploaded_paths)} 个文件",
        result={"file_count": len(uploaded_paths), "paths": uploaded_paths[:20]},
    )
    db.commit()
    return WorkspaceMultiUploadResponse(ok=True, files=responses, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/files", response_model=WorkspaceFileMutationResponse)
def upload_workspace_file(
    workspace_id: int,
    req: UploadWorkspaceFileRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    limit_bytes, limit_message = _upload_limit_for(user, member, workspace)
    root = _workspace_file_root(workspace)
    directory = _safe_relative_path(req.directory)
    _ensure_not_trash_path(directory)
    filename = _safe_name(req.filename)
    target_dir = _resolve_workspace_child(root, directory)
    if not target_dir.exists() or not target_dir.is_dir():
        raise HTTPException(status_code=400, detail="目标文件夹不存在")
    target_path = _resolve_workspace_child(root, directory / filename)
    if target_path.exists() and target_path.is_dir():
        raise HTTPException(status_code=400, detail="不能覆盖文件夹")
    try:
        content = base64.b64decode(req.content_base64, validate=True)
    except binascii.Error as exc:
        raise HTTPException(status_code=400, detail="文件内容格式不正确") from exc
    if len(content) > limit_bytes:
        _raise_with_audit(
            db,
            user.id,
            "workspace_file_upload",
            400,
            limit_message,
            _audit_detail(workspace_id, (directory / filename).as_posix(), actor_id=user.id, error="file too large"),
        )

    conflict_path = _resolve_conflict_path(target_dir, filename, req.conflict_strategy)
    if conflict_path is None:
        skipped_path = (directory / filename).as_posix()
        agent_run = _write_workspace_file_agent_run(
            db,
            user_id=user.id,
            workspace=workspace,
            source_type="workspace_file_upload",
            title="上传项目文件",
            path=skipped_path,
            status="cancelled",
            detail="目标位置已存在同名文件，按策略跳过",
            result={"rag_status": "skipped"},
        )
        db.commit()
        return WorkspaceFileMutationResponse(ok=False, path=skipped_path, rag_status="skipped", agent_run=serialize_agent_run(db, agent_run))
    target_path = _resolve_workspace_child(root, conflict_path.relative_to(root))
    if target_path.exists() and target_path.is_dir():
        raise HTTPException(status_code=400, detail="不能覆盖文件夹")
    target_path.write_bytes(content)
    rel_path = target_path.relative_to(root).as_posix()
    meta = _upsert_workspace_file(db, workspace_id, user.id, rel_path, filename, req.content_type, len(content), target_path)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_file_upload",
        _audit_detail(workspace_id, rel_path, meta.id, actor_id=user.id, size=len(content)),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_file_upload",
        title="上传项目文件",
        path=rel_path,
        detail=filename,
        result={"file_id": meta.id, "rag_status": meta.rag_status, "size": len(content)},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, file_id=meta.id, rag_status=meta.rag_status, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/folders", response_model=WorkspaceFileMutationResponse)
def create_workspace_folder(
    workspace_id: int,
    req: CreateWorkspaceFolderRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    parent_rel = _safe_relative_path(req.parent_path)
    _ensure_not_trash_path(parent_rel)
    parent = _resolve_workspace_child(root, parent_rel)
    if not parent.exists() or not parent.is_dir():
        raise HTTPException(status_code=400, detail="目标父文件夹不存在")
    folder_name = _safe_name(req.name)
    target = _resolve_workspace_child(root, parent.relative_to(root) / folder_name)
    if target.exists():
        raise HTTPException(status_code=409, detail="已存在同名文件夹")
    target.mkdir()
    rel_path = target.relative_to(root).as_posix()
    _write_workspace_audit(
        db,
        user.id,
        "workspace_folder_create",
        _audit_detail(workspace_id, rel_path, actor_id=user.id),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_folder_create",
        title="新建项目文件夹",
        path=rel_path,
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, agent_run=serialize_agent_run(db, agent_run))


# ── Meeting endpoints ────────────────────────────────────────────────────

@router.post("/{workspace_id}/meetings/folder", response_model=MeetingFolderResponse)
def create_meeting_folder(
    workspace_id: int,
    req: CreateMeetingFolderRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a single-meeting folder structure inside the workspace."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持创建会议文件夹")
    if req.meeting_type not in MEETING_TYPES:
        raise HTTPException(status_code=400, detail=f"会议类型不合法，可选值：{', '.join(MEETING_TYPES)}")

    root = _workspace_file_root(workspace)

    # Determine parent directory for meetings
    parent_rel_str = meeting_parent_path(workspace.workspace_kind)
    parent_rel = _safe_relative_path(parent_rel_str)
    _ensure_not_trash_path(parent_rel)
    parent = _resolve_workspace_child(root, parent_rel)
    parent.mkdir(parents=True, exist_ok=True)

    # Parse optional meeting time
    meeting_dt: datetime | None = None
    if req.meeting_time:
        try:
            meeting_dt = datetime.fromisoformat(req.meeting_time)
        except (ValueError, TypeError):
            raise HTTPException(status_code=400, detail="会议时间格式不合法，请使用 ISO-8601")

    folder_name = make_meeting_folder_name(meeting_dt, req.topic)
    meeting_dir = meeting_folder_collision_free(parent, folder_name)
    meeting_dir.mkdir(parents=True, exist_ok=True)

    created_dirs: list[str] = []
    for sub in MEETING_SUBDIRS:
        sub_dir = meeting_dir / sub
        sub_dir.mkdir(parents=True, exist_ok=True)
        created_dirs.append(sub_dir.relative_to(root).as_posix())

    meeting_rel = meeting_dir.relative_to(root).as_posix()
    created_dirs.insert(0, meeting_rel)

    # Write meeting metadata
    _write_meeting_meta(
        meeting_dir,
        topic=req.topic,
        meeting_time=req.meeting_time,
        meeting_type=req.meeting_type,
    )

    _write_workspace_audit(
        db,
        user.id,
        "meeting_folder_create",
        _audit_detail(
            workspace_id,
            meeting_rel,
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            meeting_folder_path=meeting_rel,
            created_dirs=created_dirs,
            gbrain_ingest=False,
        ),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="meeting_folder_create",
        title="创建会议文件夹",
        path=meeting_rel,
        detail=f"会议：{req.topic}",
    )
    db.commit()
    return MeetingFolderResponse(
        ok=True,
        meeting_folder_path=meeting_rel,
        created_dirs=created_dirs,
        created_files=[],
        gbrain_ingest=False,
        agent_run=serialize_agent_run(db, agent_run),
    )


@router.post("/{workspace_id}/meetings/transcript", response_model=SaveMeetingTranscriptResponse)
def save_meeting_transcript(
    workspace_id: int,
    req: SaveMeetingTranscriptRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Save pasted meeting text as transcript-v1.md and transcript-latest.md inside an existing meeting folder."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持保存会议转录")

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(req.folder_path)
    _ensure_not_trash_path(folder_rel)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    if not folder_dir.exists() or not folder_dir.is_dir():
        raise HTTPException(status_code=400, detail="会议文件夹不存在")

    # Validate this is a proper meeting folder with all 5 subdirectories
    missing = [sub for sub in MEETING_SUBDIRS if not (folder_dir / sub).is_dir()]
    if missing:
        raise HTTPException(status_code=400, detail="请选择会议文件夹后保存转录文本")

    # Enforce parent path whitelist per workspace kind
    expected_parent = meeting_parent_path(workspace.workspace_kind)
    folder_posix = folder_rel.as_posix()
    if folder_posix != expected_parent and not folder_posix.startswith(expected_parent + "/"):
        raise HTTPException(status_code=400, detail=f"会议文件夹必须位于 {expected_parent}/ 下")

    transcript_dir = folder_dir / "02-转录文本"

    content = req.content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="转录文本不能为空")

    # Build transcript markdown with formal template
    now_utc = datetime.now(timezone.utc)
    transcript_md = _build_transcript_markdown(
        content, now_utc,
        input_type=req.input_type,
        original_filename=req.original_filename,
    )

    # Write transcript-v1.md
    v1_path = _resolve_conflict_path(transcript_dir, "transcript-v1.md", "keep_both")
    if v1_path is None:
        raise HTTPException(status_code=500, detail="无法写入转录文件")
    v1_path.write_text(transcript_md, encoding="utf-8")
    v1_rel = v1_path.relative_to(root).as_posix()

    # Write / overwrite transcript-latest.md
    latest_path = transcript_dir / "transcript-latest.md"
    latest_path.write_text(transcript_md, encoding="utf-8")
    latest_rel = latest_path.relative_to(root).as_posix()

    # Record WorkspaceFile metadata for both files
    _upsert_workspace_file(
        db, workspace_id, user.id, v1_rel,
        "transcript-v1.md", "text/markdown", len(transcript_md.encode("utf-8")), v1_path,
    )
    _upsert_workspace_file(
        db, workspace_id, user.id, latest_rel,
        "transcript-latest.md", "text/markdown", len(transcript_md.encode("utf-8")), latest_path,
    )

    _write_workspace_audit(
        db,
        user.id,
        "meeting_transcript_save",
        _audit_detail(
            workspace_id,
            req.folder_path,
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            meeting_folder_path=req.folder_path,
            created_files=[v1_rel, latest_rel],
            gbrain_ingest=False,
        ),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="meeting_transcript_save",
        title="保存会议转录文本",
        path=req.folder_path,
        detail=f"转录：transcript-v1.md, transcript-latest.md",
    )
    db.commit()
    return SaveMeetingTranscriptResponse(
        ok=True,
        meeting_folder_path=req.folder_path,
        transcript_v1_path=v1_rel,
        transcript_latest_path=latest_rel,
        gbrain_ingest=False,
        agent_run=serialize_agent_run(db, agent_run),
    )


@router.post("/{workspace_id}/meetings/transcript/file", response_model=SaveMeetingTranscriptResponse)
async def save_meeting_transcript_from_file(
    workspace_id: int,
    folder_path: str = Form(...),
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload a TXT/MD/DOCX file as meeting transcript source."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持保存会议转录")

    filename = (file.filename or "untitled").strip()
    lower = filename.lower()

    # Determine input type and extract text
    content: str
    input_type: str
    if lower.endswith(".docx"):
        content_bytes = await file.read()
        content = _extract_text_from_docx(content_bytes, filename)
        input_type = "docx"
    elif lower.endswith(".txt"):
        content_bytes = await file.read()
        content = content_bytes.decode("utf-8", errors="replace")
        input_type = "txt"
    elif lower.endswith(".md") or lower.endswith(".markdown"):
        content_bytes = await file.read()
        content = content_bytes.decode("utf-8", errors="replace")
        input_type = "md"
    else:
        raise HTTPException(status_code=400, detail="仅支持 TXT / MD / DOCX 格式的转录文件")

    if not content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空，无法生成转录")

    # Delegate to the same logic as the paste endpoint
    return save_meeting_transcript(
        workspace_id,
        SaveMeetingTranscriptRequest(
            folder_path=folder_path,
            content=content,
            input_type=input_type,
            original_filename=filename,
        ),
        user,
        db,
    )


# ── DOCX extraction ─────────────────────────────────────────────────────

def _extract_text_from_docx(file_bytes: bytes, filename: str = "") -> str:
    """Extract plain text from a .docx file using python-docx."""
    import io
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            # Try extracting from tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX 解析组件未安装")
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"DOCX 文件解析失败：{exc}")


# ── Helper ───────────────────────────────────────────────────────────────

_TRANSCRIPT_MEDIA_INPUT_TYPES = {"mp3", "wav", "m4a", "ogg", "flac", "mp4", "mov", "avi", "wmv", "mkv", "webm"}


def _transcript_source_label(input_type: str, original_filename: str = "") -> str:
    suffix = f"（{original_filename}）" if original_filename else ""
    normalized = (input_type or "paste").strip().lower()
    if normalized == "paste":
        return "用户粘贴文本"
    if normalized == "txt":
        return f"TXT 上传{suffix}"
    if normalized == "md":
        return f"MD 上传{suffix}"
    if normalized in ("docx", "doc"):
        return f"DOCX 上传{suffix}"
    if normalized in _TRANSCRIPT_MEDIA_INPUT_TYPES:
        return f"音视频自动转录{suffix}"
    return f"{normalized.upper()} 输入{suffix}" if normalized else f"文件输入{suffix}"


def _transcript_metadata_value(transcript_text: str, field_name: str) -> str:
    pattern = re.compile(rf"^\|\s*{re.escape(field_name)}\s*\|\s*(.*?)\s*\|\s*$", re.MULTILINE)
    match = pattern.search(transcript_text or "")
    return match.group(1).strip() if match else ""


def _build_transcript_markdown(
    raw_text: str,
    now: datetime,
    input_type: str = "paste",
    original_filename: str = "",
    transcription_status: str = "completed",
    warnings: list[str] | None = None,
) -> str:
    """Build the formal five-section transcript template."""
    ts = now.strftime("%Y-%m-%d %H:%M UTC")
    source_label = _transcript_source_label(input_type, original_filename)

    # ── Basic speaker detection ──────────────────────────────────────
    speakers, segments = _detect_speakers(raw_text)
    speaker_count = len(speakers)

    # ── Speaker overview table ───────────────────────────────────────
    speaker_rows: list[str] = []
    for sp in speakers:
        speaker_rows.append(
            f"| {sp['id']} | {sp['label']} | 未映射 "
            f"| {sp.get('ratio','—')} "
            f"| {sp.get('duration','—')} "
            f"| 待确认 |"
        )

    # ── Full transcript table ────────────────────────────────────────
    transcript_rows: list[str] = []
    for seg in segments:
        transcript_rows.append(
            f"| {seg.get('line','—')} "
            f"| {seg['time']} "
            f"| {seg['speaker_id']} "
            f"| {seg['speaker_label']} "
            f"| {seg['content']} "
            f"| {seg.get('confidence','—')} "
            f"| {seg.get('flag','—')} |"
        )

    # ── Timeline ─────────────────────────────────────────────────────
    timeline_rows: list[str] = []
    for seg in segments[:20]:
        summary = seg['content'][:40].replace("\n", " ").replace("|", "/")
        timeline_rows.append(f"| {seg.get('line','—')} | {seg['time']} | {seg['speaker_id']} | {summary} |")

    return (
        "# 会议转录文本\n\n"
        "## 基本信息\n\n"
        f"| 字段 | 值 |\n"
        f"|---|---|\n"
        f"| 转录时间 | {ts} |\n"
        f"| 转录来源 | {source_label} |\n"
        f"| 输入类型 | {input_type} |\n"
        f"| 原始文件名 | {original_filename or '—'} |\n"
        f"| 转录状态 | {transcription_status or 'completed'} |\n"
        f"| 缺失片段 | {_escape_pipe('; '.join(warnings or []) if transcription_status == 'partial' else '—')} |\n"
        f"| 检测说话人数 | {speaker_count} |\n"
        "\n"
        "## 说话人概览\n\n"
        "| 说话人ID | 显示名称 | 映射状态 | 发言占比 | 发言时长 | 备注 |\n"
        "|---|---|---|---|---|---|\n"
        + "\n".join(speaker_rows) + "\n"
        "\n"
        "## 说话人时间轴\n\n"
        "| 行号 | 时间点 | 说话人ID | 内容摘要 |\n"
        "|---|---|---|---|\n"
        + "\n".join(timeline_rows) + "\n"
        "\n"
        "## 疑似术语纠错\n\n"
        "| 原识别 | 建议修正 | 类型 | 置信度 | 来源时间点 |\n"
        "|---|---|---|---|---|\n"
        "| — | — | — | — | — |\n"
        "\n"
        "## 完整转录\n\n"
        "| 行号 | 时间点 | 说话人ID | 显示名称 | 内容 | 置信度 | 标记 |\n"
        "|---|---|---|---|---|---|---|\n"
        + "\n".join(transcript_rows) + "\n"
        "\n"
        "---\n"
        "*本转录由 Project_R 自动生成。说话人映射和术语纠错为初始结果，请人工复核。*\n"
    )


def _detect_speakers(text: str) -> tuple[list[dict], list[dict]]:
    """Basic speaker detection from text patterns.

    Looks for:
    - `Name:` or `Name：` (Chinese colon) at line start
    - `Speaker N:` patterns
    - `[Name]` bracketed speaker tags

    If no patterns found, treats entire text as Speaker 1.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return (
            [{"id": "Speaker 1", "label": "Speaker 1", "ratio": "100%", "duration": "—"}],
            [{"line": 1, "time": "—", "speaker_id": "Speaker 1", "speaker_label": "Speaker 1", "content": text[:200], "confidence": "—", "flag": "—"}],
        )

    # Pattern 1: "Name:" or "Name：" at line start
    speaker_pattern = re.compile(r"^(.{1,30}?)[：:]\s*(.+)", re.UNICODE)
    # Pattern 2: "Speaker N" pattern
    speaker_n_pattern = re.compile(r"^(Speaker\s*\d+|发言人\s*[A-Za-z\d]+)\s*[：:]\s*(.+)", re.IGNORECASE)
    # Pattern 3: "[Name]" bracketed
    bracket_pattern = re.compile(r"^\[(.{1,30}?)\]\s*(.+)", re.UNICODE)

    segments: list[dict] = []
    speaker_ids: dict[str, str] = {}  # label → speaker_id
    speaker_line_counts: dict[str, int] = {}
    next_speaker_index = 1

    for line in lines:
        match = speaker_n_pattern.match(line) or bracket_pattern.match(line) or speaker_pattern.match(line)
        if match:
            raw_label = match.group(1).strip()
            content = match.group(2).strip()
        else:
            raw_label = ""
            content = line

        if raw_label:
            if raw_label not in speaker_ids:
                sid = f"Speaker {next_speaker_index}"
                speaker_ids[raw_label] = sid
                speaker_line_counts[sid] = 0
                next_speaker_index += 1
            sid = speaker_ids[raw_label]
        else:
            # Continuation — attribute to last speaker
            sid = segments[-1]["speaker_id"] if segments else f"Speaker {next_speaker_index}"
            if sid not in speaker_line_counts:
                speaker_ids[f"Speaker {next_speaker_index}"] = sid
                speaker_line_counts[sid] = 0
                next_speaker_index += 1

        speaker_line_counts[sid] = speaker_line_counts.get(sid, 0) + 1
        # Truncate long content for table cells
        cell_content = content[:200].replace("\n", " ").replace("|", "/")
        segments.append({
            "line": len(segments) + 1,
            "time": "—",
            "speaker_id": sid,
            "speaker_label": raw_label or sid,
            "content": cell_content,
            "confidence": "—",
            "flag": "—" if raw_label else "待确认",
        })

    total = sum(speaker_line_counts.values()) or 1
    speakers: list[dict] = []
    for label, sid in speaker_ids.items():
        count = speaker_line_counts.get(sid, 0)
        speakers.append({
            "id": sid,
            "label": label,
            "ratio": f"{round(count / total * 100)}%",
            "duration": "—",
        })

    if not speakers:
        speakers = [{"id": "Speaker 1", "label": "Speaker 1", "ratio": "100%", "duration": "—"}]
    if not segments:
        segments = [{"line": 1, "time": "—", "speaker_id": "Speaker 1", "speaker_label": "Speaker 1", "content": text[:200], "confidence": "—", "flag": "—"}]

    return speakers, segments


def _read_meeting_meta(folder_dir: Path) -> dict[str, str]:
    """Read meeting metadata from .meeting-meta.json in the meeting folder root.
    Returns a dict with keys like 'meeting_type', 'topic', 'meeting_time'.
    Defaults to empty values if the file doesn't exist."""
    meta_path = folder_dir / MEETING_TYPE_META_FILENAME
    if not meta_path.exists():
        return {}
    try:
        data = json.loads(meta_path.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            return {k: str(v) for k, v in data.items()}
    except (json.JSONDecodeError, OSError):
        pass
    return {}


def _write_meeting_meta(folder_dir: Path, *, topic: str, meeting_time: str | None, meeting_type: str) -> None:
    """Write meeting metadata to .meeting-meta.json in the meeting folder root."""
    meta_path = folder_dir / MEETING_TYPE_META_FILENAME
    data: dict[str, str] = {
        "topic": topic,
        "meeting_type": meeting_type,
    }
    if meeting_time:
        data["meeting_time"] = meeting_time
    meta_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


# ── Step 3: Meeting minutes & actions generation ─────────────────────────

class MeetingGenerateRequest(BaseModel):
    folder_path: str
    regenerate: bool = False  # if True, create a new version even if already generated
    allow_partial: bool = True


class MeetingGenerateResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    minutes_v_path: str
    minutes_latest_path: str
    actions_v_path: str
    actions_latest_path: str
    gbrain_ingest: bool = False
    agent_run: AgentRunResponse | None = None
    model_used: str = ""
    token_cost: int = 0


def _next_version_number(dir_path: Path, prefix: str) -> int:
    """Find the next version number for a file series like prefix-v1.md, prefix-v2.md."""
    highest = 0
    if not dir_path.exists():
        return 1
    pattern = re.compile(rf"^{re.escape(prefix)}-v(\d+)\.md$", re.IGNORECASE)
    for child in dir_path.iterdir():
        if child.is_file():
            m = pattern.match(child.name)
            if m:
                highest = max(highest, int(m.group(1)))
    return highest + 1


_MEETING_SYSTEM_PROMPT = """你是 Project_R 的企业会议纪要助手。你擅长从中文/英文会议转录文本中提取关键信息，生成结构化的会议纪要和行动项。

规则：
1. 只根据转录文本中的依据生成内容，不得编造。
2. 没有明确负责人的行动项，标记为「待确认」。
3. 没有明确截止时间的行动项，截止时间写「待确认」。
4. 没有明确依据的决策、风险、问题，标记为「待确认」。
5. 使用中文输出，专业、简洁、可操作。
6. 辅助总结只能作为整理参考；与一手转录冲突时，以一手转录为准，并把辅助总结独有信息标记为待确认或注明来源。
7. 输出格式为标准的 Markdown，严格按照用户要求的模板分段。"""


def _build_minutes_prompt(
    transcript_text: str,
    speaker_map_text: str | None = None,
    term_corrections_text: str | None = None,
    auxiliary_summaries_text: str | None = None,
    meeting_type: str | None = None,
) -> str:
    """Build the user prompt for minutes generation from transcript."""
    transcript_source = _transcript_metadata_value(transcript_text, "转录来源") or "从转录文本基本信息读取，无法判断写「待确认」"
    sections = [
        "# 会议纪要生成",
        "",
        "请根据以下会议转录文本生成正式会议纪要。",
        "如果提供了说话人映射，请使用真实名称而非 Speaker ID。",
    ]
    if speaker_map_text:
        sections.append("\n## 说话人映射参考\n\n" + speaker_map_text)
    if term_corrections_text:
        sections.append("\n## 术语纠错参考\n\n" + term_corrections_text)
    if auxiliary_summaries_text:
        sections.append(
            "\n## 辅助总结参考\n\n"
            "以下材料来自同一会议资料目录的辅助总结，只能作为二级参考；"
            "关键结论必须优先回到一手转录文本，辅助总结独有内容需标注来源或待确认。\n\n"
            + auxiliary_summaries_text
        )
    sections.append("\n## 会议转录文本\n\n" + transcript_text)
    sections.append(
        f"""

## 输出模板

请严格按以下 Markdown 模板输出。不得省略任何段落。没有内容时写「—」或「无」。

### 会议基本信息
| 字段 | 值 |
|---|---|
| 会议主题 | （从内容推断，如无法推断写「待确认」） |
| 会议时间 | （从内容或文件名推断，如无法推断写「待确认」） |
| 参会人 | （列出检测到的说话人，如无法推断写「待确认」） |
| 会议类型 | {meeting_type or '其他'} |
| 转录来源 | {transcript_source} |

### 一句话结论
（用一句话概括会议最核心的结论或决定）

### 会议摘要
（按议题或话题组织，每个议题包含：议题名称、讨论内容、结论）

### 关键决策
| ID | 决策 | 决策背景 | 影响范围 | 来源时间点 | 依据摘录 | 置信度 | 待确认 |
|---|---|---|---|---|---|---|---|
| D1 | ... | ... | ... | 00:00:00/待确认 | ... | 高/中/低 | 是/否 |

### 行动项
| ID | 行动项 | 负责人 | 协作人 | 截止时间 | 优先级 | 状态 | 来源时间点 | 待确认 |
|---|---|---|---|---|---|---|---|---|
| A1 | ... | ...（无则写待确认） | ...（无则写—） | ...（无则写待确认） | 高/中/低 | 待确认/待执行/已完成/已取消 | 00:00:00/待确认 | 是/否 |

### 风险与问题
| ID | 风险或问题 | 类型 | 影响 | 建议下一步 | 负责人 | 来源时间点 | 严重度 |
|---|---|---|---|---|---|---|---|
| R1 | ... | 技术/工期/成本/商务/客户/资料缺口 | ... | ... | ... | 00:00:00/待确认 | 高/中/低 |

### 待确认事项
| ID | 待确认事项 | 为什么需要确认 | 建议确认对象 | 来源时间点 |
|---|---|---|---|---|
| Q1 | ... | ... | ... | 00:00:00/待确认 |

### 资料与证据
| ID | 资料类型 | 文件或来源 | 来源时间点 | 依据摘录 | 说明 |
|---|---|---|---|---|---|
| E1 | 一手转录/辅助总结/用户补充/原始音视频 | ... | 00:00:00/待确认 | ... | ... |

### 可沉淀知识候选
（如果有可以沉淀为公司规则、项目经验或流程改进的知识，列出候选。如果没有写「无」）
- 类型：公司规则候选 / 项目经验候选 / 流程改进候选 / 模板候选
- 内容：...

### 生成说明
- 生成时间：当前时间
- 转录来源：{transcript_source}
- 使用模型：DeepSeek Flash
- 说话人映射：未使用 / 已使用
- 待确认项目：N 项

"""
    )
    return "\n".join(sections)


def _build_actions_prompt(
    transcript_text: str,
    speaker_map_text: str | None = None,
    term_corrections_text: str | None = None,
    auxiliary_summaries_text: str | None = None,
) -> str:
    """Build the user prompt for action items generation from transcript."""
    sections = [
        "# 行动项生成",
        "",
        "请根据以下会议转录文本提取行动项。",
        "如果提供了说话人映射，请使用真实名称而非 Speaker ID。",
    ]
    if speaker_map_text:
        sections.append("\n## 说话人映射参考\n\n" + speaker_map_text)
    if term_corrections_text:
        sections.append("\n## 术语纠错参考\n\n" + term_corrections_text)
    if auxiliary_summaries_text:
        sections.append(
            "\n## 辅助总结参考\n\n"
            "以下材料来自同一会议资料目录的辅助总结，只能作为二级参考；"
            "无法在一手转录中确认的行动项必须标记为待确认，并注明来源为辅助总结。\n\n"
            + auxiliary_summaries_text
        )
    sections.append("\n## 会议转录文本\n\n" + transcript_text)
    sections.append(
        """

## 输出模板

请严格按以下 Markdown 模板输出。不得省略任何段落。没有行动项时写「无」。

### 基本信息
| 字段 | 值 |
|---|---|
| 来源会议 | （自动填入） |
| 提取时间 | （当前时间） |
| 行动项总数 | N |

### 行动项总览
| 状态 | 数量 |
|---|---|
| 待确认 | N |
| 待执行 | N |
| 已完成 | 0 |
| 已取消 | 0 |

### 行动项清单
| ID | 状态 | 优先级 | 行动项 | 负责人 | 协作人 | 截止时间 | 依赖条件 | 来源时间点 | 依据摘录 | 待确认原因 |
|---|---|---|---|---|---|---|---|---|---|---|
| A1 | 待确认/待执行/已完成/已取消 | 高/中/低 | ... | ...（无则写待确认） | ...（无则写—） | ...（无则写待确认） | ...（无则写—） | 00:00:00/待确认 | ... | ...（无则写—） |

### 按负责人分组
（用二级标题列出每位负责人的行动项）

### 待确认行动项
（单独列出所有标记为「待确认」的行动项）

### 生成说明
- 生成时间：当前时间
- 使用模型：DeepSeek Flash
- 待确认项目：N 项
- 注意：行动项仅供参考，请人工复核后执行

"""
    )
    return "\n".join(sections)


def _read_file_safe(path: Path) -> str:
    """Read a file as string, return empty if missing."""
    if not path.exists() or not path.is_file():
        return ""
    return path.read_text(encoding="utf-8")


def _failed_transcript_reason(transcript_text: str) -> str:
    if "转录失败" not in transcript_text and "transcription_status: failed" not in transcript_text.lower():
        return ""
    match = re.search(r"\*\*错误\*\*[：:]\s*(.+)", transcript_text)
    if match:
        return match.group(1).strip()
    match = re.search(r"错误[：:]\s*(.+)", transcript_text)
    if match:
        return match.group(1).strip()
    return "音视频转录失败"


def _transcript_status_value(transcript_text: str) -> str:
    if _failed_transcript_reason(transcript_text):
        return "failed"
    statuses = [
        match.group(1).strip().lower()
        for match in re.finditer(r"^\|\s*转录状态\s*\|\s*(.*?)\s*\|\s*$", transcript_text or "", re.MULTILINE)
    ]
    if "partial" in statuses:
        return "partial"
    status = statuses[0] if statuses else ""
    if status:
        return status
    if "partial" in (transcript_text or "").lower() or "部分转录" in (transcript_text or ""):
        return "partial"
    return "completed"


def _partial_transcript_notice(transcript_text: str) -> str:
    missing = _transcript_metadata_value(transcript_text, "缺失片段") or "存在未成功转录的片段，具体时间段待确认"
    return f"| Q-PARTIAL | 转录不完整 | {missing}，纪要和行动项可能缺失上下文 | 会议组织者 | 待确认 |\n"


def _meeting_run_lock_path(root: Path, folder_dir: Path) -> Path:
    return folder_dir / ".project-r-meeting-processing.lock"


def _acquire_meeting_run_lock(root: Path, folder_dir: Path, *, operation: str, user_id: int) -> Path:
    lock_path = _meeting_run_lock_path(root, folder_dir)
    if lock_path.exists():
        raise HTTPException(status_code=409, detail="当前会议已有处理中任务，请等待完成后再操作")
    lock_path.write_text(
        json.dumps(
            {
                "operation": operation,
                "user_id": user_id,
                "created_at": serialize_datetime_utc(datetime.now(timezone.utc)),
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return lock_path


def _release_meeting_run_lock(lock_path: Path | None) -> None:
    if lock_path and lock_path.exists():
        try:
            lock_path.unlink()
        except OSError:
            pass


def _workspace_file_uploader(db: Session, workspace_id: int, rel_path: str) -> int | None:
    meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    return meta.uploaded_by if meta else None


def _notify_meeting_run_finished(
    db: Session,
    *,
    workspace: Workspace,
    actor_user_id: int,
    folder_path: str,
    title: str,
    status: str,
    detail: str,
) -> None:
    severity = "success" if status == "completed" else "warning" if status == "partial" else "critical"
    notify_user(
        db,
        actor_user_id,
        category="workspace",
        severity=severity,
        title=title,
        content=f"{workspace.name}：{detail}",
        action_status="none" if status == "completed" else "pending",
        action_kind="open_workspace",
        action_payload={"workspace_id": workspace.id, "path": folder_path},
        event_key=f"workspace:{workspace.id}:meeting:{folder_path}:{title}:{datetime.now(timezone.utc).timestamp()}",
    )


_AUXILIARY_SUMMARY_EXTENSIONS = {".md", ".txt", ".docx"}


def _filename_match_tokens(filename: str) -> set[str]:
    stem = Path(filename or "").stem.lower()
    raw_tokens = re.findall(r"[a-z0-9]+|[\u4e00-\u9fff]{2,}", stem)
    ignored = {"audio", "video", "meeting", "summary", "transcript", "纪要", "总结", "转录", "会议"}
    return {token for token in raw_tokens if len(token) >= 2 and token not in ignored}


def _read_auxiliary_summaries(folder_dir: Path, source_filename: str = "", max_chars: int = 20000) -> str:
    summary_dir = folder_dir / "03-辅助总结"
    if not summary_dir.exists() or not summary_dir.is_dir():
        return ""

    source_tokens = _filename_match_tokens(source_filename)
    candidates = [
        child for child in sorted(summary_dir.iterdir(), key=lambda item: item.name.lower())
        if child.is_file() and not child.name.startswith("~$") and child.suffix.lower() in _AUXILIARY_SUMMARY_EXTENSIONS
    ]
    if source_tokens:
        matched = [
            child for child in candidates
            if source_tokens.intersection(_filename_match_tokens(child.name))
        ]
        candidates = matched

    sections: list[str] = []
    total = 0
    for child in candidates:
        suffix = child.suffix.lower()
        try:
            if suffix == ".docx":
                text = _extract_text_from_docx(child.read_bytes(), child.name)
            else:
                text = child.read_text(encoding="utf-8")
        except Exception as exc:
            sections.append(f"### {child.name}\n\n> 辅助总结读取失败：{exc}")
            continue
        text = text.strip()
        if not text:
            continue
        remaining = max_chars - total
        if remaining <= 0:
            break
        clipped = text[:remaining]
        total += len(clipped)
        suffix_note = "\n\n> 已截断，仅保留前部内容。" if len(text) > len(clipped) else ""
        sections.append(f"### {child.name}\n\n{clipped}{suffix_note}")
    return "\n\n".join(sections)


def _parse_table_row_count(markdown_text: str, section_header: str) -> int:
    """Count data rows in a table section (headers and separators excluded)."""
    lines = markdown_text.split("\n")
    in_section = False
    data_count = 0
    for line in lines:
        if section_header in line and line.startswith("|"):
            in_section = True
            continue
        if in_section:
            if "---" in line:
                continue
            if not line.startswith("|"):
                break
            data_count += 1
    return data_count


@router.post("/{workspace_id}/meetings/generate", response_model=MeetingGenerateResponse)
def generate_meeting_minutes_and_actions(
    workspace_id: int,
    req: MeetingGenerateRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate meeting minutes and action items from transcript via LLM."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持生成会议纪要")
    if workspace.workspace_kind not in ("project", "customer"):
        raise HTTPException(status_code=400, detail="不支持的工作区类型")

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(req.folder_path)
    _ensure_not_trash_path(folder_rel)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    if not folder_dir.exists() or not folder_dir.is_dir():
        raise HTTPException(status_code=400, detail="会议文件夹不存在")

    # Validate meeting folder structure
    missing = [sub for sub in MEETING_SUBDIRS if not (folder_dir / sub).is_dir()]
    if missing:
        raise HTTPException(status_code=400, detail="请选择会议文件夹后生成纪要")

    # Enforce parent path whitelist
    expected_parent = meeting_parent_path(workspace.workspace_kind)
    folder_posix = folder_rel.as_posix()
    if folder_posix != expected_parent and not folder_posix.startswith(expected_parent + "/"):
        raise HTTPException(status_code=400, detail=f"会议文件夹必须位于 {expected_parent}/ 下")

    # Read transcript
    transcript_dir = folder_dir / "02-转录文本"
    transcript_path = transcript_dir / "transcript-latest.md"
    if not transcript_path.exists():
        raise HTTPException(status_code=400, detail="转录文件（transcript-latest.md）不存在，请先保存转录")
    transcript_rel = transcript_path.relative_to(root).as_posix()
    transcript_uploaded_by = _workspace_file_uploader(db, workspace_id, transcript_rel)
    transcript_text = transcript_path.read_text(encoding="utf-8")
    if not transcript_text.strip():
        raise HTTPException(status_code=400, detail="转录文件为空")
    failed_reason = _failed_transcript_reason(transcript_text)
    if failed_reason:
        raise HTTPException(status_code=400, detail=f"转录未成功，不能生成会议纪要。原因：{failed_reason}")
    transcription_status = _transcript_status_value(transcript_text)
    if transcription_status == "partial" and not req.allow_partial:
        raise HTTPException(status_code=400, detail="转录状态为 partial，请确认允许基于成功片段生成纪要后再继续")
    lock_path = _acquire_meeting_run_lock(root, folder_dir, operation="generate_minutes", user_id=user.id)

    # Read optional speaker map and term corrections
    speaker_map_path = folder_dir / "02-转录文本" / "speaker-map-latest.md"
    speaker_map_text = _read_file_safe(speaker_map_path)
    term_corrections_path = folder_dir / "02-转录文本" / "term-corrections-latest.md"
    term_corrections_text = _read_file_safe(term_corrections_path)
    original_filename = _transcript_metadata_value(transcript_text, "原始文件名")
    auxiliary_summaries_text = _read_auxiliary_summaries(folder_dir, source_filename=original_filename)

    # Determine version numbers
    minutes_dir = folder_dir / "04-会议纪要"
    actions_dir = folder_dir / "05-行动项"
    minutes_dir.mkdir(parents=True, exist_ok=True)
    actions_dir.mkdir(parents=True, exist_ok=True)

    minutes_ver = _next_version_number(minutes_dir, "minutes")
    actions_ver = _next_version_number(actions_dir, "actions")

    # If regenerate is False and already exists, don't overwrite
    if not req.regenerate:
        if (minutes_dir / "minutes-latest.md").exists() or (actions_dir / "actions-latest.md").exists():
            raise HTTPException(
                status_code=409,
                detail="已存在纪要与行动项。如需重新生成，请设置 regenerate=True 或先删除已有文件",
            )

    try:
        # LLM generation
        from app.shared.llm.client import get_llm_client

        minutes_md: str = ""
        actions_md: str = ""
        token_input: int = 0
        token_output: int = 0
        model_used: str = "template-fallback"

        meeting_type = _read_meeting_meta(folder_dir).get("meeting_type", "")

        _minutes_prompt = _build_minutes_prompt(
            transcript_text,
            speaker_map_text,
            term_corrections_text,
            auxiliary_summaries_text,
            meeting_type=meeting_type,
        )
        _actions_prompt = _build_actions_prompt(
            transcript_text,
            speaker_map_text,
            term_corrections_text,
            auxiliary_summaries_text,
        )

        client = get_llm_client("deepseek-flash")
        model_used = "deepseek-flash"

        # Generate minutes
        minutes_response = client.complete(
            [{"role": "user", "content": _minutes_prompt}],
            system_prompt=_MEETING_SYSTEM_PROMPT,
            temperature=0.3,
        )
        minutes_md = minutes_response.text.strip() if minutes_response.text else ""
        if minutes_response.usage:
            token_input += minutes_response.usage.get("input_tokens", 0)
            token_output += minutes_response.usage.get("output_tokens", 0)

        # Generate actions
        actions_response = client.complete(
            [{"role": "user", "content": _actions_prompt}],
            system_prompt=_MEETING_SYSTEM_PROMPT,
            temperature=0.3,
        )
        actions_md = actions_response.text.strip() if actions_response.text else ""
        if actions_response.usage:
            token_input += actions_response.usage.get("input_tokens", 0)
            token_output += actions_response.usage.get("output_tokens", 0)

    except HTTPException:
        raise
    except Exception as exc:
        # If LLM fails, fall back to template-based generation
        model_used = "template-fallback"
        from app.shared.time.utils import serialize_datetime_utc
        now_ts = serialize_datetime_utc(datetime.now(timezone.utc))
        minutes_md = _build_fallback_minutes(transcript_text, now_ts, str(exc))
        actions_md = _build_fallback_actions(now_ts)

    finally:
        _release_meeting_run_lock(lock_path)

    if transcription_status == "partial":
        partial_block = (
            "\n\n> 转录状态：partial。以下纪要和行动项仅基于已成功转录片段生成。"
            "缺失片段可能导致结论不完整，请人工复核。\n"
        )
        if "转录状态：partial" not in minutes_md:
            minutes_md = minutes_md.rstrip() + partial_block

    # Save minutes
    minutes_v_path = minutes_dir / f"minutes-v{minutes_ver}.md"
    minutes_v_path.write_text(minutes_md, encoding="utf-8")
    minutes_v_rel = minutes_v_path.relative_to(root).as_posix()

    minutes_latest_path = minutes_dir / "minutes-latest.md"
    minutes_latest_path.write_text(minutes_md, encoding="utf-8")
    minutes_latest_rel = minutes_latest_path.relative_to(root).as_posix()

    # Save actions
    actions_v_path = actions_dir / f"actions-v{actions_ver}.md"
    actions_v_path.write_text(actions_md, encoding="utf-8")
    actions_v_rel = actions_v_path.relative_to(root).as_posix()

    actions_latest_path = actions_dir / "actions-latest.md"
    actions_latest_path.write_text(actions_md, encoding="utf-8")
    actions_latest_rel = actions_latest_path.relative_to(root).as_posix()

    # Record WorkspaceFile metadata
    _upsert_workspace_file(
        db, workspace_id, user.id, minutes_v_rel,
        f"minutes-v{minutes_ver}.md", "text/markdown", len(minutes_md.encode("utf-8")), minutes_v_path,
    ).rag_status = "partial" if transcription_status == "partial" else "not_ingested"
    # Mark previously synced files as needs_reingest
    for prefix_dir, prefix_name in [(minutes_dir, "minutes"), (actions_dir, "actions")]:
        if not prefix_dir.exists():
            continue
        for child in prefix_dir.iterdir():
            if child.is_file() and child.name.startswith(prefix_name) and child != minutes_v_path and child != minutes_latest_path:
                cr = child.relative_to(root).as_posix()
                cmeta = db.query(WorkspaceFile).filter(
                    WorkspaceFile.workspace_id == workspace_id,
                    WorkspaceFile.relative_path == cr).first()
                if cmeta and cmeta.rag_status in ("synced", "gbrain_ready", "sync_pending"):
                    cmeta.rag_status = "needs_reingest"

    _upsert_workspace_file(
        db, workspace_id, user.id, minutes_latest_rel,
        "minutes-latest.md", "text/markdown", len(minutes_md.encode("utf-8")), minutes_latest_path,
    ).rag_status = "partial" if transcription_status == "partial" else "not_ingested"
    _upsert_workspace_file(
        db, workspace_id, user.id, actions_v_rel,
        f"actions-v{actions_ver}.md", "text/markdown", len(actions_md.encode("utf-8")), actions_v_path,
    ).rag_status = "partial" if transcription_status == "partial" else "not_ingested"
    _upsert_workspace_file(
        db, workspace_id, user.id, actions_latest_rel,
        "actions-latest.md", "text/markdown", len(actions_md.encode("utf-8")), actions_latest_path,
    ).rag_status = "partial" if transcription_status == "partial" else "not_ingested"

    total_tokens = token_input + token_output

    _write_workspace_audit(
        db,
        user.id,
        "meeting_minutes_generate",
        _audit_detail(
            workspace_id,
            req.folder_path,
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            meeting_folder_path=req.folder_path,
            input_files=[{"path": transcript_rel, "uploaded_by": transcript_uploaded_by}],
            run_by=user.id,
            created_files=[minutes_v_rel, minutes_latest_rel, actions_v_rel, actions_latest_rel],
            model=model_used,
            token_cost=total_tokens,
            transcription_status=transcription_status,
            gbrain_ingest=False,
        ),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="meeting_minutes_generate",
        title="生成会议纪要与行动项",
        path=req.folder_path,
        detail=f"纪要：minutes-v{minutes_ver}.md, actions-v{actions_ver}.md",
        status="completed" if transcription_status != "partial" else "completed",
        result={"transcription_status": transcription_status},
    )
    _notify_meeting_run_finished(
        db,
        workspace=workspace,
        actor_user_id=user.id,
        folder_path=req.folder_path,
        title="会议纪要生成完成" if transcription_status != "partial" else "会议纪要基于部分转录生成",
        status="completed" if transcription_status != "partial" else "partial",
        detail=f"{req.folder_path} 已生成 minutes-v{minutes_ver}.md / actions-v{actions_ver}.md",
    )
    db.commit()
    return MeetingGenerateResponse(
        ok=True,
        meeting_folder_path=req.folder_path,
        minutes_v_path=minutes_v_rel,
        minutes_latest_path=minutes_latest_rel,
        actions_v_path=actions_v_rel,
        actions_latest_path=actions_latest_rel,
        gbrain_ingest=False,
        agent_run=serialize_agent_run(db, agent_run),
        model_used=model_used,
        token_cost=total_tokens,
    )


def _build_fallback_minutes(transcript_text: str, timestamp: str, error: str = "") -> str:
    """Template-based fallback when LLM is unavailable."""
    transcript_source = _transcript_metadata_value(transcript_text, "转录来源") or "待确认"
    transcription_status = _transcript_status_value(transcript_text)
    partial_note = _partial_transcript_notice(transcript_text) if transcription_status == "partial" else ""
    return f"""# 会议纪要

## 会议基本信息

| 字段 | 值 |
|---|---|
| 会议主题 | 待确认 |
| 会议时间 | 待确认 |
| 参会人 | 待确认 |
| 会议类型 | 其他 |
| 转录来源 | {transcript_source} |
| 转录状态 | {transcription_status} |

## 一句话结论

待确认（LLM 暂不可用：{error}）

## 会议摘要

待确认

## 关键决策

| ID | 决策 | 决策背景 | 影响范围 | 来源时间点 | 依据摘录 | 置信度 | 待确认 |
|---|---|---|---|---|---|---|---|
| — | — | — | — | — | — | — | 是 |

## 行动项

| ID | 行动项 | 负责人 | 协作人 | 截止时间 | 优先级 | 状态 | 来源时间点 | 待确认 |
|---|---|---|---|---|---|---|---|---|
| — | — | 待确认 | — | 待确认 | — | 待确认 | — | 是 |

## 风险与问题

| ID | 风险或问题 | 类型 | 影响 | 建议下一步 | 负责人 | 来源时间点 | 严重度 |
|---|---|---|---|---|---|---|---|
| — | — | 资料缺口 | — | 人工复核转录和纪要 | 待确认 | — | 中 |

## 待确认事项

| ID | 待确认事项 | 为什么需要确认 | 建议确认对象 | 来源时间点 |
|---|---|---|---|---|
| Q1 | 全部内容 | LLM 暂不可用，请人工编写纪要 | 会议组织者 | — |
{partial_note}

## 资料与证据

| ID | 资料类型 | 文件或来源 | 来源时间点 | 依据摘录 | 说明 |
|---|---|---|---|---|---|
| E1 | 一手转录 | {transcript_source} | — | — | fallback 仅保留证据入口，需人工复核 |

## 可沉淀知识候选

无（LLM 暂不可用）

## 生成说明

- 生成时间：{timestamp}
- 转录来源：{transcript_source}
- 转录状态：{transcription_status}
- 使用模型：template-fallback
- 待确认项目：全部
"""


def _build_fallback_actions(timestamp: str) -> str:
    """Template-based fallback for action items when LLM is unavailable."""
    return f"""# 行动项

## 基本信息

| 字段 | 值 |
|---|---|
| 来源会议 | 待确认 |
| 提取时间 | {timestamp} |
| 行动项总数 | 0 |

## 行动项总览

| 状态 | 数量 |
|---|---|
| 待确认 | 0 |
| 待执行 | 0 |
| 已完成 | 0 |
| 已取消 | 0 |

## 行动项清单

| ID | 状态 | 优先级 | 行动项 | 负责人 | 协作人 | 截止时间 | 依赖条件 | 来源时间点 | 依据摘录 | 待确认原因 |
|---|---|---|---|---|---|---|---|---|---|---|
| — | 待确认 | — | — | 待确认 | — | 待确认 | — | — | — | LLM 暂不可用，请人工从转录文本提取 |

## 待确认行动项

全部行动项需人工复核。

## 生成说明

- 生成时间：{timestamp}
- 使用模型：template-fallback
- 待确认项目：全部
- 注意：行动项仅供参考，请人工复核后执行
"""


# ── Step 4: Speaker map & term corrections ───────────────────────────────

class SpeakerMapItem(BaseModel):
    speaker_id: str  # e.g. "Speaker 1"
    display_name: str  # e.g. "张三"


class SaveSpeakerMapRequest(BaseModel):
    folder_path: str
    speakers: list[SpeakerMapItem]


class SpeakerMapResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    speaker_map_path: str
    gbrain_ingest: bool = False


class TermCorrectionItem(BaseModel):
    original: str
    corrected: str
    type: str = "general"  # general / name / technical / acronym
    confidence: str = "中"


class SaveTermCorrectionsRequest(BaseModel):
    folder_path: str
    corrections: list[TermCorrectionItem]


class TermCorrectionsResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    corrections_path: str
    gbrain_ingest: bool = False


class DetectedSpeaker(BaseModel):
    speaker_id: str
    display_name: str
    ratio: str
    duration: str = "—"


class MeetingSpeakersResponse(BaseModel):
    ok: bool
    detected_speakers: list[DetectedSpeaker]


def _parse_speakers_from_transcript(transcript_text: str) -> list[DetectedSpeaker]:
    """Parse the 说话人概览 section of a transcript to extract speaker info."""
    speakers: list[DetectedSpeaker] = []
    in_section = False
    for line in transcript_text.split("\n"):
        if "说话人概览" in line and "##" in line:
            in_section = True
            continue
        if in_section:
            if "## " in line and "说话人概览" not in line:
                break
            if line.startswith("|") and "---" not in line and "说话人ID" not in line:
                # Parse table row: | Speaker 1 | Speaker 1 | 未映射 | 60% | — | 待确认 |
                parts = [p.strip() for p in line.split("|") if p.strip()]
                if len(parts) >= 4:
                    speakers.append(DetectedSpeaker(
                        speaker_id=parts[0],
                        display_name=parts[1] if len(parts) > 1 else parts[0],
                        ratio=parts[3] if len(parts) > 3 else "—",
                        duration=parts[4] if len(parts) > 4 else "—",
                    ))
    return speakers


def _speaker_timeline_rows(transcript_text: str, limit: int = 30) -> list[str]:
    rows: list[str] = []
    in_section = False
    for line in transcript_text.splitlines():
        if line.startswith("## ") and "说话人时间轴" in line:
            in_section = True
            continue
        if in_section and line.startswith("## "):
            break
        if not in_section or not line.startswith("|") or "---" in line or "行号" in line:
            continue
        parts = [part.strip() for part in line.split("|") if part.strip()]
        if len(parts) >= 4:
            rows.append(
                f"| {_escape_pipe(parts[0])} | {_escape_pipe(parts[1])} | {_escape_pipe(parts[2])} | {_escape_pipe(parts[3])} |"
            )
        if len(rows) >= limit:
            break
    return rows


@router.get("/{workspace_id}/meetings/speakers", response_model=MeetingSpeakersResponse)
def get_meeting_speakers(
    workspace_id: int,
    folder_path: str = Query(..., min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Parse detected speakers from the transcript."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持此操作")
    _validate_meeting_folder(workspace, folder_path)

    root = _workspace_file_root(workspace)
    folder_dir = _resolve_workspace_child(root, _safe_relative_path(folder_path))

    transcript_path = folder_dir / "02-转录文本" / "transcript-latest.md"
    if not transcript_path.exists():
        raise HTTPException(status_code=400, detail="转录文件不存在")

    text = transcript_path.read_text(encoding="utf-8")
    speakers = _parse_speakers_from_transcript(text)
    return MeetingSpeakersResponse(ok=True, detected_speakers=speakers)


def _build_speaker_map_markdown(
    speakers: list[SpeakerMapItem],
    author: str,
    timestamp: str,
    timeline_rows: list[str] | None = None,
) -> str:
    """Build speaker-map-latest.md formal content."""
    rows = "\n".join(
        f"| {_escape_pipe(s.speaker_id)} | {_escape_pipe(s.display_name)} | 已映射 | {_escape_pipe(author)} | {timestamp} |"
        for s in speakers
    )
    timeline = "\n".join(timeline_rows or ["| — | — | — | — |"])
    return (
        "# 说话人映射\n\n"
        "## 映射状态\n\n"
        f"- 修改人：{author}\n"
        f"- 修改时间：{timestamp}\n"
        f"- 映射状态：已确认\n\n"
        "## 说话人映射表\n\n"
        "| 说话人ID | 显示名称 | 映射状态 | 修改人 | 修改时间 |\n"
        "|---|---|---|---|---|\n"
        f"{rows}\n\n"
        "## 时间轴辅助信息\n\n"
        "| 行号 | 时间点 | 说话人ID | 内容摘要 |\n"
        "|---|---|---|---|\n"
        f"{timeline}\n"
    )


@router.post("/{workspace_id}/meetings/speaker-map", response_model=SpeakerMapResponse)
def save_meeting_speaker_map(
    workspace_id: int,
    req: SaveSpeakerMapRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Save speaker mapping as speaker-map-latest.md."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    _validate_meeting_folder(workspace, req.folder_path)

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(req.folder_path)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    transcript_dir = folder_dir / "02-转录文本"
    transcript_dir.mkdir(parents=True, exist_ok=True)

    now_ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    author_name = user.nickname or user.username

    # Version numbering
    ver = _next_version_number(transcript_dir, "speaker-map")

    transcript_text = _read_file_safe(folder_dir / "02-转录文本" / "transcript-latest.md")
    timeline_rows = _speaker_timeline_rows(transcript_text)
    md = _build_speaker_map_markdown(req.speakers, author_name, now_ts, timeline_rows)

    v_path = transcript_dir / f"speaker-map-v{ver}.md"
    v_path.write_text(md, encoding="utf-8")
    latest_path = transcript_dir / "speaker-map-latest.md"
    latest_path.write_text(md, encoding="utf-8")

    v_rel = v_path.relative_to(root).as_posix()
    latest_rel = latest_path.relative_to(root).as_posix()

    _upsert_workspace_file(db, workspace_id, user.id, v_rel,
                           f"speaker-map-v{ver}.md", "text/markdown", len(md.encode("utf-8")), v_path)
    _upsert_workspace_file(db, workspace_id, user.id, latest_rel,
                           "speaker-map-latest.md", "text/markdown", len(md.encode("utf-8")), latest_path)

    _write_workspace_audit(db, user.id, "meeting_speaker_map_save",
                           _audit_detail(workspace_id, req.folder_path, actor_id=user.id,
                                         gbrain_ingest=False))
    db.commit()
    return SpeakerMapResponse(ok=True, meeting_folder_path=req.folder_path,
                               speaker_map_path=latest_rel, gbrain_ingest=False)


def _build_term_corrections_markdown(corrections: list[TermCorrectionItem], timestamp: str) -> str:
    rows = "\n".join(
        f"| {_escape_pipe(c.original)} | {_escape_pipe(c.corrected)} | {_escape_pipe(c.type)} | {c.confidence} | 已确认 |"
        for c in corrections
    )
    return (
        "# 术语纠错\n\n"
        f"- 修改时间：{timestamp}\n"
        f"- 纠错数：{len(corrections)}\n\n"
        "## 术语纠错表\n\n"
        "| 原识别 | 建议修正 | 类型 | 置信度 | 状态 |\n"
        "|---|---|---|---|---|\n"
        f"{rows}\n"
    )


@router.post("/{workspace_id}/meetings/term-corrections", response_model=TermCorrectionsResponse)
def save_meeting_term_corrections(
    workspace_id: int,
    req: SaveTermCorrectionsRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Save term corrections as term-corrections-latest.md."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    _validate_meeting_folder(workspace, req.folder_path)

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(req.folder_path)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    transcript_dir = folder_dir / "02-转录文本"
    transcript_dir.mkdir(parents=True, exist_ok=True)

    now_ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    ver = _next_version_number(transcript_dir, "term-corrections")

    md = _build_term_corrections_markdown(req.corrections, now_ts)

    v_path = transcript_dir / f"term-corrections-v{ver}.md"
    v_path.write_text(md, encoding="utf-8")
    latest_path = transcript_dir / "term-corrections-latest.md"
    latest_path.write_text(md, encoding="utf-8")

    v_rel = v_path.relative_to(root).as_posix()
    latest_rel = latest_path.relative_to(root).as_posix()

    _upsert_workspace_file(db, workspace_id, user.id, v_rel,
                           f"term-corrections-v{ver}.md", "text/markdown", len(md.encode("utf-8")), v_path)
    _upsert_workspace_file(db, workspace_id, user.id, latest_rel,
                           "term-corrections-latest.md", "text/markdown", len(md.encode("utf-8")), latest_path)

    _write_workspace_audit(db, user.id, "meeting_term_corrections_save",
                           _audit_detail(workspace_id, req.folder_path, actor_id=user.id,
                                         gbrain_ingest=False))
    db.commit()
    return TermCorrectionsResponse(ok=True, meeting_folder_path=req.folder_path,
                                    corrections_path=latest_rel, gbrain_ingest=False)


def _escape_pipe(text: str) -> str:
    """Escape | characters in Markdown table cell content."""
    return text.replace("|", "&#124;")


def _validate_meeting_folder(workspace: "Workspace", folder_path: str) -> None:
    """Shared validation: check that folder_path is a valid meeting folder."""
    from app.features.workspaces.files.service import MEETING_SUBDIRS, meeting_parent_path
    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(folder_path)
    _ensure_not_trash_path(folder_rel)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    if not folder_dir.exists() or not folder_dir.is_dir():
        raise HTTPException(status_code=400, detail="会议文件夹不存在")
    missing = [sub for sub in MEETING_SUBDIRS if not (folder_dir / sub).is_dir()]
    if missing:
        raise HTTPException(status_code=400, detail="请选择会议文件夹")
    expected_parent = meeting_parent_path(workspace.workspace_kind)
    folder_posix = folder_rel.as_posix()
    if folder_posix != expected_parent and not folder_posix.startswith(expected_parent + "/"):
        raise HTTPException(status_code=400, detail=f"会议文件夹必须位于 {expected_parent}/ 下")


# ── Step 5: Media transcription ───────────────────────────────────────────

_SUPPORTED_MEDIA_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".mp4", ".mov", ".avi", ".wmv", ".mkv", ".webm"}


class MediaTranscribeResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    media_path: str
    transcript_v1_path: str
    transcript_latest_path: str
    transcription_status: str
    segment_count: int = 1
    warnings: list[str] = []
    gbrain_ingest: bool = False
    agent_run: AgentRunResponse | None = None
    token_cost: int = 0


class MediaTranscribePreflightRequest(BaseModel):
    folder_path: str
    filename: str
    size_bytes: int
    content_type: str = "application/octet-stream"


class MediaTranscribePreflightResponse(BaseModel):
    ok: bool
    filename: str
    size_mb: float
    estimated_duration_minutes: int | None = None
    is_long_media: bool = False
    estimated_segments: int = 1
    estimated_cost_note: str = ""
    warnings: list[str] = []
    model: str = "MiMo V2.5"


class MeetingRetryRequest(BaseModel):
    folder_path: str
    operation: str = "transcribe"  # transcribe / generate_minutes


class MeetingRetryResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    operation: str
    status: str  # queued / completed / partial / failed
    message: str = ""
    agent_run: AgentRunResponse | None = None


def _size_mb(path: Path) -> float:
    return path.stat().st_size / (1024 * 1024) if path.exists() else 0.0


def _duration_minutes(path: Path) -> int | None:
    """Estimate media duration in minutes via ffprobe. Returns None if unavailable."""
    try:
        proc = subprocess.run(
            ["ffprobe", "-v", "quiet", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", str(path)],
            capture_output=True, text=True, timeout=15,
        )
        return max(1, round(float(proc.stdout.strip()) / 60))
    except Exception:
        return None


def _estimate_media_info(size_bytes: int, filename: str) -> dict:
    """Estimate media cost/duration from file size and type. Returns preflight info."""
    size_mb = size_bytes / (1024 * 1024)
    is_audio_only = Path(filename).suffix.lower() in {".mp3", ".wav", ".m4a", ".ogg", ".flac"}
    # Rough estimate: audio ~1 MB/min, video ~8 MB/min
    if is_audio_only:
        est_minutes = max(1, round(size_mb / 1.0))
    else:
        est_minutes = max(1, round(size_mb / 8.0))
    is_long = est_minutes > 30
    seg_count = max(1, (est_minutes + 299) // 300)  # 300s segments
    warnings: list[str] = []
    if is_long:
        warnings.append(f"媒体时长超过 30 分钟（预估 {est_minutes} 分钟），将自动分段转录（{seg_count} 段）")
    if size_mb > 500:
        warnings.append("文件超过 500 MB，转录时间较长，请耐心等待")
    cost_note = f"预估 {est_minutes} 分钟，将使用 MiMo V2.5 模型转录。{'长视频将自动分段处理。' if is_long else ''}"
    return {
        "size_mb": round(size_mb, 1),
        "estimated_duration_minutes": est_minutes,
        "is_long_media": is_long,
        "estimated_segments": seg_count,
        "estimated_cost_note": cost_note,
        "warnings": warnings,
    }


@router.post("/{workspace_id}/meetings/transcribe/media", response_model=MediaTranscribeResponse)
async def transcribe_meeting_media(
    workspace_id: int,
    folder_path: str = Form(...),
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Upload and transcribe meeting audio/video via MiMo V2.5."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持音视频转录")
    _validate_meeting_folder(workspace, folder_path)

    filename = (file.filename or "recording").strip()
    lower = filename.lower()
    ext = Path(filename).suffix.lower()
    if ext not in _SUPPORTED_MEDIA_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"仅支持音视频格式：{', '.join(sorted(_SUPPORTED_MEDIA_EXTENSIONS))}")

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(folder_path)
    folder_dir = _resolve_workspace_child(root, folder_rel)
    lock_path = _acquire_meeting_run_lock(root, folder_dir, operation="media_transcribe", user_id=user.id)

    try:
        # Save raw media to 01-原始资料/
        raw_dir = folder_dir / "01-原始资料"
        raw_dir.mkdir(parents=True, exist_ok=True)
        media_path = _resolve_conflict_path(raw_dir, _safe_name(filename), "keep_both")
        if media_path is None:
            raise HTTPException(status_code=500, detail="无法保存媒体文件")

        content_bytes = await file.read()
        media_path.write_bytes(content_bytes)
        media_rel = media_path.relative_to(root).as_posix()

        # Warn for long videos > 30 min
        duration_min = _duration_minutes(media_path)
        if duration_min is not None and duration_min > 30:
            pass  # caller/frontend handles confirmation; this is informational

        # Transcribe via app.features.preprocessing.media_transcription
        from app.features.preprocessing.media_transcription import transcribe_media_to_markdown, load_media_transcription_options
        from app.shared.llm.client import get_llm_client

        token_cost = 0
        options = load_media_transcription_options()
        transcription_client = get_llm_client(options.model_profile)
        result = transcribe_media_to_markdown(media_path, options=options, llm_client=transcription_client)
        transcript_text = result.transcript_text
        transcription_status = result.transcription_status
        segment_count = result.segment_count
        warnings_list = list(result.warnings) if result.warnings else []
        if result.token_usage:
            token_cost += result.token_usage.get("input_tokens", 0) + result.token_usage.get("output_tokens", 0)
        if result.refinement_token_usage:
            token_cost += result.refinement_token_usage.get("input_tokens", 0) + result.refinement_token_usage.get("output_tokens", 0)
    except Exception as exc:
        transcription_status = "failed"
        segment_count = 0
        warnings_list = [str(exc)]
        token_cost = 0
        media_rel = media_path.relative_to(root).as_posix() if "media_path" in locals() and media_path.exists() else ""
        content_bytes = content_bytes if "content_bytes" in locals() else b""
        transcript_text = (
            f"# 会议转录文本 - 转录失败\n\n"
            f"**转录状态**：failed\n\n"
            f"**错误**：{exc}\n\n"
            f"请检查媒体文件是否有效，或联系管理员。\n\n"
            f"> **注意**：当前版本仅支持整体重试转录（通过「重试转录」按钮），不支持单片段重跑。\n"
        )
    finally:
        _release_meeting_run_lock(lock_path)

    # Save transcript
    transcript_dir = folder_dir / "02-转录文本"
    transcript_dir.mkdir(parents=True, exist_ok=True)

    # Use the formal template wrapper
    now_utc = datetime.now(timezone.utc)
    if transcription_status == "failed":
        final_md = transcript_text
    else:
        final_md = _build_transcript_markdown(
            transcript_text,
            now_utc,
            input_type=ext.lstrip("."),
            original_filename=filename,
            transcription_status=transcription_status,
            warnings=warnings_list,
        )

    v1_path = _resolve_conflict_path(transcript_dir, "transcript-v1.md", "keep_both")
    if v1_path is None:
        raise HTTPException(status_code=500, detail="无法写入转录文件")
    v1_path.write_text(final_md, encoding="utf-8")
    v1_rel = v1_path.relative_to(root).as_posix()

    latest_path = transcript_dir / "transcript-latest.md"
    latest_path.write_text(final_md, encoding="utf-8")
    latest_rel = latest_path.relative_to(root).as_posix()

    # Record WorkspaceFile metadata
    media_meta = _upsert_workspace_file(db, workspace_id, user.id, media_rel,
                           filename, file.content_type or "application/octet-stream", len(content_bytes), media_path)
    media_meta.rag_status = "pending_transcription" if transcription_status == "failed" else "not_ingested"
    v1_meta = _upsert_workspace_file(db, workspace_id, user.id, v1_rel,
                           "transcript-v1.md", "text/markdown", len(final_md.encode("utf-8")), v1_path)
    latest_meta = _upsert_workspace_file(db, workspace_id, user.id, latest_rel,
                           "transcript-latest.md", "text/markdown", len(final_md.encode("utf-8")), latest_path)
    transcript_rag_status = "failed" if transcription_status == "failed" else "partial" if transcription_status == "partial" else "not_ingested"
    v1_meta.rag_status = transcript_rag_status
    latest_meta.rag_status = transcript_rag_status

    _write_workspace_audit(db, user.id, "meeting_media_transcribe",
                           _audit_detail(workspace_id, folder_path, actor_id=user.id,
                                         media_file=media_rel,
                                         transcript=v1_rel,
                                         status=transcription_status,
                                         segments=segment_count,
                                         gbrain_ingest=False),
                           success=transcription_status != "failed")
    agent_run = _write_workspace_file_agent_run(
        db, user_id=user.id, workspace=workspace,
        source_type="meeting_media_transcribe",
        title="会议音视频转录",
        path=folder_path,
        detail=f"转录：{filename}（{segment_count}段，{transcription_status}）",
        status="failed" if transcription_status == "failed" else "completed",
    )
    _notify_meeting_run_finished(
        db,
        workspace=workspace,
        actor_user_id=user.id,
        folder_path=folder_path,
        title="会议音视频转录完成" if transcription_status != "failed" else "会议音视频转录失败",
        status="partial" if transcription_status == "partial" else "failed" if transcription_status == "failed" else "completed",
        detail=f"{filename}：{transcription_status}，生成 {latest_rel}",
    )
    db.commit()
    return MediaTranscribeResponse(
        ok=True,
        meeting_folder_path=folder_path,
        media_path=media_rel,
        transcript_v1_path=v1_rel,
        transcript_latest_path=latest_rel,
        transcription_status=transcription_status,
        segment_count=segment_count,
        warnings=warnings_list,
        gbrain_ingest=False,
        agent_run=serialize_agent_run(db, agent_run),
        token_cost=token_cost,
    )


# ── Step 5b: Media preflight ───────────────────────────────────────────────

@router.post("/{workspace_id}/meetings/transcribe/media/preflight", response_model=MediaTranscribePreflightResponse)
def preflight_meeting_media_transcribe(
    workspace_id: int,
    req: MediaTranscribePreflightRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Estimate media cost/duration before transcription. Does not upload or transcribe."""
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持音视频转录")
    _validate_meeting_folder(workspace, req.folder_path)

    ext = Path(req.filename).suffix.lower()
    if ext not in _SUPPORTED_MEDIA_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"仅支持音视频格式：{', '.join(sorted(_SUPPORTED_MEDIA_EXTENSIONS))}")

    info = _estimate_media_info(req.size_bytes, req.filename)
    return MediaTranscribePreflightResponse(
        ok=True,
        filename=req.filename,
        size_mb=info["size_mb"],
        estimated_duration_minutes=info["estimated_duration_minutes"],
        is_long_media=info["is_long_media"],
        estimated_segments=info["estimated_segments"],
        estimated_cost_note=info["estimated_cost_note"],
        warnings=info["warnings"],
        model="MiMo V2.5",
    )


# ── Step 5c: Meeting retry ─────────────────────────────────────────────────

@router.post("/{workspace_id}/meetings/retry", response_model=MeetingRetryResponse)
def retry_meeting_operation(
    workspace_id: int,
    req: MeetingRetryRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retry a failed meeting operation (transcription or minutes generation).

    Re-runs the same operation after verifying that:
    - The previous run failed
    - No active lock exists
    - The meeting folder still has the necessary input files
    """
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    if workspace.workspace_kind == "user":
        raise HTTPException(status_code=400, detail="个人工作台不支持此操作")
    if workspace.workspace_kind not in ("project", "customer"):
        raise HTTPException(status_code=400, detail="不支持的工作区类型")

    root = _workspace_file_root(workspace)
    folder_rel = _safe_relative_path(req.folder_path)
    _ensure_not_trash_path(folder_rel)
    folder_dir = _resolve_workspace_child(root, folder_rel)

    # Validate this is a meeting folder
    missing = [sub for sub in MEETING_SUBDIRS if not (folder_dir / sub).is_dir()]
    if missing:
        raise HTTPException(status_code=400, detail="请选择完整的会议文件夹")

    expected_parent = meeting_parent_path(workspace.workspace_kind)
    folder_posix = folder_rel.as_posix()
    if folder_posix != expected_parent and not folder_posix.startswith(expected_parent + "/"):
        raise HTTPException(status_code=400, detail=f"会议文件夹必须位于 {expected_parent}/ 下")

    lock_path = _meeting_run_lock_path(root, folder_dir)
    if lock_path.exists():
        raise HTTPException(status_code=409, detail="当前会议已有处理中任务，请等待完成后再操作")

    if req.operation == "transcribe":
        # Check for failed transcript
        transcript_latest = folder_dir / "02-转录文本" / "transcript-latest.md"
        if transcript_latest.exists():
            text = transcript_latest.read_text(encoding="utf-8")
            if not _failed_transcript_reason(text):
                raise HTTPException(status_code=400, detail="转录已成功完成，无需重试。如需重新转录请先删除旧转录文件。")

        # Find raw media in 01-原始资料 to re-transcribe
        raw_dir = folder_dir / "01-原始资料"
        if not raw_dir.exists() or not raw_dir.is_dir():
            raise HTTPException(status_code=400, detail="原始资料目录不存在，无法重试转录。请重新上传音视频文件。")

        media_files = [
            child for child in sorted(raw_dir.iterdir())
            if child.is_file() and child.suffix.lower() in _SUPPORTED_MEDIA_EXTENSIONS
        ]
        if not media_files:
            raise HTTPException(status_code=400, detail="原始资料中没有音视频文件，无法重试转录。请重新上传。")

        # Use the first (or most recent) media file
        media_path = max(media_files, key=lambda p: p.stat().st_mtime)
        _acquire_meeting_run_lock(root, folder_dir, operation="media_transcribe_retry", user_id=user.id)

        try:
            from app.features.preprocessing.media_transcription import transcribe_media_to_markdown, load_media_transcription_options
            from app.shared.llm.client import get_llm_client

            token_cost = 0
            options = load_media_transcription_options()
            transcription_client = get_llm_client(options.model_profile)
            result = transcribe_media_to_markdown(media_path, options=options, llm_client=transcription_client)
            transcript_text = result.transcript_text
            transcription_status = result.transcription_status
            segment_count = result.segment_count
            warnings_list = list(result.warnings) if result.warnings else []
            if result.token_usage:
                token_cost += result.token_usage.get("input_tokens", 0) + result.token_usage.get("output_tokens", 0)
            if result.refinement_token_usage:
                token_cost += result.refinement_token_usage.get("input_tokens", 0) + result.refinement_token_usage.get("output_tokens", 0)
        except Exception as exc:
            transcription_status = "failed"
            segment_count = 0
            warnings_list = [str(exc)]
            token_cost = 0
            transcript_text = (
                f"# 会议转录文本 - 转录失败\n\n"
                f"**转录状态**：failed\n\n"
                f"**错误**：{exc}\n\n"
                f"请检查媒体文件是否有效，或联系管理员。\n\n"
                f"> **注意**：当前版本仅支持整体重试转录（通过「重试转录」按钮），不支持单片段重跑。\n"
            )
        finally:
            _release_meeting_run_lock(lock_path)

        # Save transcript (overwrite latest, keep history)
        transcript_dir = folder_dir / "02-转录文本"
        transcript_dir.mkdir(parents=True, exist_ok=True)

        now_utc = datetime.now(timezone.utc)
        if transcription_status != "failed":
            final_md = _build_transcript_markdown(
                transcript_text, now_utc,
                input_type=media_path.suffix.lstrip("."),
                original_filename=media_path.name,
                transcription_status=transcription_status,
                warnings=warnings_list,
            )
        else:
            final_md = transcript_text

        v_ver = _next_version_number(transcript_dir, "transcript")
        v_path = transcript_dir / f"transcript-v{v_ver}.md"
        v_path.write_text(final_md, encoding="utf-8")
        latest_path = transcript_dir / "transcript-latest.md"
        latest_path.write_text(final_md, encoding="utf-8")

        v_rel = v_path.relative_to(root).as_posix()
        latest_rel = latest_path.relative_to(root).as_posix()

        _upsert_workspace_file(db, workspace_id, user.id, v_rel,
                               f"transcript-v{v_ver}.md", "text/markdown", len(final_md.encode("utf-8")), v_path)
        _upsert_workspace_file(db, workspace_id, user.id, latest_rel,
                               "transcript-latest.md", "text/markdown", len(final_md.encode("utf-8")), latest_path)

        _write_workspace_audit(db, user.id, "meeting_media_transcribe_retry",
                               _audit_detail(workspace_id, req.folder_path, actor_id=user.id,
                                             media_file=str(media_path), transcript=latest_rel,
                                             status=transcription_status, segments=segment_count))
        agent_run = _write_workspace_file_agent_run(
            db, user_id=user.id, workspace=workspace,
            source_type="meeting_media_transcribe_retry",
            title="重试会议音视频转录",
            path=req.folder_path,
            detail=f"转录重试：{media_path.name}（{segment_count}段，{transcription_status}）",
            status="failed" if transcription_status == "failed" else "completed",
        )
        _notify_meeting_run_finished(
            db, workspace=workspace, actor_user_id=user.id,
            folder_path=req.folder_path,
            title="会议音视频转录重试完成" if transcription_status != "failed" else "会议音视频转录重试失败",
            status="partial" if transcription_status == "partial" else "failed" if transcription_status == "failed" else "completed",
            detail=f"{media_path.name}：{transcription_status}，生成 {latest_rel}",
        )
        db.commit()
        return MeetingRetryResponse(
            ok=True,
            meeting_folder_path=req.folder_path,
            operation=req.operation,
            status=transcription_status,
            message=f"转录{'部分完成' if transcription_status == 'partial' else '完成' if transcription_status != 'failed' else '失败'}（{segment_count}段）",
            agent_run=serialize_agent_run(db, agent_run),
        )

    elif req.operation == "generate_minutes":
        # Check for failed or partial minutes
        minutes_latest = folder_dir / "04-会议纪要" / "minutes-latest.md"
        transcript_latest = folder_dir / "02-转录文本" / "transcript-latest.md"
        if not transcript_latest.exists():
            raise HTTPException(status_code=400, detail="转录文件不存在，无法重试生成纪要")

        transcript_text = transcript_latest.read_text(encoding="utf-8")
        failed_reason = _failed_transcript_reason(transcript_text)
        if failed_reason:
            raise HTTPException(status_code=400, detail=f"转录未成功，需要先重试转录。原因：{failed_reason}")

        # Delegate to generate endpoint with regenerate=True (in-process via function call)
        from app.shared.llm.client import get_llm_client

        speaker_map_text = _read_file_safe(folder_dir / "02-转录文本" / "speaker-map-latest.md")
        term_corrections_text = _read_file_safe(folder_dir / "02-转录文本" / "term-corrections-latest.md")
        original_filename = _transcript_metadata_value(transcript_text, "原始文件名")
        auxiliary_summaries_text = _read_auxiliary_summaries(folder_dir, source_filename=original_filename)
        transcription_status = _transcript_status_value(transcript_text)

        lock_path = _acquire_meeting_run_lock(root, folder_dir, operation="generate_minutes_retry", user_id=user.id)

        try:
            minutes_ver = _next_version_number(folder_dir / "04-会议纪要", "minutes")
            actions_ver = _next_version_number(folder_dir / "05-行动项", "actions")
            (folder_dir / "04-会议纪要").mkdir(parents=True, exist_ok=True)
            (folder_dir / "05-行动项").mkdir(parents=True, exist_ok=True)

            try:
                client = get_llm_client("deepseek-flash")
                model_used = "deepseek-flash"
                retry_meeting_type = _read_meeting_meta(folder_dir).get("meeting_type", "")
                minutes_response = client.complete(
                    [{"role": "user", "content": _build_minutes_prompt(transcript_text, speaker_map_text, term_corrections_text, auxiliary_summaries_text, meeting_type=retry_meeting_type)}],
                    system_prompt=_MEETING_SYSTEM_PROMPT,
                    temperature=0.3,
                )
                minutes_md = minutes_response.text.strip() if minutes_response.text else ""
                actions_response = client.complete(
                    [{"role": "user", "content": _build_actions_prompt(transcript_text, speaker_map_text, term_corrections_text, auxiliary_summaries_text)}],
                    system_prompt=_MEETING_SYSTEM_PROMPT,
                    temperature=0.3,
                )
                actions_md = actions_response.text.strip() if actions_response.text else ""
                token_cost = (
                    (minutes_response.usage or {}).get("input_tokens", 0) + (minutes_response.usage or {}).get("output_tokens", 0)
                    + (actions_response.usage or {}).get("input_tokens", 0) + (actions_response.usage or {}).get("output_tokens", 0)
                )
            except Exception as exc:
                model_used = "template-fallback"
                now_ts = serialize_datetime_utc(datetime.now(timezone.utc))
                minutes_md = _build_fallback_minutes(transcript_text, now_ts, str(exc))
                actions_md = _build_fallback_actions(now_ts)
                token_cost = 0
        finally:
            _release_meeting_run_lock(lock_path)

        if transcription_status == "partial":
            partial_block = "\n\n> 转录状态：partial。以下纪要和行动项仅基于已成功转录片段生成。缺失片段可能导致结论不完整，请人工复核。\n"
            if "转录状态：partial" not in minutes_md:
                minutes_md = minutes_md.rstrip() + partial_block

        minutes_v_path = folder_dir / "04-会议纪要" / f"minutes-v{minutes_ver}.md"
        minutes_v_path.write_text(minutes_md, encoding="utf-8")
        minutes_latest_path = folder_dir / "04-会议纪要" / "minutes-latest.md"
        minutes_latest_path.write_text(minutes_md, encoding="utf-8")
        actions_v_path = folder_dir / "05-行动项" / f"actions-v{actions_ver}.md"
        actions_v_path.write_text(actions_md, encoding="utf-8")
        actions_latest_path = folder_dir / "05-行动项" / "actions-latest.md"
        actions_latest_path.write_text(actions_md, encoding="utf-8")

        minutes_v_rel = minutes_v_path.relative_to(root).as_posix()
        minutes_latest_rel = minutes_latest_path.relative_to(root).as_posix()
        actions_v_rel = actions_v_path.relative_to(root).as_posix()
        actions_latest_rel = actions_latest_path.relative_to(root).as_posix()

        _upsert_workspace_file(db, workspace_id, user.id, minutes_v_rel,
                               f"minutes-v{minutes_ver}.md", "text/markdown", len(minutes_md.encode("utf-8")), minutes_v_path)
        _upsert_workspace_file(db, workspace_id, user.id, minutes_latest_rel,
                               "minutes-latest.md", "text/markdown", len(minutes_md.encode("utf-8")), minutes_latest_path)
        _upsert_workspace_file(db, workspace_id, user.id, actions_v_rel,
                               f"actions-v{actions_ver}.md", "text/markdown", len(actions_md.encode("utf-8")), actions_v_path)
        _upsert_workspace_file(db, workspace_id, user.id, actions_latest_rel,
                               "actions-latest.md", "text/markdown", len(actions_md.encode("utf-8")), actions_latest_path)

        _write_workspace_audit(db, user.id, "meeting_minutes_generate_retry",
                               _audit_detail(workspace_id, req.folder_path, actor_id=user.id,
                                             created_files=[minutes_v_rel, minutes_latest_rel, actions_v_rel, actions_latest_rel],
                                             model=model_used, token_cost=token_cost))
        agent_run = _write_workspace_file_agent_run(
            db, user_id=user.id, workspace=workspace,
            source_type="meeting_minutes_generate_retry",
            title="重试生成会议纪要与行动项",
            path=req.folder_path,
            detail=f"纪要重试：minutes-v{minutes_ver}.md / actions-v{actions_ver}.md（{model_used}）",
        )
        _notify_meeting_run_finished(
            db, workspace=workspace, actor_user_id=user.id,
            folder_path=req.folder_path,
            title="会议纪要重试生成完成",
            status="completed" if transcription_status != "partial" else "partial",
            detail=f"{req.folder_path} 已重新生成 minutes-v{minutes_ver}.md / actions-v{actions_ver}.md",
        )
        db.commit()
        return MeetingRetryResponse(
            ok=True,
            meeting_folder_path=req.folder_path,
            operation=req.operation,
            status="completed" if transcription_status != "partial" else "partial",
            message=f"重新生成纪要与行动项（v{minutes_ver}），模型：{model_used}，token：{token_cost}",
            agent_run=serialize_agent_run(db, agent_run),
        )

    else:
        raise HTTPException(status_code=400, detail=f"不支持的重试操作：{req.operation}")


# ── Step 6: GBrain meeting ingest ────────────────────────────────────────

class MeetingIngestRequest(BaseModel):
    folder_path: str
    recursive: bool = True
    single_file_path: str | None = None  # For single-file actions-only ingest


class MeetingIngestResponse(BaseModel):
    ok: bool
    meeting_folder_path: str
    gbrain_ready_path: str
    source_id: str
    source_scope: str
    ingested_files: list[str]
    skipped_files: list[str]
    gbrain_ingest: bool = True
    agent_run: AgentRunResponse | None = None
    warning: str = ""


def _gbrain_ready_compose(
    meeting_folder_name: str,
    minutes_md: str,
    transcript_md: str,
    actions_md: str = "",
    *,
    source_scope: str = "",
    source_context: str = "full_meeting",
) -> str:
    """Compose a GBrain-ready combined page from meeting output files.

    When source_context is 'action_items_only', the page contains only action items
    with lower context priority.
    """
    generated_at = serialize_datetime_utc(datetime.now(timezone.utc))
    is_actions_only = source_context == "action_items_only"

    lines = [
        "---",
        "schema: project_r_meeting_gbrain_ready_v1",
        f"title: {meeting_folder_name}",
        f"source_context: {source_context}",
        f"source_scope: {source_scope or 'workspace'}",
        "source_priority: transcript_first" if not is_actions_only else "source_priority: actions_only",
        "generated_by: Project_R meeting workflow",
        f"generated_at: {generated_at}",
        "---",
        "",
        f"# {meeting_folder_name}",
    ]

    if is_actions_only:
        lines.extend([
            "",
            "> ⚠️ 本页面仅包含行动项，不包含完整会议纪要和转录文本。",
            "> 如需要完整会议知识，建议录入完整会议资料（minutes-latest.md + transcript-latest.md）。",
            "",
            "---",
            "",
            "## Source Context / 来源说明",
            "",
            f"- source_context: `{source_context}` — 仅行动项，低上下文完整度",
            f"- source_scope: `{source_scope or 'workspace'}`",
            "- 行动项由 Project_R 从会议转录文本或辅助总结中提炼生成。",
            "- 没有转录文本和会议纪要上下文，行动项的负责人和截止时间均以原始文件标注为准。",
            "",
            "## 行动项内容",
            "",
        ])
        if actions_md:
            lines.append(actions_md.lstrip("# ").strip())
        else:
            lines.append("（无行动项内容）")
    else:
        lines.extend([
            "",
            "> 本页面由 Project_R 自动编译生成。来源：会议文件夹中的 latest 版本。会议纪要是整理结果，不是一手转录；事实判断优先回到一手转录证据。",
            "",
            "---",
            "",
            "## Source Context / 来源说明",
            "",
            f"- source_context: `{source_context}`",
            f"- source_scope: `{source_scope or 'workspace'}`",
            "- 一手证据：`transcript-latest.md`",
            "- 整理结果：`minutes-latest.md`",
            "- 行动项辅助：`actions-latest.md`",
            "- 原始音视频不直接作为 GBrain 正文；如需核验，应回到工作区原始资料。",
            "",
            "## 会议摘要",
            "",
            "摘要、决策、行动项、风险和待确认事项来自下方会议纪要与行动项结构化内容；所有无法在转录中确认的内容应以待确认处理。",
            "",
            "## 会议纪要",
            "",
            minutes_md.lstrip("# ").strip() if minutes_md else "（无纪要内容）",
            "",
            "## 决策 / 行动项 / 风险 / 待确认事项",
            "",
            "详见会议纪要和行动项章节中的结构化表格；引用时应优先使用表格中的来源时间点和依据摘录。",
            "",
            "## 转录文本",
            "",
            transcript_md.lstrip("# ").strip() if transcript_md else "（无转录）",
            "",
            "## 一手转录来源引用",
            "",
            "- 文件：`transcript-latest.md`",
            "- 引用粒度：时间戳 / 说话人 / 内容行",
        ])
        if actions_md:
            lines.extend(["", "---", "", "## 行动项（辅助参考）", "", actions_md.lstrip("# ").strip()])

    lines.append("")
    return "\n\n".join(lines)


@router.post("/{workspace_id}/meetings/ingest", response_model=MeetingIngestResponse)
def ingest_meeting_to_gbrain(
    workspace_id: int,
    req: MeetingIngestRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Compile meeting outputs into a GBrain-ready page and ingest to the workspace source.

    Supports both full meeting ingest (default) and single-file actions-only ingest
    (when single_file_path points to actions-latest.md).
    """
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)

    is_single_file_actions = (
        req.single_file_path is not None
        and req.single_file_path.rstrip("/").endswith("actions-latest.md")
    )

    if is_single_file_actions:
        # Single-file actions-only: skip full meeting folder validation
        pass
    else:
        _validate_meeting_folder(workspace, req.folder_path)

    member = db.query(WorkspaceMember).filter(
        WorkspaceMember.workspace_id == workspace_id, WorkspaceMember.user_id == user.id).first()
    is_admin = member and member.role == "admin"
    is_system_admin = user.role == "admin"

    # Permission check
    if workspace.workspace_kind == "customer":
        if not (is_admin or is_system_admin):
            raise HTTPException(status_code=403, detail="仅客户工作区管理员可录入会议资料")
    elif workspace.workspace_kind == "project":
        if not (is_admin or is_system_admin):
            raise HTTPException(status_code=403, detail="仅项目管理员可录入会议文件夹")

    root = _workspace_file_root(workspace)
    folder_dir = _resolve_workspace_child(root, _safe_relative_path(req.folder_path))

    # Determine source scope
    if workspace.workspace_kind == "project":
        from core.gbrain import project_source_id_for_workspace, project_source_paths_for_workspace
        source_id = project_source_id_for_workspace(workspace)
        paths = project_source_paths_for_workspace(workspace)
        source_scope = "project"
    else:
        from core.gbrain import customer_source_id_for_workspace, customer_source_paths_for_workspace
        source_id = customer_source_id_for_workspace(workspace)
        paths = customer_source_paths_for_workspace(workspace)
        source_scope = "customer"

    gbrain_ready_dir = paths.get("gbrain_ready", Path(""))
    if isinstance(gbrain_ready_dir, str):
        gbrain_ready_dir = Path(gbrain_ready_dir)
    gbrain_ready_dir.mkdir(parents=True, exist_ok=True)

    ingested: list[str] = []
    skipped: list[str] = []
    warning: str = ""

    if is_single_file_actions:
        # ── Single-file actions-only ingest ──────────────────────────────
        # Only read the actions-latest.md file
        actions_dir = folder_dir / "05-行动项"
        actions_latest = actions_dir / "actions-latest.md"
        if not actions_latest.exists():
            raise HTTPException(status_code=400, detail="行动项文件 actions-latest.md 不存在")

        actions_md = actions_latest.read_text(encoding="utf-8")
        lr = actions_latest.relative_to(root).as_posix()
        ingested.append(lr)

        # Check if full meeting files exist and warn
        has_minutes = (folder_dir / "04-会议纪要" / "minutes-latest.md").exists()
        has_transcript = (folder_dir / "02-转录文本" / "transcript-latest.md").exists()
        if has_minutes and has_transcript:
            warning = "该会议存在完整的纪要和转录文件。建议改为录入完整会议资料以获取更全面的知识上下文。"

        meeting_name = folder_dir.name
        gb_md = _gbrain_ready_compose(
            meeting_name,
            "",
            "",
            actions_md,
            source_scope=source_scope,
            source_context="action_items_only",
        )
        gb_path = _resolve_conflict_path(gbrain_ready_dir, f"{_safe_name(meeting_name)}.md", "keep_both")
        if gb_path is None:
            raise HTTPException(status_code=500, detail="无法写入 GBrain-ready 文件")
        gb_path.write_text(gb_md, encoding="utf-8")

        # Update ingested file status
        imeta = db.query(WorkspaceFile).filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == lr).first()
        if imeta:
            imeta.rag_status = "gbrain_ready"

    else:
        # ── Full meeting ingest (existing behavior) ─────────────────────
        sections = {
            "04-会议纪要": "minutes",
            "02-转录文本": "transcript",
            "05-行动项": "actions",
        }
        collected: dict[str, str] = {}
        for subdir, prefix in sections.items():
            dir_p = folder_dir / subdir
            if not dir_p.exists():
                continue
            latest_path = dir_p / f"{prefix}-latest.md"
            if not latest_path.exists():
                continue
            lr = latest_path.relative_to(root).as_posix()
            ingested.append(lr)
            collected[subdir] = latest_path.read_text(encoding="utf-8")

            # Mark vN files as superseded
            vn_pattern = re.compile(rf"^{re.escape(prefix)}-v(\d+)\.md$", re.IGNORECASE)
            for child in dir_p.iterdir():
                if not child.is_file():
                    continue
                if child.name == f"{prefix}-latest.md":
                    continue
                if vn_pattern.match(child.name):
                    sr = child.relative_to(root).as_posix()
                    skipped.append(sr)
                    sp_meta = db.query(WorkspaceFile).filter(
                        WorkspaceFile.workspace_id == workspace_id,
                        WorkspaceFile.relative_path == sr).first()
                    if sp_meta:
                        sp_meta.rag_status = "skipped_superseded_version"

        if not collected:
            raise HTTPException(status_code=400, detail="没有可录入的会议文件。请先生成纪要和转录。")

        meeting_name = folder_dir.name
        gb_md = _gbrain_ready_compose(
            meeting_name,
            collected.get("04-会议纪要", ""),
            collected.get("02-转录文本", ""),
            collected.get("05-行动项", ""),
            source_scope=source_scope,
            source_context="full_meeting",
        )
        gb_path = _resolve_conflict_path(gbrain_ready_dir, f"{_safe_name(meeting_name)}.md", "keep_both")
        if gb_path is None:
            raise HTTPException(status_code=500, detail="无法写入 GBrain-ready 文件")
        gb_path.write_text(gb_md, encoding="utf-8")

        # Update ingested files' statuses
        for ipath in ingested:
            imeta = db.query(WorkspaceFile).filter(
                WorkspaceFile.workspace_id == workspace_id,
                WorkspaceFile.relative_path == ipath).first()
            if imeta:
                imeta.rag_status = "gbrain_ready"

    _write_workspace_audit(db, user.id, "meeting_gbrain_ingest",
                           _audit_detail(workspace_id, req.folder_path, actor_id=user.id,
                                         source_id=source_id, source_scope=source_scope,
                                         ingested=ingested, skipped=skipped,
                                         gbrain_ready_path=str(gb_path.resolve()),
                                         gbrain_ready_generated=True,
                                         gbrain_synced=False,
                                         single_file_actions_only=is_single_file_actions,
                                         warning=warning if warning else None))
    agent_run = _write_workspace_file_agent_run(
        db, user_id=user.id, workspace=workspace,
        source_type="meeting_gbrain_ingest",
        title="会议行动项录入 GBrain" if is_single_file_actions else "会议资料录入 GBrain",
        path=req.folder_path,
        detail=f"已生成 {len(ingested)} 个 GBrain-ready 文件，跳过 {len(skipped)} 个旧版本",
    )
    db.commit()
    return MeetingIngestResponse(
        ok=True,
        meeting_folder_path=req.folder_path,
        gbrain_ready_path=str(gb_path.resolve()),
        source_id=source_id,
        source_scope=source_scope,
        ingested_files=ingested,
        skipped_files=skipped,
        gbrain_ingest=True,
        agent_run=serialize_agent_run(db, agent_run),
        warning=warning,
    )


# ── End meeting endpoints ────────────────────────────────────────────────


@router.delete("/{workspace_id}/files", response_model=WorkspaceFileMutationResponse)
def delete_workspace_file(
    workspace_id: int,
    path: str = Query(..., min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(path)
    target = _resolve_workspace_child(root, rel)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    rel_path = target.relative_to(root).as_posix()
    meta = (
        db.query(WorkspaceFile)
        .filter(
            WorkspaceFile.workspace_id == workspace_id,
            WorkspaceFile.relative_path == rel_path,
            WorkspaceFile.deleted_at.is_(None),
        )
        .first()
    )
    if not _member_can_mutate_file(member, user.id, meta, user.role):
        _raise_with_audit(
            db,
            user.id,
            "workspace_file_delete",
            403,
            "只有上传人或管理员可以删除该文件",
            _audit_detail(workspace_id, rel_path, meta.id if meta else None, actor_id=user.id, error="permission denied"),
        )
    if not meta:
        meta = _upsert_workspace_file(
            db,
            workspace_id,
            user.id,
            rel_path,
            target.name,
            "application/octet-stream",
            target.stat().st_size,
            target,
        )
    trash_path = _trash_target(root, meta, rel_path)
    shutil.move(str(target), str(trash_path))
    now = datetime.now(timezone.utc)
    meta.deleted_at = now
    meta.deleted_by = user.id
    meta.trash_path = trash_path.relative_to(root).as_posix()
    meta.rag_status = "source_deleted"
    meta.updated_at = now
    _mark_workspace_rag_pending(db, workspace_id)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_file_delete",
        _audit_detail(workspace_id, rel_path, meta.id, actor_id=user.id, trash_path=meta.trash_path),
    )
    if member.role == "admin" and meta.uploaded_by != user.id:
        notify_user(
            db,
            meta.uploaded_by,
            category="workspace",
            severity="warning",
            title="项目文件已被管理员删除",
            content=f"{workspace.name} 中的 {rel_path} 已由项目管理员删除并移入回收区。",
            action_status="none",
            action_kind="open_workspace",
            action_payload={"workspace_id": workspace.id},
            event_key=f"workspace:{workspace.id}:file_deleted:{meta.id}",
        )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_file_delete",
        title="移入项目回收区",
        path=rel_path,
        detail=meta.trash_path,
        result={"file_id": meta.id, "rag_status": meta.rag_status, "trash_path": meta.trash_path},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, file_id=meta.id, rag_status=meta.rag_status, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/files/restore", response_model=WorkspaceFileMutationResponse)
def restore_workspace_file(
    workspace_id: int,
    req: RestoreWorkspaceFileRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    meta = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.id == req.file_id)
        .first()
    )
    if not meta or not meta.deleted_at:
        raise HTTPException(status_code=404, detail="回收区文件不存在")
    source = _resolve_workspace_child(root, _safe_relative_path(meta.trash_path))
    target = _resolve_workspace_child(root, _safe_relative_path(meta.relative_path))
    if not source.exists() or not source.is_file():
        raise HTTPException(status_code=404, detail="回收区文件不存在")
    if target.exists():
        _raise_with_audit(
            db,
            user.id,
            "workspace_file_restore",
            409,
            "原路径已存在同名文件，请先处理冲突",
            _audit_detail(workspace_id, meta.relative_path, meta.id, actor_id=user.id, error="target exists"),
        )
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.move(str(source), str(target))
    now = datetime.now(timezone.utc)
    meta.deleted_at = None
    meta.deleted_by = None
    meta.trash_path = ""
    meta.rag_status = "new"
    _record_file_signature(meta, target)
    meta.updated_at = now
    _mark_workspace_rag_pending(db, workspace_id)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_file_restore",
        _audit_detail(workspace_id, meta.relative_path, meta.id, actor_id=user.id),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_file_restore",
        title="恢复项目文件",
        path=meta.relative_path,
        result={"file_id": meta.id, "rag_status": meta.rag_status},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=meta.relative_path, file_id=meta.id, rag_status=meta.rag_status, agent_run=serialize_agent_run(db, agent_run))


@router.delete("/{workspace_id}/files/permanent", response_model=WorkspaceFileMutationResponse)
def permanently_delete_workspace_file(
    workspace_id: int,
    file_id: int = Query(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    meta = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.id == file_id)
        .first()
    )
    if not meta or not meta.deleted_at:
        raise HTTPException(status_code=404, detail="回收区文件不存在")
    if not _member_can_restore_file(member, user.id, meta, user.role):
        _raise_with_audit(
            db,
            user.id,
            "workspace_file_permanent_delete",
            403,
            "只有上传人或管理员可以永久删除该文件",
            _audit_detail(workspace_id, meta.relative_path, meta.id, actor_id=user.id, error="permission denied"),
        )
    if meta.trash_path:
        trash_path = _resolve_workspace_child(root, _safe_relative_path(meta.trash_path))
        if trash_path.exists() and trash_path.is_file():
            trash_path.unlink()
    rel_path = meta.relative_path
    db.delete(meta)
    _mark_workspace_rag_pending(db, workspace_id)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_file_permanent_delete",
        _audit_detail(workspace_id, rel_path, file_id, actor_id=user.id),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_file_permanent_delete",
        title="永久删除项目文件",
        path=rel_path,
        result={"file_id": file_id, "rag_status": "pending"},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, file_id=file_id, rag_status="pending", agent_run=serialize_agent_run(db, agent_run))


@router.delete("/{workspace_id}/files/trash", response_model=WorkspaceTrashClearResponse)
def clear_workspace_trash(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    query = db.query(WorkspaceFile).filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.deleted_at.is_not(None))
    if member.role != "admin":
        query = query.filter(WorkspaceFile.uploaded_by == user.id, WorkspaceFile.deleted_by == user.id)
    deleted = 0
    for meta in query.all():
        if meta.trash_path:
            trash_path = _resolve_workspace_child(root, _safe_relative_path(meta.trash_path))
            if trash_path.exists() and trash_path.is_file():
                trash_path.unlink()
        db.delete(meta)
        deleted += 1
    _write_workspace_audit(db, user.id, "workspace_trash_clear", _audit_detail(workspace_id, actor_id=user.id, deleted_files=deleted))
    notify_workspace_bulk_delete_risk(db, workspace=workspace, actor_user_id=user.id, deleted_files=deleted)
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_trash_clear",
        title="清空项目回收区",
        path=".trash",
        result={"deleted_files": deleted},
    )
    db.commit()
    return {"ok": True, "deleted_files": deleted, "agent_run": serialize_agent_run(db, agent_run)}


@router.post("/{workspace_id}/attachments/save", response_model=WorkspaceFileMutationResponse)
def save_attachment_to_workspace(
    workspace_id: int,
    req: SaveAttachmentToWorkspaceRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    attachment = (
        db.query(SessionAttachment)
        .filter(
            SessionAttachment.id == req.attachment_id,
            SessionAttachment.session_id == req.session_id,
            SessionAttachment.user_id == user.id,
        )
        .first()
    )
    if not attachment:
        raise HTTPException(status_code=404, detail="附件不存在")
    source = Path(attachment.stored_path)
    if not source.exists() or not source.is_file():
        raise HTTPException(status_code=404, detail="附件文件不存在")
    root = _workspace_file_root(workspace)
    target_dir = _resolve_workspace_child(root, Path(DEFAULT_UNFILED_DIR))
    target_dir.mkdir(exist_ok=True)
    filename = _safe_name(attachment.original_name)
    conflict_path = _resolve_conflict_path(target_dir, filename, req.conflict_strategy)
    if conflict_path is None:
        skipped_path = f"{DEFAULT_UNFILED_DIR}/{filename}"
        agent_run = _write_workspace_file_agent_run(
            db,
            user_id=user.id,
            workspace=workspace,
            source_type="workspace_attachment_save",
            title="保存会话附件到项目",
            path=skipped_path,
            status="cancelled",
            detail="目标位置已存在同名文件，按策略跳过",
            result={"attachment_id": attachment.id, "rag_status": "skipped"},
        )
        db.commit()
        return WorkspaceFileMutationResponse(ok=False, path=skipped_path, rag_status="skipped", agent_run=serialize_agent_run(db, agent_run))
    target = _resolve_workspace_child(root, conflict_path.relative_to(root))
    shutil.copy2(source, target)
    rel_path = target.relative_to(root).as_posix()
    meta = _upsert_workspace_file(
        db,
        workspace_id,
        user.id,
        rel_path,
        target.name,
        attachment.content_type,
        target.stat().st_size,
        target,
    )
    _write_workspace_audit(db, user.id, "workspace_attachment_save", _audit_detail(workspace_id, rel_path, meta.id, actor_id=user.id, attachment_id=attachment.id))
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_attachment_save",
        title="保存会话附件到项目",
        path=rel_path,
        result={"file_id": meta.id, "attachment_id": attachment.id, "rag_status": meta.rag_status},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, file_id=meta.id, rag_status=meta.rag_status, agent_run=serialize_agent_run(db, agent_run))


@router.delete("/{workspace_id}/folders", response_model=WorkspaceFileMutationResponse)
def delete_workspace_folder(
    workspace_id: int,
    path: str = Query(..., min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(path)
    _ensure_not_trash_path(rel)
    if _is_template_root(rel):
        raise HTTPException(status_code=400, detail="默认模板文件夹不能删除")
    target = _resolve_workspace_child(root, rel)
    if not target.exists() or not target.is_dir():
        raise HTTPException(status_code=404, detail="文件夹不存在")
    if any(target.iterdir()):
        raise HTTPException(status_code=400, detail="只能删除空文件夹")
    target.rmdir()
    rel_path = target.relative_to(root).as_posix()
    _write_workspace_audit(
        db,
        user.id,
        "workspace_folder_delete",
        _audit_detail(workspace_id, rel_path, actor_id=user.id),
    )
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_folder_delete",
        title="删除项目文件夹",
        path=rel_path,
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=rel_path, agent_run=serialize_agent_run(db, agent_run))


@router.put("/{workspace_id}/paths/rename", response_model=WorkspaceFileMutationResponse)
def rename_workspace_path(
    workspace_id: int,
    req: RenameWorkspacePathRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(req.path)
    _ensure_not_trash_path(rel)
    if _is_template_root(rel):
        raise HTTPException(status_code=400, detail="默认模板文件夹不能重命名")
    source = _resolve_workspace_child(root, rel)
    if not source.exists():
        raise HTTPException(status_code=404, detail="文件或文件夹不存在")
    source_is_file = source.is_file()
    new_name = _safe_name(req.new_name)
    target = _resolve_workspace_child(root, rel.parent / new_name)
    if target.exists():
        raise HTTPException(status_code=409, detail="目标位置已存在同名项目")

    rel_path = source.relative_to(root).as_posix()
    meta = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.relative_path == rel_path, WorkspaceFile.deleted_at.is_(None))
        .first()
    )
    if source_is_file and not _member_can_mutate_file(member, user.id, meta, user.role):
        raise HTTPException(status_code=403, detail="只有上传人或管理员可以重命名该文件")

    shutil.move(str(source), str(target))
    new_rel_path = target.relative_to(root).as_posix()
    now = datetime.now(timezone.utc)
    if source_is_file:
        if meta:
            meta.relative_path = new_rel_path
            meta.original_name = target.name
            meta.rag_status = "pending"
            meta.updated_at = now
    else:
        _sync_file_descendant_paths(db, workspace_id, rel_path, new_rel_path)
    _write_workspace_audit(db, user.id, "workspace_path_rename", _audit_detail(workspace_id, new_rel_path, meta.id if meta else None, actor_id=user.id, old_path=rel_path))
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_path_rename",
        title="重命名项目路径",
        path=new_rel_path,
        detail=f"{rel_path} -> {new_rel_path}",
        result={"file_id": meta.id if meta else None, "old_path": rel_path, "rag_status": meta.rag_status if meta else None},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=new_rel_path, file_id=meta.id if meta else None, rag_status=meta.rag_status if meta else None, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/paths/move", response_model=WorkspaceFileMutationResponse)
def move_workspace_path(
    workspace_id: int,
    req: MoveWorkspacePathRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    member = _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(req.path)
    _ensure_not_trash_path(rel)
    if _is_template_root(rel):
        raise HTTPException(status_code=400, detail="默认模板文件夹不能移动")
    source = _resolve_workspace_child(root, rel)
    if not source.exists():
        raise HTTPException(status_code=404, detail="文件或文件夹不存在")
    source_is_file = source.is_file()
    target_dir_rel = _safe_relative_path(req.target_directory)
    _ensure_not_trash_path(target_dir_rel)
    target_dir = _resolve_workspace_child(root, target_dir_rel)
    if not target_dir.exists() or not target_dir.is_dir():
        raise HTTPException(status_code=400, detail="目标文件夹不存在")
    if not source_is_file and target_dir.resolve().is_relative_to(source.resolve()):
        raise HTTPException(status_code=400, detail="不能移动到自身下级目录")

    rel_path = source.relative_to(root).as_posix()
    meta = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.relative_path == rel_path, WorkspaceFile.deleted_at.is_(None))
        .first()
    )
    if source_is_file and not _member_can_mutate_file(member, user.id, meta, user.role):
        raise HTTPException(status_code=403, detail="只有上传人或管理员可以移动该文件")

    conflict_path = _resolve_conflict_path(target_dir, source.name, req.conflict_strategy)
    if conflict_path is None:
        agent_run = _write_workspace_file_agent_run(
            db,
            user_id=user.id,
            workspace=workspace,
            source_type="workspace_path_move",
            title="移动项目路径",
            path=rel_path,
            status="cancelled",
            detail="目标位置已存在同名路径，按策略跳过",
            result={"file_id": meta.id if meta else None, "rag_status": "skipped"},
        )
        db.commit()
        return WorkspaceFileMutationResponse(ok=False, path=rel_path, file_id=meta.id if meta else None, rag_status="skipped", agent_run=serialize_agent_run(db, agent_run))
    target = _resolve_workspace_child(root, conflict_path.relative_to(root))
    if target.exists() and target.is_dir() and source_is_file:
        raise HTTPException(status_code=400, detail="不能覆盖文件夹")
    if target.exists() and not source_is_file:
        raise HTTPException(status_code=409, detail="目标位置已存在同名文件夹")
    if target.exists() and req.conflict_strategy == "replace":
        existing_meta = (
            db.query(WorkspaceFile)
            .filter(
                WorkspaceFile.workspace_id == workspace_id,
                WorkspaceFile.relative_path == target.relative_to(root).as_posix(),
                WorkspaceFile.deleted_at.is_(None),
            )
            .first()
        )
        if existing_meta:
            db.delete(existing_meta)
        target.unlink()

    shutil.move(str(source), str(target))
    new_rel_path = target.relative_to(root).as_posix()
    now = datetime.now(timezone.utc)
    if source_is_file:
        if meta:
            meta.relative_path = new_rel_path
            meta.rag_status = "pending"
            meta.updated_at = now
    else:
        _sync_file_descendant_paths(db, workspace_id, rel_path, new_rel_path)
    _write_workspace_audit(db, user.id, "workspace_path_move", _audit_detail(workspace_id, new_rel_path, meta.id if meta else None, actor_id=user.id, old_path=rel_path))
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_path_move",
        title="移动项目路径",
        path=new_rel_path,
        detail=f"{rel_path} -> {new_rel_path}",
        result={"file_id": meta.id if meta else None, "old_path": rel_path, "rag_status": meta.rag_status if meta else None},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=new_rel_path, file_id=meta.id if meta else None, rag_status=meta.rag_status if meta else None, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/paths/copy", response_model=WorkspaceFileMutationResponse)
def copy_workspace_path(
    workspace_id: int,
    req: CopyWorkspacePathRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(req.path)
    _ensure_not_trash_path(rel)
    if _is_template_root(rel):
        raise HTTPException(status_code=400, detail="默认模板文件夹不能复制")
    source = _resolve_workspace_child(root, rel)
    if not source.exists():
        raise HTTPException(status_code=404, detail="文件或文件夹不存在")
    source_is_file = source.is_file()
    target_dir_rel = _safe_relative_path(req.target_directory)
    _ensure_not_trash_path(target_dir_rel)
    target_dir = _resolve_workspace_child(root, target_dir_rel)
    if not target_dir.exists() or not target_dir.is_dir():
        raise HTTPException(status_code=400, detail="目标文件夹不存在")
    if not source_is_file and target_dir.resolve().is_relative_to(source.resolve()):
        raise HTTPException(status_code=400, detail="不能复制到自身下级目录")

    rel_path = source.relative_to(root).as_posix()
    source_meta = (
        db.query(WorkspaceFile)
        .filter(WorkspaceFile.workspace_id == workspace_id, WorkspaceFile.relative_path == rel_path, WorkspaceFile.deleted_at.is_(None))
        .first()
    )
    conflict_path = _resolve_conflict_path(target_dir, source.name, req.conflict_strategy)
    if conflict_path is None:
        agent_run = _write_workspace_file_agent_run(
            db,
            user_id=user.id,
            workspace=workspace,
            source_type="workspace_path_copy",
            title="复制项目路径",
            path=rel_path,
            status="cancelled",
            detail="目标位置已存在同名路径，按策略跳过",
            result={"file_id": source_meta.id if source_meta else None, "rag_status": "skipped"},
        )
        db.commit()
        return WorkspaceFileMutationResponse(ok=False, path=rel_path, file_id=source_meta.id if source_meta else None, rag_status="skipped", agent_run=serialize_agent_run(db, agent_run))
    target = _resolve_workspace_child(root, conflict_path.relative_to(root))
    if target.exists() and target.is_dir() and source_is_file:
        raise HTTPException(status_code=400, detail="不能覆盖文件夹")
    if target.exists() and not source_is_file:
        raise HTTPException(status_code=409, detail="目标位置已存在同名文件夹")
    if target.exists() and req.conflict_strategy == "replace":
        existing_meta = (
            db.query(WorkspaceFile)
            .filter(
                WorkspaceFile.workspace_id == workspace_id,
                WorkspaceFile.relative_path == target.relative_to(root).as_posix(),
                WorkspaceFile.deleted_at.is_(None),
            )
            .first()
        )
        if existing_meta:
            db.delete(existing_meta)
        target.unlink()

    if source_is_file:
        shutil.copy2(source, target)
        meta = _create_copied_file_metadata(
            db,
            workspace_id=workspace_id,
            source_rel_path=rel_path,
            target_file=target,
            root=root,
            user_id=user.id,
        )
    else:
        shutil.copytree(source, target)
        _copy_descendant_file_metadata(
            db,
            workspace_id=workspace_id,
            source_dir=source,
            target_dir=target,
            root=root,
            user_id=user.id,
        )
        meta = None

    new_rel_path = target.relative_to(root).as_posix()
    _write_workspace_audit(db, user.id, "workspace_path_copy", _audit_detail(workspace_id, new_rel_path, meta.id if meta else None, actor_id=user.id, old_path=rel_path))
    agent_run = _write_workspace_file_agent_run(
        db,
        user_id=user.id,
        workspace=workspace,
        source_type="workspace_path_copy",
        title="复制项目路径",
        path=new_rel_path,
        detail=f"{rel_path} -> {new_rel_path}",
        result={"file_id": meta.id if meta else None, "old_path": rel_path, "rag_status": meta.rag_status if meta else None},
    )
    db.commit()
    return WorkspaceFileMutationResponse(ok=True, path=new_rel_path, file_id=meta.id if meta else None, rag_status=meta.rag_status if meta else None, agent_run=serialize_agent_run(db, agent_run))


@router.post("/{workspace_id}/knowledge/ingest", response_model=WorkspaceKnowledgeRefreshResponse)
@router.post("/{workspace_id}/knowledge/refresh", response_model=WorkspaceKnowledgeRefreshResponse)
def refresh_workspace_knowledge(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    req: WorkspaceKnowledgeIngestRequest | None = None,
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    if workspace.workspace_kind not in {"project", "customer"}:
        raise HTTPException(status_code=400, detail="当前工作区类型暂不支持一键录入知识库")
    ingest_request = _normalize_workspace_ingest_request(db, workspace, user, req)
    payload = _execute_workspace_knowledge_ingest(
        db,
        workspace,
        user.id,
        source_path=ingest_request["path"],
        recursive=ingest_request["recursive"],
    )
    agent_run = _write_immediate_workspace_ingest_agent_run(db, workspace, user.id, payload)
    payload["agent_run"] = serialize_agent_run(db, agent_run)
    db.commit()
    return WorkspaceKnowledgeRefreshResponse(**payload)


@router.post("/{workspace_id}/knowledge/ingest/async", response_model=WorkspaceKnowledgeIngestJobResponse)
def enqueue_workspace_knowledge_ingest(
    workspace_id: int,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    req: WorkspaceKnowledgeIngestRequest | None = None,
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    _ensure_member(db, user.id, workspace_id)
    if workspace.workspace_kind not in {"project", "customer"}:
        raise HTTPException(status_code=400, detail="当前工作区类型暂不支持一键录入知识库")
    ingest_request = _normalize_workspace_ingest_request(db, workspace, user, req)
    job = WorkspaceIngestJob(
        workspace_id=workspace_id,
        requested_by=user.id,
        status="queued",
        result_json=json.dumps({"request": ingest_request}, ensure_ascii=False),
    )
    db.add(job)
    db.flush()
    run_id = _workspace_ingest_job_run_id(job)
    mark_workspace_ingest_job_queued(job, workspace=workspace, ingest_request=ingest_request, run_id=run_id)
    create_queued_workspace_ingest_agent_run(db, job, workspace, ingest_request)
    notify_workspace_ingest_queued(
        db,
        workspace=workspace,
        actor_user_id=user.id,
        job_id=job.id,
        request_label=_workspace_ingest_request_label(ingest_request),
    )
    db.commit()
    db.refresh(job)
    background_tasks.add_task(_run_workspace_knowledge_ingest_job, job.id)
    return _serialize_ingest_job(db, job)


@router.get("/{workspace_id}/knowledge/ingest/jobs/{job_id}", response_model=WorkspaceKnowledgeIngestJobResponse)
def get_workspace_knowledge_ingest_job(
    workspace_id: int,
    job_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _ensure_member(db, user.id, workspace_id)
    job = (
        db.query(WorkspaceIngestJob)
        .filter(WorkspaceIngestJob.id == job_id, WorkspaceIngestJob.workspace_id == workspace_id)
        .first()
    )
    if not job:
        raise HTTPException(status_code=404, detail="录入任务不存在")
    return _serialize_ingest_job(db, job)


@router.get("/{workspace_id}/knowledge/graph", response_model=WorkspaceKnowledgeGraphResponse)
def workspace_knowledge_graph(
    workspace_id: int,
    focus: str | None = Query(default=None),
    entity_type: str | None = Query(default=None),
    limit: int = Query(default=120, ge=1, le=500),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id, Workspace.is_archived == False).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_can_open_workspace(db, user, workspace)
    scope = _workspace_gbrain_graph_scope(workspace)
    source_id = str(scope["source_id"])
    derived_path = scope["derived_path"]
    if not isinstance(derived_path, Path):
        raise HTTPException(status_code=500, detail="工作区 GBrain 路径配置错误")
    graph = build_source_graph(
        source_id,
        derived_path=derived_path,
        focus=focus,
        entity_type=entity_type,
        limit=limit,
    )
    _write_workspace_audit(
        db,
        user.id,
        "workspace_gbrain_graph_view",
        _audit_detail(
            workspace.id,
            "knowledge/graph",
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            source_id=source_id,
            nodes=len(graph.get("nodes") or []),
            edges=len(graph.get("edges") or []),
            events=len(graph.get("events") or []),
        ),
    )
    db.commit()
    return WorkspaceKnowledgeGraphResponse(
        ok=bool(graph.get("ok")),
        workspace_id=workspace.id,
        workspace_name=workspace.name,
        workspace_kind=workspace.workspace_kind,
        source_id=source_id,
        source_scope=str(scope["source_scope"]),
        intelligence_kind=str(scope["intelligence_kind"]),
        derived_path=str(graph.get("derived_path") or derived_path),
        focus=focus,
        entity_type=entity_type,
        nodes=list(graph.get("nodes") or []),
        edges=list(graph.get("edges") or []),
        events=list(graph.get("events") or []),
        profile_cards=_workspace_profile_cards(graph),
        stats=graph.get("stats") if isinstance(graph.get("stats"), dict) else None,
        warnings=[str(item) for item in graph.get("warnings") or []],
    )


@router.get("/{workspace_id}/knowledge/entity-merge-candidates", response_model=WorkspaceEntityMergeCandidatesResponse)
def workspace_entity_merge_candidates(
    workspace_id: int,
    focus: str | None = Query(default=None),
    limit: int = Query(default=80, ge=1, le=200),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id, Workspace.is_archived == False).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    if not _is_workspace_admin(db, user, workspace_id):
        raise HTTPException(status_code=403, detail="仅系统管理员或工作区管理员可查看实体候选")
    scope = _workspace_gbrain_graph_scope(workspace)
    source_id = str(scope["source_id"])
    derived_path = scope["derived_path"]
    if not isinstance(derived_path, Path):
        raise HTTPException(status_code=500, detail="工作区 GBrain 路径配置错误")
    result = build_entity_merge_candidates(source_id, derived_path=derived_path, focus=focus, limit=limit)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_gbrain_entity_merge_candidates_view",
        _audit_detail(
            workspace.id,
            "knowledge/entity-merge-candidates",
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            source_id=source_id,
            candidates=len(result.get("candidates") or []),
        ),
    )
    db.commit()
    return WorkspaceEntityMergeCandidatesResponse(
        ok=bool(result.get("ok")),
        workspace_id=workspace.id,
        workspace_name=workspace.name,
        workspace_kind=workspace.workspace_kind,
        source_id=source_id,
        source_scope=str(scope["source_scope"]),
        derived_path=str(result.get("derived_path") or derived_path),
        focus=focus,
        candidates=list(result.get("candidates") or []),
        stats=result.get("stats") if isinstance(result.get("stats"), dict) else None,
        warnings=[str(item) for item in result.get("warnings") or []],
    )


@router.post("/{workspace_id}/knowledge/entity-merge-candidates/action")
def workspace_entity_merge_candidate_action(
    workspace_id: int,
    request: WorkspaceEntityMergeActionRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id, Workspace.is_archived == False).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    if not _is_workspace_admin(db, user, workspace_id):
        raise HTTPException(status_code=403, detail="仅系统管理员或工作区管理员可处理实体候选")
    if request.action not in {"create_entity_page", "dismiss", "record_alias", "apply_relink_changes"}:
        raise HTTPException(status_code=400, detail="实体候选操作不合法")
    scope = _workspace_gbrain_graph_scope(workspace)
    source_id = str(scope["source_id"])
    derived_path = scope["derived_path"]
    if not isinstance(derived_path, Path):
        raise HTTPException(status_code=500, detail="工作区 GBrain 路径配置错误")
    result = apply_entity_merge_candidate_action(
        source_id,
        request.candidate_id,
        request.action,
        derived_path=derived_path,
        actor=user.username,
    )
    sync_result: dict | None = None
    if result.get("ok") and result.get("status") in {
        "created",
        "already_exists",
        "alias_recorded",
        "alias_already_exists",
        "relink_applied",
    }:
        sync_result = GBrainAdapter().sync_source(source_id=source_id, repo_path=derived_path, no_pull=True)
        result["sync"] = sync_result
    _write_workspace_audit(
        db,
        user.id,
        "workspace_gbrain_entity_merge_candidate_action",
        _audit_detail(
            workspace.id,
            "knowledge/entity-merge-candidates/action",
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            source_id=source_id,
            action=request.action,
            status=result.get("status"),
            candidate_id=request.candidate_id[:160],
            sync=(sync_result or {}).get("status") or "",
        ),
    )
    db.commit()
    if not result.get("ok"):
        raise HTTPException(status_code=400, detail=result.get("error") or "实体候选操作失败")
    return result


@router.get("/{workspace_id}/knowledge/entity-merge-candidates/preview")
def workspace_entity_merge_candidate_preview(
    workspace_id: int,
    candidate_id: str = Query(..., min_length=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id, Workspace.is_archived == False).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    if not _is_workspace_admin(db, user, workspace_id):
        raise HTTPException(status_code=403, detail="仅系统管理员或工作区管理员可预览实体候选")
    scope = _workspace_gbrain_graph_scope(workspace)
    source_id = str(scope["source_id"])
    derived_path = scope["derived_path"]
    if not isinstance(derived_path, Path):
        raise HTTPException(status_code=500, detail="工作区 GBrain 路径配置错误")
    result = build_entity_merge_candidate_preview(source_id, candidate_id, derived_path=derived_path)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_gbrain_entity_merge_candidate_preview",
        _audit_detail(
            workspace.id,
            "knowledge/entity-merge-candidates/preview",
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            source_id=source_id,
            status=result.get("status"),
            candidate_id=candidate_id[:160],
            changes=((result.get("stats") or {}).get("planned_relink_changes") or 0),
        ),
    )
    db.commit()
    if not result.get("ok"):
        raise HTTPException(status_code=400, detail=result.get("error") or "实体候选预览失败")
    return result


@router.get("/{workspace_id}/knowledge/graph/native-context")
def workspace_native_graph_context(
    workspace_id: int,
    slug: str = Query(..., min_length=1),
    depth: int = Query(default=2, ge=1, le=10),
    direction: str = Query(default="both"),
    link_type: str | None = Query(default=None),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id, Workspace.is_archived == False).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_can_open_workspace(db, user, workspace)
    scope = _workspace_gbrain_graph_scope(workspace)
    source_id = str(scope["source_id"])
    result = GBrainAdapter().graph_context(
        slug,
        source_id=source_id,
        depth=depth,
        direction=direction,
        link_type=link_type,
    )
    _write_workspace_audit(
        db,
        user.id,
        "workspace_gbrain_native_graph_context",
        _audit_detail(
            workspace.id,
            "knowledge/graph/native-context",
            actor_id=user.id,
            workspace_kind=workspace.workspace_kind,
            source_id=source_id,
            slug=slug[:160],
            status=result.get("status"),
        ),
    )
    db.commit()
    return result


def _normalize_workspace_ingest_request(
    db: Session,
    workspace: Workspace,
    user: User,
    req: WorkspaceKnowledgeIngestRequest | None,
) -> dict:
    requested_path = (req.path if req else "").replace("\\", "/").strip("/")
    recursive = True if req is None else bool(req.recursive)
    root = _workspace_file_root(workspace)
    rel = _safe_relative_path(requested_path) if requested_path else Path()
    _ensure_not_trash_path(rel)
    target = _resolve_workspace_child(root, rel)
    if not target.exists():
        raise HTTPException(status_code=404, detail="录入路径不存在")
    rel_path = target.relative_to(root).as_posix()
    if rel_path == ".":
        rel_path = ""
    is_file = target.is_file()
    is_directory = target.is_dir()
    if not is_file and not is_directory:
        raise HTTPException(status_code=400, detail="录入路径不是文件或文件夹")

    is_admin = _is_workspace_admin(db, user, workspace.id)
    if workspace.workspace_kind == "customer" and not is_admin:
        raise HTTPException(status_code=403, detail="客户资料录入仅允许系统管理员或客户工作区管理员执行")
    if workspace.workspace_kind == "project":
        if is_admin:
            pass
        elif is_file and not recursive:
            meta = (
                db.query(WorkspaceFile)
                .filter(
                    WorkspaceFile.workspace_id == workspace.id,
                    WorkspaceFile.relative_path == rel_path,
                    WorkspaceFile.deleted_at.is_(None),
                )
                .first()
            )
            if not meta or meta.uploaded_by != user.id:
                raise HTTPException(status_code=403, detail="普通成员只能录入自己上传的单个文件")
        else:
            raise HTTPException(status_code=403, detail="递归录入文件夹仅允许系统管理员或项目管理员执行")
    if not is_directory:
        recursive = False
    return {
        "path": rel_path,
        "recursive": recursive,
        "target_type": "directory" if is_directory else "file",
    }


def _compile_project_workspace_sources_for_request(workspace: Workspace, source_path: str, recursive: bool) -> dict:
    try:
        return compile_project_workspace_sources(workspace, source_path=source_path, recursive=recursive)
    except TypeError:
        if source_path or not recursive:
            raise
        return compile_project_workspace_sources(workspace)


def _compile_customer_workspace_sources_for_request(workspace: Workspace, source_path: str, recursive: bool) -> dict:
    try:
        return compile_customer_workspace_sources(workspace, source_path=source_path, recursive=recursive)
    except TypeError:
        if source_path or not recursive:
            raise
        return compile_customer_workspace_sources(workspace)


def _new_workspace_ingest_run_id(workspace: Workspace) -> str:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    return f"workspace-{workspace.id}-{stamp}-{uuid.uuid4().hex[:8]}"


def _execute_workspace_knowledge_ingest(
    db: Session,
    workspace: Workspace,
    actor_user_id: int,
    *,
    source_path: str = "",
    recursive: bool = True,
    run_id: str | None = None,
    initial_status_history: list[dict] | None = None,
) -> dict:
    run_id = run_id or _new_workspace_ingest_run_id(workspace)
    started_at = datetime.now(timezone.utc)
    status_history = list(initial_status_history or [])
    if not any(item.get("status") == "preprocessing" for item in status_history if isinstance(item, dict)):
        status_history.append(_workspace_ingest_status_event("preprocessing", "开始预处理源文件", started_at))
    payload = execute_workspace_ingest_core(
        db,
        workspace,
        actor_user_id,
        source_path=source_path,
        recursive=recursive,
        run_id=run_id,
        started_at=started_at,
        status_history=status_history,
        compile_project=_compile_project_workspace_sources_for_request,
        compile_customer=_compile_customer_workspace_sources_for_request,
        adapter_factory=GBrainAdapter,
    )
    _write_workspace_audit(
        db,
        actor_user_id,
        "workspace_knowledge_refresh",
        _audit_detail(
            workspace.id,
            actor_id=actor_user_id,
            **workspace_ingest_audit_fields(payload, source_path=source_path, recursive=recursive),
        ),
    )
    if workspace.workspace_kind in {"project", "customer"}:
        _notify_workspace_ingest_finished(
            db,
            workspace=workspace,
            actor_user_id=actor_user_id,
            ok=bool(payload.get("ok")),
            indexed_files=int(payload.get("indexed_files", 0)),
            failed_files=int(payload.get("failed_files", 0)),
            pending_extractor_capability_files=int(payload.get("pending_extractor_capability_files", 0)),
            pending_transcription_files=int(payload.get("pending_transcription_files", 0)),
            gbrain_error=payload.get("gbrain_error"),
        )
    return payload


def _run_workspace_knowledge_ingest_job(job_id: int) -> None:
    db = SessionLocal()
    try:
        job = db.query(WorkspaceIngestJob).filter(WorkspaceIngestJob.id == job_id).first()
        if not job or job.status not in {"queued", "failed"}:
            return
        workspace = db.query(Workspace).filter(Workspace.id == job.workspace_id).first()
        ingest_request = _workspace_ingest_request_from_job(job)
        run_id = _workspace_ingest_run_id_from_job(job)
        initial_history = mark_workspace_ingest_job_running(
            job,
            workspace=workspace,
            ingest_request=ingest_request,
            run_id=run_id,
        )
        agent_run = _get_or_create_workspace_ingest_agent_run(db, job, workspace)
        agent_run.status = "running"
        add_workspace_ingest_started_event(db, agent_run, workspace, ingest_request)
        db.commit()

        if not workspace:
            raise ValueError("workspace no longer exists")
        payload = _execute_workspace_knowledge_ingest(
            db,
            workspace,
            job.requested_by,
            source_path=str(ingest_request.get("path") or ""),
            recursive=bool(ingest_request.get("recursive", True)),
            run_id=run_id,
            initial_status_history=initial_history,
        )
        mark_workspace_ingest_job_completed(job, payload)
        add_workspace_ingest_result_event(db, agent_run, payload)
        finish_workspace_ingest_agent_run(db, agent_run, payload)
        db.commit()
    except Exception as exc:
        db.rollback()
        job = db.query(WorkspaceIngestJob).filter(WorkspaceIngestJob.id == job_id).first()
        if job:
            workspace = db.query(Workspace).filter(Workspace.id == job.workspace_id).first()
            request = _workspace_ingest_request_from_job(job)
            run_id = _workspace_ingest_run_id_from_job(job)
            mark_workspace_ingest_job_failed(
                job,
                workspace=workspace,
                request=request,
                run_id=run_id,
                error=str(exc),
            )
            agent_run = _get_or_create_workspace_ingest_agent_run(db, job, workspace)
            fail_workspace_ingest_agent_run(db, agent_run, workspace_id=job.workspace_id, error=str(exc))
            notify_workspace_ingest_failed(
                db,
                job=job,
                workspace=workspace,
                error=str(exc),
            )
            db.commit()
    finally:
        db.close()


def _serialize_ingest_job(db: Session, job: WorkspaceIngestJob) -> WorkspaceKnowledgeIngestJobResponse:
    try:
        result = json.loads(job.result_json or "{}")
    except json.JSONDecodeError:
        result = {}
    return WorkspaceKnowledgeIngestJobResponse(
        id=job.id,
        workspace_id=job.workspace_id,
        requested_by=job.requested_by,
        status=job.status,
        result=result if isinstance(result, dict) else {},
        error_message=job.error_message,
        created_at=job.created_at,
        started_at=job.started_at,
        finished_at=job.finished_at,
        agent_run=_serialize_workspace_ingest_agent_run(db, job),
    )


@router.put("/{workspace_id}", response_model=WorkspaceResponse)
def update_workspace(
    workspace_id: int,
    req: UpdateWorkspaceRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")

    _ensure_workspace_admin(db, user, workspace_id)
    if workspace.is_default and (req.name is not None or req.description is not None):
        raise HTTPException(status_code=400, detail="默认工作区不能重命名或修改")

    if req.name is not None:
        name = req.name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="项目名称不能为空")
        if len(name) > 128:
            raise HTTPException(status_code=400, detail="项目名称不能超过 128 个字符")
        workspace.name = name
    if req.description is not None:
        workspace.description = req.description.strip()
    if req.is_hidden is not None:
        if workspace.workspace_kind == "user":
            raise HTTPException(status_code=400, detail="个人工作台不支持隐藏项目设置")
        workspace.is_hidden = req.is_hidden

    db.commit()
    db.refresh(workspace)
    return _workspace_response(db, workspace, user=user)


@router.get("/{workspace_id}/member-candidates", response_model=list[WorkspaceMemberCandidateResponse])
def list_workspace_member_candidates(
    workspace_id: int,
    q: str = Query(default=""),
    limit: int = Query(default=30, ge=1, le=80),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    text = q.strip().lower()
    member_roles = {
        member.user_id: member.role
        for member in db.query(WorkspaceMember).filter(WorkspaceMember.workspace_id == workspace_id).all()
    }
    users_query = db.query(User).filter(User.is_active == True)
    if text:
        needle = f"%{text}%"
        users_query = users_query.filter(
            or_(
                User.username.ilike(needle),
                User.nickname.ilike(needle),
                User.work_group.ilike(needle),
            )
        )
    users = users_query.order_by(User.username.asc()).limit(limit).all()
    return [
        WorkspaceMemberCandidateResponse(
            user_id=item.id,
            username=item.username,
            nickname=item.nickname,
            work_group=item.work_group,
            role=item.role,
            is_member=item.id in member_roles,
            member_role=member_roles.get(item.id),
        )
        for item in users
    ]


@router.get("/{workspace_id}/group-candidates", response_model=list[WorkspaceGroupCandidateResponse])
def list_workspace_group_candidates(
    workspace_id: int,
    q: str = Query(default=""),
    limit: int = Query(default=30, ge=1, le=80),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    authorized = set(_workspace_access_groups(db, workspace_id))
    groups: dict[str, WorkspaceGroupCandidateResponse] = {}
    for (group_name,) in db.query(User.work_group).filter(User.work_group != "").distinct().all():
        normalized = _normalize_group_name(group_name)
        if normalized:
            groups[normalized] = WorkspaceGroupCandidateResponse(
                group_name=normalized,
                source="user",
                is_authorized=normalized in authorized,
            )
    for group_name in authorized:
        groups[group_name] = WorkspaceGroupCandidateResponse(
            group_name=group_name,
            source=groups.get(group_name).source if group_name in groups else "workspace",
            is_authorized=True,
        )
    text = q.strip().lower()
    items = [
        item
        for item in groups.values()
        if not text or text in item.group_name.lower()
    ]
    return sorted(items, key=lambda item: (not item.is_authorized, item.group_name.lower()))[:limit]


@router.post("/{workspace_id}/members", response_model=MemberResponse)
def upsert_workspace_member(
    workspace_id: int,
    req: UpsertWorkspaceMemberRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    role = _validate_workspace_member_role(req.role)
    target_user = _resolve_member_target_user(db, req)
    member = _workspace_membership(db, target_user.id, workspace_id)
    action = "workspace_member_update" if member else "workspace_member_invite"
    if member:
        member.role = role
    else:
        member = WorkspaceMember(workspace_id=workspace_id, user_id=target_user.id, role=role)
        db.add(member)
    db.flush()
    _write_workspace_audit(
        db,
        user.id,
        action,
        _audit_detail(workspace_id, actor_id=user.id, target_user_id=target_user.id, role=role),
    )
    notify_user(
        db,
        target_user.id,
        category="workspace",
        severity="info",
        title="工作区权限已更新",
        content=f"{workspace.name}：你已被设置为{'工作区管理员' if role == 'admin' else '成员'}。",
        action_status="none",
        action_kind="open_workspace",
        action_payload={"workspace_id": workspace.id},
        event_key=f"workspace:{workspace.id}:member:{target_user.id}:{role}",
    )
    db.commit()
    db.refresh(member)
    db.refresh(target_user)
    return _serialize_workspace_member(member, target_user)


@router.put("/{workspace_id}/members/{target_user_id}", response_model=MemberResponse)
def update_workspace_member_role(
    workspace_id: int,
    target_user_id: int,
    req: UpdateWorkspaceMemberRoleRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    role = _validate_workspace_member_role(req.role)
    member = _workspace_membership(db, target_user_id, workspace_id)
    if not member:
        raise HTTPException(status_code=404, detail="成员不存在")
    if member.role == "admin" and role != "admin" and _local_workspace_admin_count(db, workspace_id) <= 1:
        raise HTTPException(status_code=400, detail="至少需要保留一名工作区管理员")
    target_user = db.query(User).filter(User.id == target_user_id).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="用户不存在")

    member.role = role
    _write_workspace_audit(
        db,
        user.id,
        "workspace_member_role_update",
        _audit_detail(workspace_id, actor_id=user.id, target_user_id=target_user_id, role=role),
    )
    notify_user(
        db,
        target_user.id,
        category="workspace",
        severity="info",
        title="工作区角色已更新",
        content=f"{workspace.name}：你的角色已变更为{'工作区管理员' if role == 'admin' else '成员'}。",
        action_status="none",
        action_kind="open_workspace",
        action_payload={"workspace_id": workspace.id},
        event_key=f"workspace:{workspace.id}:role:{target_user.id}:{role}",
    )
    db.commit()
    db.refresh(member)
    return _serialize_workspace_member(member, target_user)


@router.delete("/{workspace_id}/members/{target_user_id}")
def remove_workspace_member(
    workspace_id: int,
    target_user_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    member = _workspace_membership(db, target_user_id, workspace_id)
    if not member:
        raise HTTPException(status_code=404, detail="成员不存在")
    if member.role == "admin" and _local_workspace_admin_count(db, workspace_id) <= 1:
        raise HTTPException(status_code=400, detail="至少需要保留一名工作区管理员")
    db.delete(member)
    _write_workspace_audit(
        db,
        user.id,
        "workspace_member_remove",
        _audit_detail(workspace_id, actor_id=user.id, target_user_id=target_user_id),
    )
    notify_user(
        db,
        target_user_id,
        category="workspace",
        severity="info",
        title="工作区访问权限已移除",
        content=f"{workspace.name}：你已不再是该工作区成员。",
        action_status="none",
        action_kind="open_workspace",
        action_payload={"workspace_id": workspace.id},
        event_key=f"workspace:{workspace.id}:remove:{target_user_id}",
    )
    db.commit()
    return {"ok": True}


@router.post("/{workspace_id}/groups", response_model=WorkspaceGroupResponse)
def upsert_workspace_group(
    workspace_id: int,
    req: UpsertWorkspaceGroupRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    group_name = _normalize_group_name(req.group_name)
    if not group_name:
        raise HTTPException(status_code=400, detail="组别不能为空")
    existing = (
        db.query(WorkspaceGroupAccess)
        .filter(WorkspaceGroupAccess.workspace_id == workspace_id, WorkspaceGroupAccess.group_name == group_name)
        .first()
    )
    if not existing:
        db.add(WorkspaceGroupAccess(workspace_id=workspace_id, group_name=group_name))
        _write_workspace_audit(
            db,
            user.id,
            "workspace_group_access_add",
            _audit_detail(workspace_id, actor_id=user.id, group_name=group_name),
        )
        db.commit()
    return WorkspaceGroupResponse(group_name=group_name)


@router.delete("/{workspace_id}/groups/{group_name}")
def remove_workspace_group(
    workspace_id: int,
    group_name: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="工作区不存在")
    _ensure_mutable_membership_workspace(workspace)
    _ensure_workspace_admin(db, user, workspace_id)

    normalized = _normalize_group_name(group_name)
    db.query(WorkspaceGroupAccess).filter(
        WorkspaceGroupAccess.workspace_id == workspace_id,
        WorkspaceGroupAccess.group_name == normalized,
    ).delete()
    _write_workspace_audit(
        db,
        user.id,
        "workspace_group_access_remove",
        _audit_detail(workspace_id, actor_id=user.id, group_name=normalized),
    )
    db.commit()
    return {"ok": True}


@router.post("/{workspace_id}/join")
def join_workspace(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")

    existing = (
        db.query(WorkspaceMember)
        .filter(
            WorkspaceMember.workspace_id == workspace_id,
            WorkspaceMember.user_id == user.id,
        )
        .first()
    )
    if existing:
        return {"ok": True, "message": "已是项目成员"}

    if user.role == "admin" and workspace.workspace_kind != "user":
        return {"ok": True, "message": "系统管理员无需加入即可访问"}
    if workspace.workspace_kind == "project" and not workspace.is_hidden:
        return {"ok": True, "message": "开放项目无需加入即可访问"}
    if workspace.workspace_kind in {"project", "customer"} and _has_workspace_group_access(db, user, workspace_id):
        return {"ok": True, "message": "你的组别已获授权访问"}

    raise HTTPException(status_code=403, detail="受限工作区只能通过工作区管理员添加人员或组别授权")


@router.delete("/{workspace_id}")
def delete_workspace(
    workspace_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
    if not workspace:
        raise HTTPException(status_code=404, detail="项目不存在")
    if workspace.is_default:
        raise HTTPException(status_code=400, detail="默认工作区不能删除")

    if not _is_workspace_admin(db, user, workspace_id):
        raise HTTPException(status_code=403, detail="仅项目管理员可删除")

    db.query(WorkspaceMember).filter(
        WorkspaceMember.workspace_id == workspace_id
    ).delete()
    db.query(WorkspaceGroupAccess).filter(
        WorkspaceGroupAccess.workspace_id == workspace_id
    ).delete()
    db.delete(workspace)
    db.commit()
    return {"ok": True}


def _ensure_member(db: Session, user_id: int, workspace_id: int):
    member = _workspace_membership(db, user_id, workspace_id)
    if member:
        return member
    user = db.query(User).filter(User.id == user_id).first()
    if user and user.role == "admin":
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        if workspace and workspace.workspace_kind != "user":
            return WorkspaceMember(workspace_id=workspace_id, user_id=user_id, role="admin")
    if user:
        workspace = db.query(Workspace).filter(Workspace.id == workspace_id).first()
        if workspace and workspace.workspace_kind == "project":
            if not workspace.is_hidden or _has_workspace_group_access(db, user, workspace_id):
                return WorkspaceMember(workspace_id=workspace_id, user_id=user_id, role="member")
        if workspace and workspace.workspace_kind == "customer" and _has_workspace_group_access(db, user, workspace_id):
            return WorkspaceMember(workspace_id=workspace_id, user_id=user_id, role="member")
    if not member:
        raise HTTPException(status_code=403, detail="你尚未加入该工作区")
    return member

